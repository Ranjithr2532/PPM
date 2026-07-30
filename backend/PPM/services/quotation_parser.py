import io
import re
from typing import Dict, Any, Optional, List

try:
    from docx import Document
except ImportError:
    Document = None


def parse_docx_quotation(file_bytes: bytes, filename: str = "") -> Dict[str, Any]:
    """
    Parses a Microsoft Word (.docx) proposal/quotation document and extracts
    key proposal fields to prefill the proposal creation form.
    
    Extracts:
      - enquiry_date / quote_date
      - customer_name
      - address
      - email (Primary email only, excluding CC emails)
      - phone_no
      - alternate_contact_details (Kind Attention / Contact Person)
      - quote_reference
      - quote_description (Subject / Scope summary)
      - quote_amount
      - center / quotation_given_by_department
      - quotation_given_by_name
    """
    if not filename.lower().endswith(".docx") and not file_bytes.startswith(b"PK\x03\x04"):
        raise ValueError("Invalid file format. Only Microsoft Word (.docx) files are supported.")

    if Document is None:
        raise RuntimeError("python-docx package is missing on server.")

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Could not parse document. Please ensure it is a valid .docx file. ({str(e)})")

    # Collect all paragraph text lines and table cell contents
    all_lines: List[str] = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            # Normalize formatting stars and extra spaces
            clean_txt = re.sub(r"\*+", "", txt).strip()
            if clean_txt:
                all_lines.append(clean_txt)

    table_cells: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                ctxt = cell.text.strip()
                if ctxt:
                    table_cells.append(ctxt)

    extracted: Dict[str, Any] = {
        "enquiry_date": "",
        "customer_name": "",
        "address": "",
        "email": "",
        "phone_no": "",
        "alternate_contact_details": "",
        "email_reference": "",
        "quote_reference": "",
        "quote_description": "",
        "quote_amount": "",
        "center": "",
        "quotation_given_by_department": "",
        "payment_terms_and_condition": "",
        "payment_terms_and_conditions": "",
        "payment_terms": "",
        "terms_and_conditions": [],
    }

    # Helper regex patterns
    email_regex = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")
    date_regex = re.compile(
        r"\b(\d{1,2}[./-]\d{1,2}[./-]\d{2,4}|\d{4}[./-]\d{1,2}[./-]\d{1,2}|(?:\d{1,2}\s+)?(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{2,4})\b",
        re.IGNORECASE,
    )
    phone_regex = re.compile(r"\b(\+?\d{1,3}[-.\s]?)?\(?\d{2,5}\)?[-.\s]?\d{3,5}[-.\s]?\d{3,5}\b")
    amount_regex = re.compile(r"\b(?:Rs\.?|INR|Total\s*Cost|Total\s*Amount|Total|Cost|Amount|PRICE)\b[^\d\w]*([\d,]+(?:\.\d{2})?)", re.IGNORECASE)

    # 1. Parse Primary Email
    # Search lines for "Email:" or "Email - To:", strictly ignoring "Cc:" lines
    primary_email = ""
    for line in all_lines:
        line_lower = line.lower()
        if "cc:" in line_lower:
            continue
        if "email" in line_lower or "mail" in line_lower:
            match = email_regex.search(line)
            if match:
                primary_email = match.group(0)
                break
    if not primary_email:
        # General email fallback across lines (excluding cc)
        for line in all_lines:
            if "cc:" not in line.lower():
                match = email_regex.search(line)
                if match:
                    primary_email = match.group(0)
                    break
    extracted["email"] = primary_email

    # 2. Parse Date
    for line in all_lines:
        line_lower = line.lower()
        if any(kw in line_lower for kw in ["date:", "enquiry date:", "quote date:", "dated", "date"]):
            if ":" in line:
                val_after_colon = line.split(":", 1)[1].strip()
                match = date_regex.search(val_after_colon)
                if match:
                    extracted["enquiry_date"] = match.group(0)
                    extracted["quote_date"] = match.group(0)
                    break
                elif val_after_colon:
                    extracted["enquiry_date"] = val_after_colon
                    extracted["quote_date"] = val_after_colon
                    break
            match = date_regex.search(line)
            if match:
                extracted["enquiry_date"] = match.group(0)
                extracted["quote_date"] = match.group(0)
                break
    if not extracted.get("enquiry_date") and all_lines:
        for line in all_lines[:5]:
            match = date_regex.search(line)
            if match:
                extracted["enquiry_date"] = match.group(0)
                extracted["quote_date"] = match.group(0)
                break

    # 3. Parse Department / Center
    for line in all_lines:
        if "dept:" in line.lower() or "department:" in line.lower() or "center:" in line.lower():
            parts = re.split(r":", line, maxsplit=1)
            if len(parts) > 1 and parts[1].strip():
                dept_val = parts[1].strip()
                extracted["center"] = dept_val
                extracted["quotation_given_by_department"] = dept_val
                break

    # 4. Parse Customer & Address
    # In quotation_generator format, "Customer:" is followed by Line 1 (Name) and subsequent lines (Address)
    customer_name = ""
    address_lines = []
    in_customer_block = False

    for idx, line in enumerate(all_lines):
        line_clean = line.strip()
        if re.match(r"^Customer\s*:", line_clean, re.IGNORECASE) or re.match(r"^Customer Name\s*:", line_clean, re.IGNORECASE) or re.match(r"^M/s\.?", line_clean, re.IGNORECASE):
            parts = line_clean.split(":", 1)
            first_val = parts[1].strip() if len(parts) > 1 else line_clean
            if first_val:
                customer_name = first_val
            in_customer_block = True
            continue

        if in_customer_block:
            # Stop customer block when reaching next section header
            if any(line_clean.lower().startswith(stop_kw) for stop_kw in ["kind attention", "attention", "reference", "subject", "sac code", "email", "scope of work", "terms"]):
                in_customer_block = False
            else:
                address_lines.append(line_clean)

    if not customer_name:
        for line in all_lines:
            if any(kw in line.lower() for kw in ["customer name:", "company name:", "client name:", "client:"]):
                parts = line.split(":", 1)
                if len(parts) > 1 and parts[1].strip():
                    customer_name = parts[1].strip()
                    break

    if not customer_name:
        for idx, line in enumerate(all_lines):
            line_clean = line.strip()
            if any(line_clean.lower().startswith(kw) for kw in ["kind attn", "kind attention", "sub:", "subject:", "ref:", "dear sir"]):
                cust_candidates = []
                for prev_line in reversed(all_lines[:idx]):
                    p_clean = prev_line.strip()
                    if any(p_clean.lower().startswith(h) for h in ["ref no", "date:", "quotation", "iso 9001", "dept:"]):
                        break
                    cust_candidates.insert(0, p_clean)
                if cust_candidates:
                    customer_name = cust_candidates[0]
                    address_lines = cust_candidates[1:]
                break

    extracted["customer_name"] = customer_name
    extracted["address"] = ", ".join(address_lines) if address_lines else ""

    # 5. Parse Kind Attention / Contact Person
    for line in all_lines:
        if any(kw in line.lower() for kw in ["kind attention:", "kind attn:", "kind attention", "kind attn", "attention:", "contact person:", "contact:"]):
            parts = line.split(":", 1)
            if len(parts) > 1 and parts[1].strip():
                extracted["alternate_contact_details"] = parts[1].strip()
                break

    # 6. Parse Phone Number
    for line in all_lines:
        if any(kw in line.lower() for kw in ["phone:", "phone no:", "mobile:", "tel:"]):
            match = phone_regex.search(line)
            if match:
                extracted["phone_no"] = match.group(0).strip()
                break

    # 7. Parse Reference (Populates email_reference)
    for line in all_lines:
        if any(kw in line.lower() for kw in ["reference:", "ref:", "email reference:", "proposal no:", "quote ref:"]):
            parts = line.split(":", 1)
            if len(parts) > 1 and parts[1].strip():
                ref_val = parts[1].strip()
                extracted["email_reference"] = ref_val
                extracted["quote_reference"] = ref_val
                break

    # 8. Parse Subject / Scope Summary -> quote_description
    subject_val = ""
    for line in all_lines:
        if any(kw in line.lower() for kw in ["subject:", "sub:", "product details:", "scope of work:"]):
            parts = line.split(":", 1)
            if len(parts) > 1 and parts[1].strip():
                subject_val = parts[1].strip()
                break
    extracted["quote_description"] = subject_val

    # 9. Parse Quote Amount
    amount_found = ""
    # Search paragraphs & tables for monetary amounts or numbers in cost tables
    combined_texts = all_lines + table_cells
    for txt in combined_texts:
        # Match explicit currency patterns first
        match = amount_regex.search(txt)
        if match:
            num_str = match.group(1).replace(",", "").strip()
            try:
                val = float(num_str)
                if val > 0:
                    amount_found = str(int(val)) if val.is_integer() else str(val)
                    break
            except ValueError:
                pass
    if not amount_found:
        # Fallback: find pure numerical values in table cells (excluding small sl_nos and years)
        for cell_txt in table_cells:
            cell_clean = cell_txt.replace(",", "").strip()
            if cell_clean.isdigit():
                val = int(cell_clean)
                if val >= 1000 and not (1900 <= val <= 2099):  # Exclude small item indices and years
                    amount_found = str(val)
                    break
    extracted["quote_amount"] = amount_found

    # 10. Parse Terms & Conditions / Payment Terms and Conditions
    terms_items: List[str] = []
    payment_terms_val = ""
    in_terms_section = False

    for idx, line in enumerate(all_lines):
        line_clean = line.strip()
        line_lower = line_clean.lower()

        if any(h in line_lower for h in ["payment terms", "terms and conditions", "terms & conditions"]):
            in_terms_section = True
            if ":" in line_clean:
                after_colon = line_clean.split(":", 1)[1].strip()
                if after_colon:
                    payment_terms_val = after_colon
                    terms_items.append(after_colon)
            continue

        if in_terms_section:
            is_stop_kw = any(stop_kw in line_lower for stop_kw in [
                "bank details:", "gst charges:", "governing law", "jurisdiction", "our pan",
                "thanking you", "yours sincerely", "for central manufacturing", "scientist",
                "project planning", "centre head", "group head", "director", "cmti"
            ])
            is_name_line = bool(re.match(r"^[A-Z\s\.]{3,},?$", line_clean) and not any(kw in line_lower for kw in ["terms", "conditions", "quotation", "payment", "scope", "delivery"]))

            if is_stop_kw or is_name_line:
                in_terms_section = False
            elif line_clean:
                clean_item = re.sub(r"^[\s•\-\*\d\.]+\s*", "", line_clean).strip()
                if clean_item and not clean_item.lower().startswith("with reference to"):
                    terms_items.append(clean_item)
                    if not payment_terms_val:
                        payment_terms_val = clean_item

    extracted_terms = "; ".join(terms_items) if terms_items else payment_terms_val
    extracted["payment_terms_and_condition"] = extracted_terms
    extracted["payment_terms_and_conditions"] = extracted_terms
    extracted["payment_terms"] = extracted_terms or payment_terms_val
    extracted["terms_and_conditions"] = terms_items

    # 11. Parse Signatory Name (e.g. NARENDRA REDDY T,)
    for idx, line in enumerate(all_lines):
        line_clean = line.strip()
        line_lower = line_clean.lower()

        if any(close_kw in line_lower for close_kw in ["thanking you", "yours sincerely", "for central manufacturing"]):
            for next_line in all_lines[idx + 1:idx + 6]:
                nl_clean = next_line.strip().rstrip(",")
                nl_lower = nl_clean.lower()
                if nl_clean and not any(skip in nl_lower for skip in ["thanking you", "yours sincerely", "for central manufacturing", "dept:", "center:", "date:"]):
                    if re.match(r"^[A-Za-z\s\.]+$", nl_clean) and len(nl_clean) >= 3:
                        extracted["quotation_given_by_name"] = nl_clean
                        extracted["signatory_name"] = nl_clean
                        break
        elif re.match(r"^[A-Z\s\.]{3,},?$", line_clean) and not any(kw in line_lower for kw in ["quotation", "proposal", "terms", "conditions", "payment", "delivery", "scope", "dept", "center", "iso"]):
            clean_name = line_clean.rstrip(",")
            if len(clean_name) >= 3 and not extracted.get("quotation_given_by_name"):
                extracted["quotation_given_by_name"] = clean_name
                extracted["signatory_name"] = clean_name

    # 11. Document Validity Check
    full_doc_text = " ".join(all_lines).lower()
    has_proposal_keywords = any(
        kw in full_doc_text for kw in ["quotation", "proposal", "enquiry date:", "subject:", "reference:", "ref:", "kind attention", "terms & conditions", "terms and conditions"]
    )
    has_key_extracted_fields = any([
        extracted.get("customer_name"),
        extracted.get("email_reference"),
        extracted.get("quote_reference"),
        extracted.get("quote_description"),
    ])

    if not (has_proposal_keywords or has_key_extracted_fields):
        raise ValueError("Invalid proposal document format. Please upload a valid document generated by the Word Document Generator.")

    return extracted


def parse_cmti_quotation_docx(file_bytes: bytes, filename: str = "") -> Dict[str, Any]:
    """
    Parses an uploaded .docx proposal/quotation document specifically to extract
    all details for the CMTI Quotation Generator template.
    """
    if not filename.lower().endswith(".docx") and not file_bytes.startswith(b"PK\x03\x04"):
        raise ValueError("Invalid file format. Only Microsoft Word (.docx) files are supported.")

    if Document is None:
        raise RuntimeError("python-docx package is missing on server.")

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Could not parse document. Please ensure it is a valid .docx file. ({str(e)})")

    all_lines: List[str] = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            clean_txt = re.sub(r"\*+", "", txt).strip()
            if clean_txt:
                for line in clean_txt.splitlines():
                    if line.strip():
                        all_lines.append(line.strip())

    table_cells: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                ctxt = cell.text.strip()
                if ctxt:
                    table_cells.append(ctxt)

    extracted: Dict[str, Any] = {
        "header_code": "ISO 9001-2015 CMTI/PPBD/001/Rev-00",
        "ref_no": "",
        "date": "",
        "customer_lines": [],
        "kind_attention": "",
        "salutation": "Dear Sir,",
        "subject": "",
        "email_ref": "",
        "item_description": "",
        "quote_amount": "",
        "scope_of_work": [],
        "validity": "",
        "payment_terms": "100% after completion of work & submission of report.",
        "delivery": "1 Month from the date of acceptance of PO.",
        "contact_details": "",
        "commercial_contact": "Mrs.Kusuma A, OS(Gr-I) (PP&BD), (Contact No.: 080-22188233 / 242).",
        "signatory_name": "",
        "signatory_designation": "Scientist 'F' & Centre Head, Project Planning and Business Development.",
    }

    # 1. Ref No & Date
    for line in all_lines:
        line_clean = line.strip()
        if re.search(r"ref\.?\s*no\.?:?", line_clean, re.IGNORECASE):
            parts = line_clean.split(":", 1)
            if len(parts) > 1 and parts[1].strip():
                extracted["ref_no"] = parts[1].strip()
        elif re.search(r"date\s*:", line_clean, re.IGNORECASE):
            parts = line_clean.split(":", 1)
            if len(parts) > 1 and parts[1].strip():
                extracted["date"] = parts[1].strip()

    # 2. Customer Lines & Kind Attention
    cust_lines = []
    in_cust = False
    for line in all_lines:
        line_clean = line.strip()
        if line_clean.startswith("M/s.") or line_clean.startswith("Customer:") or line_clean.startswith("To,"):
            in_cust = True
            cust_lines.append(line_clean)
            continue
        if in_cust:
            if any(kw in line_clean.lower() for kw in ["kind attn", "dear sir", "sub:", "subject:", "ref:"]):
                in_cust = False
            else:
                cust_lines.append(line_clean)

        if "kind attn:" in line_clean.lower() or "kind attention:" in line_clean.lower():
            parts = line_clean.split(":", 1)
            if len(parts) > 1:
                extracted["kind_attention"] = parts[1].strip()

    extracted["customer_lines"] = cust_lines

    # 3. Subject & Email Ref
    for line in all_lines:
        line_clean = line.strip()
        if line_clean.lower().startswith("sub:") or line_clean.lower().startswith("subject:"):
            parts = line_clean.split(":", 1)
            if len(parts) > 1:
                extracted["subject"] = parts[1].strip()
                m = re.search(r"Quotation for [“\"']?(.*?)[”\"']?\.?$", parts[1].strip(), re.IGNORECASE)
                if m and m.group(1):
                    extracted["item_description"] = m.group(1).strip()
        elif line_clean.lower().startswith("ref:") and "email" in line_clean.lower():
            parts = line_clean.split(":", 1)
            if len(parts) > 1:
                extracted["email_ref"] = parts[1].strip()

    if not extracted["item_description"] and extracted["subject"]:
        extracted["item_description"] = extracted["subject"].replace("Quotation for", "").strip(' "“’')

    # 4. Scope of Work, Terms and Conditions, & Payment Terms
    scope_stop_pattern = re.compile(
        r"^(?:\d+\.\s*)?(?:terms\s*(?:and|&)\s*conditions|quotation\s*validity|payment\s*terms|delivery|gst|bank\s*details|cost\s*break-?up|cost\s*estimation|governing\s*law|jurisdiction|our\s*pan|thanking\s*you|yours\s*sincerely|for\s*central\s*manufacturing|scientist|special\s*manufacturing|cmti,)",
        re.IGNORECASE
    )

    in_scope = False
    in_terms = False
    scope_items = []
    terms_items = []

    for line in all_lines:
        line_clean = line.strip()
        line_lower = line_clean.lower()

        # Check start of Scope of Work
        if "scope of work" in line_lower or "scope:" in line_lower:
            in_scope = True
            in_terms = False
            continue

        # Check start of Terms and Conditions
        if "terms and conditions" in line_lower or "terms & conditions" in line_lower:
            in_scope = False
            in_terms = True
            continue

        # Check section stop keywords for Scope
        if in_scope and scope_stop_pattern.search(line_clean):
            in_scope = False
            if "terms and conditions" in line_lower or "terms & conditions" in line_lower:
                in_terms = True

        if in_scope:
            if line_clean and not line_lower.startswith("with reference to"):
                # Clean leading bullet symbols or numbers (e.g. •, -, *, 1.)
                clean_item = re.sub(r"^[\s•\-\*\d\.]+\s*", "", line_clean).strip()
                if clean_item:
                    scope_items.append(clean_item)

        if in_terms:
            is_stop_kw = any(stop_kw in line_lower for stop_kw in [
                "cost break-up", "cost breakup", "scientist", "cmti,", "thanking you",
                "yours sincerely", "for central manufacturing", "bank details:", "gst charges:",
                "governing law", "jurisdiction", "our pan"
            ])
            is_name_line = bool(re.match(r"^[A-Z\s\.]{3,},?$", line_clean) and not any(kw in line_lower for kw in ["terms", "conditions", "quotation", "payment", "scope", "delivery"]))

            if is_stop_kw or is_name_line:
                in_terms = False
            elif line_clean:
                clean_term = re.sub(r"^[\s•\-\*\d\.]+\s*", "", line_clean).strip()
                if clean_term:
                    terms_items.append(clean_term)

        # Separate extraction for validity, payment, delivery
        if "valid till" in line_lower or "validity:" in line_lower:
            extracted["validity"] = line_clean
        elif "payment" in line_lower and ("terms" in line_lower or "%" in line_lower or "delivery" in line_lower):
            if not extracted["payment_terms"]:
                parts = line_clean.split(":", 1)
                extracted["payment_terms"] = parts[1].strip() if len(parts) > 1 else line_clean
        elif "delivery:" in line_lower or ("month" in line_lower and "po" in line_lower):
            if not extracted["delivery"]:
                parts = line_clean.split(":", 1)
                extracted["delivery"] = parts[1].strip() if len(parts) > 1 else line_clean

    extracted["scope_of_work"] = scope_items

    if terms_items:
        extracted["payment_terms"] = "; ".join(terms_items)
        extracted["payment_terms_and_condition"] = "; ".join(terms_items)
        extracted["payment_terms_and_conditions"] = "; ".join(terms_items)
    else:
        extracted["payment_terms_and_condition"] = extracted.get("payment_terms", "")
        extracted["payment_terms_and_conditions"] = extracted.get("payment_terms", "")

    extracted["terms_and_conditions"] = terms_items

    # 5. Contact Details & Signatory
    for idx, line in enumerate(all_lines):
        line_clean = line.strip()
        line_lower = line_clean.lower()
        if "please feel free to contact" in line_lower:
            extracted["contact_details"] = line_clean

        if any(close_kw in line_lower for close_kw in ["thanking you", "yours sincerely", "for central manufacturing"]):
            for next_line in all_lines[idx + 1:idx + 6]:
                nl_clean = next_line.strip().rstrip(",")
                nl_lower = nl_clean.lower()
                if nl_clean and not any(skip in nl_lower for skip in ["thanking you", "yours sincerely", "for central manufacturing", "dept:", "center:", "date:"]):
                    if re.match(r"^[A-Za-z\s\.]+$", nl_clean) and len(nl_clean) >= 3:
                        extracted["signatory_name"] = nl_clean
                        break
        elif re.match(r"^[A-Z\s\.]{3,},?$", line_clean) and not any(kw in line_lower for kw in ["quotation", "proposal", "terms", "conditions", "payment", "delivery", "scope", "dept", "center", "iso"]):
            clean_name = line_clean.rstrip(",")
            if len(clean_name) >= 3 and not extracted.get("signatory_name"):
                extracted["signatory_name"] = clean_name

    return extracted
