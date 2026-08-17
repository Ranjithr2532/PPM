import io
import re
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Mm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    raise ImportError("python-docx package is required for document generation.")

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def set_a4_page(document):
    sec = document.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_rich_text(paragraph, text):
    pos = 0
    for match in BOLD_PATTERN.finditer(text):
        start, end = match.span()
        if start > pos:
            paragraph.add_run(text[pos:start])
        r = paragraph.add_run(match.group(1))
        r.bold = True
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_section_heading(document, title, center=False):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Calibri"
    p.paragraph_format.keep_with_next = True
    return p


def add_email_field(document, label, addresses):
    if not addresses:
        return
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

    r_lbl = p.add_run(f"{label}: ")
    r_lbl.bold = True
    r_lbl.font.size = Pt(11)

    for i, addr in enumerate(addresses):
        if i > 0:
            p.add_run(", ")
        add_hyperlink(p, f"mailto:{addr}", addr)


def add_bullets(document, items):
    for item in items:
        item = item.strip()
        if not item:
            continue
        clean_item = re.sub(r"^[\s•\-\*–]+", "", item).strip()
        p = document.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_rich_text(p, clean_item)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_table_block(document, table_spec):
    title = table_spec.get("title", "")
    headers = table_spec.get("headers", [])
    rows = table_spec.get("rows", [])
    if title:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(title).bold = True

    if not headers:
        return

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        shade_cell(hdr_cells[i], "D9E2F3")

    for row_vals in rows:
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            if isinstance(row_vals, dict):
                raw_v = row_vals.get(h, "")
                if isinstance(raw_v, (int, float)):
                    cells[i].text = f"₹ {raw_v:,.2f}"
                else:
                    cells[i].text = str(raw_v if raw_v is not None else "")
            elif isinstance(row_vals, list):
                raw_v = row_vals[i] if i < len(row_vals) else ""
                if isinstance(raw_v, (int, float)):
                    cells[i].text = f"₹ {raw_v:,.2f}"
                else:
                    cells[i].text = str(raw_v if raw_v is not None else "")
            else:
                cells[i].text = str(row_vals or "")

    document.add_paragraph()


def clear_table_borders(table):
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for child in list(tcPr):
                if child.tag.endswith('tcBorders'):
                    tcPr.remove(child)

            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_signature_block(document, name=None, designation_lines=None, signatories=None):
    valid_signatories = []
    if signatories and len(signatories) > 0:
        for sig in signatories:
            if isinstance(sig, dict):
                s_name = (sig.get("name") or "").strip()
                raw_lines = sig.get("lines") or []
                s_lines = [l.strip() for l in raw_lines if l.strip()]
            elif hasattr(sig, "dict"):
                s_dict = sig.dict()
                s_name = (s_dict.get("name") or "").strip()
                s_lines = [l.strip() for l in (s_dict.get("lines") or []) if l.strip()]
            elif hasattr(sig, "model_dump"):
                s_dict = sig.model_dump()
                s_name = (s_dict.get("name") or "").strip()
                s_lines = [l.strip() for l in (s_dict.get("lines") or []) if l.strip()]
            else:
                continue
            if s_name or s_lines:
                valid_signatories.append({"name": s_name, "lines": s_lines})

    if not valid_signatories:
        s_name = (name or "").strip()
        s_lines = [l.strip() for l in (designation_lines or []) if l.strip()]
        if s_name or s_lines:
            valid_signatories.append({"name": s_name, "lines": s_lines})

    if not valid_signatories:
        return

    document.add_paragraph()

    if len(valid_signatories) == 1:
        sig = valid_signatories[0]
        if sig["name"]:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(sig["name"] + ",")
            run.bold = True
        for line in sig["lines"]:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(line)
    else:
        num_cols = 2
        num_rows = (len(valid_signatories) + 1) // 2
        table = document.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Normal Table'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        clear_table_borders(table)

        for idx, sig in enumerate(valid_signatories):
            r_idx = idx // num_cols
            c_idx = idx % num_cols
            cell = table.rows[r_idx].cells[c_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if sig["name"]:
                run = p.add_run(sig["name"] + ",\n")
                run.bold = True
            for l_idx, line in enumerate(sig["lines"]):
                is_last = (l_idx == len(sig["lines"]) - 1)
                p.add_run(line + ("\n" if not is_last else ""))


def generate_proposal_docx(data: dict) -> io.BytesIO:
    document = build_proposal_document(data)
    target = io.BytesIO()
    document.save(target)
    target.seek(0)
    return target


def build_proposal_document(data: dict) -> Document:
    if not isinstance(data, dict) and hasattr(data, "dict"):
        data = data.dict()
    elif not isinstance(data, dict) and hasattr(data, "model_dump"):
        data = data.model_dump()

    header_code = data.get("header_code") or "ISO 9001-2015 CMTI/PPBD/001/Rev-00"
    date = data.get("date") or datetime.now().strftime("%d/%m/%Y")
    dept = data.get("dept") or ""
    email_to = data.get("email_to") or []
    email_cc = data.get("email_cc") or []
    customer_lines = data.get("customer_lines") or []
    kind_attention = data.get("kind_attention") or ""
    reference = data.get("reference") or data.get("email_ref") or ""
    subject = data.get("subject") or ""
    sac_code = data.get("sac_code") or ""
    scope_intro = data.get("scope_intro") or ""
    scope_items = data.get("scope_items") or data.get("scope_of_work") or []
    terms_items = data.get("terms_items") or []
    tables = data.get("tables") or []
    internal_cost_tables = data.get("internal_cost_tables") or []
    signatories = data.get("signatories") or []
    signatory_name = data.get("signatory_name")
    signatory_lines = data.get("signatory_lines")

    document = Document()
    set_a4_page(document)

    # 1. Header Code in Section Header (Appears on every page's top-right corner)
    section = document.sections[0]
    header = section.header
    p_hdr = header.paragraphs[0]
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_hdr.paragraph_format.space_after = Pt(0)
    r_hdr = p_hdr.add_run(header_code)
    r_hdr.font.name = "Calibri"
    r_hdr.font.size = Pt(9)
    r_hdr.font.bold = True

    # Date & Dept Block
    p_date = document.add_paragraph()
    p_date.paragraph_format.space_after = Pt(2)
    r_date = p_date.add_run(f"Date: {date}")
    r_date.font.size = Pt(11)

    if dept:
        p_dept = document.add_paragraph()
        p_dept.paragraph_format.space_after = Pt(6)
        r_dept = p_dept.add_run(f"Dept: {dept}")
        r_dept.font.size = Pt(11)

    if email_to:
        add_email_field(document, "Email", email_to)
    if email_cc:
        add_email_field(document, "Cc", email_cc)

    # Quotation Information Block
    p_qinfo = document.add_paragraph()
    p_qinfo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_qinfo.paragraph_format.space_before = Pt(8)
    p_qinfo.paragraph_format.space_after = Pt(4)
    r_qinfo = p_qinfo.add_run("Quotation Information:")
    r_qinfo.font.bold = True
    r_qinfo.font.size = Pt(11)

    if customer_lines:
        p_cust_lbl = document.add_paragraph()
        p_cust_lbl.paragraph_format.space_after = Pt(2)
        r_cust_lbl = p_cust_lbl.add_run("Customer: ")
        r_cust_lbl.font.bold = True
        r_cust_lbl.font.size = Pt(11)
        for idx, cline in enumerate(customer_lines):
            if idx == 0:
                r_c = p_cust_lbl.add_run(cline)
                r_c.font.size = Pt(11)
            else:
                p_c = document.add_paragraph()
                p_c.paragraph_format.space_after = Pt(2)
                r_c = p_c.add_run(cline)
                r_c.font.size = Pt(11)

    if kind_attention:
        p_ka = document.add_paragraph()
        p_ka.paragraph_format.space_after = Pt(2)
        r_ka_lbl = p_ka.add_run("Kind Attention: ")
        r_ka_lbl.font.bold = True
        r_ka_lbl.font.size = Pt(11)
        r_ka = p_ka.add_run(kind_attention)
        r_ka.font.size = Pt(11)

    if reference:
        p_ref = document.add_paragraph()
        p_ref.paragraph_format.space_after = Pt(2)
        r_ref_lbl = p_ref.add_run("Reference: ")
        r_ref_lbl.font.bold = True
        r_ref_lbl.font.size = Pt(11)
        r_ref = p_ref.add_run(reference)
        r_ref.font.size = Pt(11)

    if subject:
        p_sub = document.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(2)
        r_sub_lbl = p_sub.add_run("Subject: ")
        r_sub_lbl.font.bold = True
        r_sub_lbl.font.size = Pt(11)
        r_sub = p_sub.add_run(subject)
        r_sub.font.size = Pt(11)

    if sac_code:
        p_sac = document.add_paragraph()
        p_sac.paragraph_format.space_after = Pt(6)
        r_sac_lbl = p_sac.add_run("SAC Code: ")
        r_sac_lbl.font.bold = True
        r_sac_lbl.font.size = Pt(11)
        r_sac = p_sac.add_run(sac_code)
        r_sac.font.size = Pt(11)

    if scope_items or scope_intro:
        p_sc_h = document.add_paragraph()
        p_sc_h.paragraph_format.space_before = Pt(8)
        p_sc_h.paragraph_format.space_after = Pt(4)
        r_sc_h = p_sc_h.add_run("Scope of work:")
        r_sc_h.font.bold = True
        r_sc_h.font.size = Pt(11)
        if scope_intro:
            p_si = document.add_paragraph()
            p_si.paragraph_format.space_after = Pt(4)
            p_si.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_si.add_run(scope_intro)
        if scope_items:
            add_bullets(document, scope_items)

    if terms_items:
        p_tm_h = document.add_paragraph()
        p_tm_h.paragraph_format.space_before = Pt(8)
        p_tm_h.paragraph_format.space_after = Pt(4)
        r_tm_h = p_tm_h.add_run("Terms and conditions:")
        r_tm_h.font.bold = True
        r_tm_h.font.size = Pt(11)
        add_bullets(document, terms_items)

    for t in tables:
        if isinstance(t, dict):
            add_table_block(document, t)
        elif hasattr(t, "dict"):
            add_table_block(document, t.dict())
        elif hasattr(t, "model_dump"):
            add_table_block(document, t.model_dump())

    # Internal Cost Estimation section generated on a NEW PAGE before signature block
    if internal_cost_tables:
        document.add_page_break()
        add_section_heading(document, "INTERNAL COST ESTIMATION", center=True)
        for t in internal_cost_tables:
            if isinstance(t, dict):
                add_table_block(document, t)
            elif hasattr(t, "dict"):
                add_table_block(document, t.dict())
            elif hasattr(t, "model_dump"):
                add_table_block(document, t.model_dump())

    add_signature_block(document, name=signatory_name, designation_lines=signatory_lines, signatories=signatories)

    return document
