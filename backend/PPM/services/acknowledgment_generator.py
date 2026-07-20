"""
Recreates "design_and_devpt_of_hydrostatic__1_.docx" using python-docx.

Matches the original file's:
  - Page size (A4) and margins
  - Font (Garamond, 12pt / 13pt as used in the original)
  - Bold placement
  - Paragraph alignment (left / right / justify)
  - Hanging indents on the "Sub:" and "Ref:" lines
  - Line spacing per paragraph

Install dependency first:
    pip install python-docx
"""

from docx import Document
from docx.shared import Emu, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime, date

FONT_NAME = "Garamond"
SIZE_13 = Pt(13)
SIZE_12 = Pt(12)


def set_run_font(run, font_name=FONT_NAME, size=SIZE_13, bold=False):
    """Apply font name/size/bold to a run, including the East Asian /
    complex-script font slots so Word doesn't silently fall back."""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def add_run(paragraph, text, size=SIZE_13, bold=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return run


def blank_paragraph(doc, line_spacing=1.0, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    return p

DATE_FMT = "%d-%m-%Y"


def format_date(value):
    """Accepts a datetime/date object or a string, returns DD-MM-YYYY."""
    if isinstance(value, (datetime, date)):
        return value.strftime(DATE_FMT)
    if isinstance(value, str):
        # try to parse common incoming formats, adjust as needed
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y"):
            try:
                return datetime.strptime(value, fmt).strftime(DATE_FMT)
            except ValueError:
                continue
        # already in desired format or unparseable — return as-is
        return value
    raise TypeError(f"Unsupported date type: {type(value)}")


def build_document(
    ref_no: str,
    current_date: str,
    company_name: str,
    company_address: str,
    kind_attn: str,
    subject: str,
    purchase_order_no: str,
    purchase_order_date: str,
):
    doc = Document()
    
    current_date = format_date(current_date)
    purchase_order_date = format_date(purchase_order_date)

    # ---- Page size (A4) and margins, matching the original ----
    section = doc.sections[0]
    section.page_width = Emu(7560310)
    section.page_height = Emu(10692130)
    section.left_margin = Emu(914400)    # 1"
    section.right_margin = Emu(661035)   # ~0.72"
    section.top_margin = Emu(180340)     # ~0.20"
    section.bottom_margin = Emu(90170)   # ~0.10"

    # ---- Base "Normal" style ----
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = SIZE_13

    # Four blank lines at top (space reserved for letterhead), 2x line spacing
    for _ in range(4):
        blank_paragraph(doc, line_spacing=2.0)

    # Ref No. line — right aligned
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    add_run(p, f"Ref No.: {ref_no}")


    # Date line — right aligned
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    add_run(p, f"Date:{current_date}")

    blank_paragraph(doc)

    # Recipient block
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "M/s. ", size=SIZE_12, bold=True)
    add_run(p, company_name, size=SIZE_12, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, company_address, size=SIZE_12)

    blank_paragraph(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, f"Kind Attn.: {kind_attn}")

    blank_paragraph(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Dear Sir,")

    blank_paragraph(doc)

    # Subject line — bold, hanging indent
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Emu(800100)
    p.paragraph_format.first_line_indent = Emu(-342900)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Sub: ", bold=True)
    add_run(
        p,
        subject,
        size=SIZE_12,
        bold=True,
    )

    blank_paragraph(doc)

    # Ref line — bold "Ref:", hanging indent, 1.5 line spacing
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Emu(2250440)
    p.paragraph_format.first_line_indent = Emu(-1800225)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Ref: ", bold=True)
    add_run(p, " Purchase order No.")
    add_run(
        p,
        f"{purchase_order_no} dated {purchase_order_date}",
        size=SIZE_12,
    )

    blank_paragraph(doc, line_spacing=1.5)

    # Body paragraph — justified
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    add_run(
        p,
        "We acknowledge with thanks for the above Purchase order. We "
        "hereby accept the above Purchase order as per its terms and "
        "conditions",
    )

    blank_paragraph(doc)
    blank_paragraph(doc)
    blank_paragraph(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Thanking You,")

    blank_paragraph(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Yours faithfully")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "For Central Manufacturing Technology Institute")

    blank_paragraph(doc)
    blank_paragraph(doc)
    blank_paragraph(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "(Krishna Rathod)")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Scientist \u2018F\u2019 & Centre Head")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Project Planning and Business Development")

    blank_paragraph(doc, line_spacing=2.5)
    blank_paragraph(doc, line_spacing=2.0)

    return doc