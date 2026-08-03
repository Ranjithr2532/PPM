

import io
import re
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Mm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    sys.exit(
        "This script needs the 'python-docx' package.\n"
        "Install it first with:\n\n    pip install python-docx\n"
    )

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


# ============================================================================
# 1. INPUT HELPERS  (everything that talks to the user)
# ============================================================================

def ask(label, default="", required=False):
    """Single-line text prompt. Returns default if the user just hits Enter."""
    suffix = f" [{default}]" if default else ""
    while True:
        value = input(f"{label}{suffix}: ").strip()
        if not value:
            value = default
        if value or not required:
            return value
        print("   -> This field is required, please enter a value.")


def ask_yes_no(label, default=False):
    """y/n prompt. Returns a bool."""
    suffix = " [Y/n]" if default else " [y/N]"
    value = input(f"{label}{suffix}: ").strip().lower()
    if not value:
        return default
    return value.startswith("y")


def ask_list(label, hint="one point per line, blank line to finish"):
    """Collects multiple lines until the user enters a blank line."""
    print(f"{label}  ({hint})")
    items = []
    while True:
        line = input("   > ").strip()
        if not line:
            break
        items.append(line)
    return items


def ask_int(label, default=None, min_value=1):
    suffix = f" [{default}]" if default is not None else ""
    while True:
        value = input(f"{label}{suffix}: ").strip()
        if not value and default is not None:
            return default
        try:
            n = int(value)
            if n >= min_value:
                return n
        except ValueError:
            pass
        print(f"   -> Please enter a whole number >= {min_value}.")


def section(title):
    print("\n" + "-" * 60)
    print(title)
    print("-" * 60)


# ============================================================================
# 2. DOCX BUILDING HELPERS  (everything that talks to python-docx)
# ============================================================================

def set_a4_page(document):
    sec = document.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)


def add_hyperlink(paragraph, url, text):
    """Insert a real, clickable hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_rich_text(paragraph, text, bold_all=False):
    """
    Add text to a paragraph, honouring **bold** markers so users can
    emphasise figures (e.g. **Rs. 9,20,000**) the same way the source
    template does.
    """
    pos = 0
    for m in BOLD_PATTERN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()]).bold = bold_all
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:]).bold = bold_all


def add_right_aligned(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(text)
    return p


def add_section_heading(document, text, center=False, space_before=12):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text).bold = True
    return p


def add_field(document, label, value):
    """One 'Label: value' paragraph, e.g. Kind Attention: Mr Arun."""
    value = (value or "").strip()
    if not value:
        return
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(f"{label}: ").bold = True
    add_rich_text(p, value)


def add_multiline_field(document, label, lines):
    """
    'Label:' followed by one or more lines of text (e.g. a name +
    address block for Customer). Continuation lines are indented so
    they sit under the first line of text.
    """
    lines = [l.strip() for l in lines if l.strip()]
    if not lines:
        return
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(f"{label}: ").bold = True
    add_rich_text(p, lines[0])
    for extra in lines[1:]:
        p2 = document.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.left_indent = Inches(0.95)
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_rich_text(p2, extra)


def add_email_field(document, label, addresses):
    addresses = [a.strip() for a in addresses if a.strip()]
    if not addresses:
        return
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(f"{label}: ").bold = True
    for i, addr in enumerate(addresses):
        if i > 0:
            p.add_run(", ")
        add_hyperlink(p, f"mailto:{addr}", addr)


def add_bullets(document, items):
    for item in items:
        item = item.strip()
        if not item:
            continue
        clean_item = re.sub(r"^[\s•\-\*–]+", "", item).strip()
        p = document.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_rich_text(p, clean_item)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_table_block(document, table_spec):
    title, headers, rows = table_spec["title"], table_spec["headers"], table_spec["rows"]
    if title:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(title).bold = True

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        shade_cell(hdr_cells[i], "D9E2F3")

    for row_vals in rows:
        cells = table.add_row().cells
        for i in range(len(headers)):
            cells[i].text = row_vals[i] if i < len(row_vals) else ""

    document.add_paragraph()  # small spacer after the table


def clear_table_borders(table):
    """Remove all table and cell borders completely to prevent dotted lines in Word."""
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for child in list(tcPr):
                if child.tag.endswith('tcBorders'):
                    tcPr.remove(child)

            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_signature_block(document, name=None, designation_lines=None, signatories=None):
    valid_signatories = []
    if signatories and len(signatories) > 0:
        for sig in signatories:
            if isinstance(sig, dict):
                s_name = (sig.get("name") or "").strip()
                raw_lines = sig.get("lines") or []
                s_lines = [l.strip() for l in raw_lines if l.strip()]
            elif hasattr(sig, "dict"):
                s_dict = sig.dict()
                s_name = (s_dict.get("name") or "").strip()
                s_lines = [l.strip() for l in (s_dict.get("lines") or []) if l.strip()]
            elif hasattr(sig, "model_dump"):
                s_dict = sig.model_dump()
                s_name = (s_dict.get("name") or "").strip()
                s_lines = [l.strip() for l in (s_dict.get("lines") or []) if l.strip()]
            else:
                continue
            if s_name or s_lines:
                valid_signatories.append({"name": s_name, "lines": s_lines})

    if not valid_signatories:
        s_name = (name or "").strip()
        s_lines = [l.strip() for l in (designation_lines or []) if l.strip()]
        if s_name or s_lines:
            valid_signatories.append({"name": s_name, "lines": s_lines})

    if not valid_signatories:
        return

    document.add_paragraph()  # spacer

    if len(valid_signatories) == 1:
        sig = valid_signatories[0]
        if sig["name"]:
            add_right_aligned(document, sig["name"] + ",")
        for line in sig["lines"]:
            add_right_aligned(document, line)
    else:
        # Layout signatories in a grid with MAX 2 columns per row
        num_cols = 2
        num_rows = (len(valid_signatories) + 1) // 2
        table = document.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Normal Table'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        clear_table_borders(table)

        for idx, sig in enumerate(valid_signatories):
            r_idx = idx // num_cols
            c_idx = idx % num_cols
            cell = table.rows[r_idx].cells[c_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if sig["name"]:
                run = p.add_run(sig["name"] + ",\n")
                run.bold = True
            for l_idx, line in enumerate(sig["lines"]):
                is_last = (l_idx == len(sig["lines"]) - 1)
                p.add_run(line + ("\n" if not is_last else ""))


# ============================================================================
# 3. MISC HELPERS
# ============================================================================

def slugify(text, max_len=40):
    text = re.sub(r"[^\w\s-]", "", text or "").strip().lower()
    text = re.sub(r"[\s_-]+", "_", text)
    return text[:max_len] or "quotation"


def default_filename(customer_lines, date_str):
    cust = customer_lines[0] if customer_lines else "quotation"
    digits = re.sub(r"[^\d]", "", date_str) or datetime.now().strftime("%d%m%Y")
    return f"Quotation_{slugify(cust)}_{digits}.docx"


def collect_table():
    print("\n   -- New table --")
    title = ask("   Table title/heading (optional, e.g. 'Cost Break-up')")
    ncols = ask_int("   Number of columns", default=2, min_value=1)
    headers = []
    for i in range(ncols):
        headers.append(ask(f"   Column {i + 1} header", default=f"Column {i + 1}"))
    print(f"   Now enter row data — {ncols} value(s) per row, separated by '|'.")
    rows = ask_list("   Rows", hint="blank line to finish")
    parsed_rows = []
    for row in rows:
        parts = [p.strip() for p in row.split("|")]
        if len(parts) < ncols:
            parts += [""] * (ncols - len(parts))
        parsed_rows.append(parts[:ncols])
    return {"title": title, "headers": headers, "rows": parsed_rows}


# ============================================================================
# 4. MAIN INTERACTIVE FLOW
# ============================================================================

def main():
    print("=" * 60)
    print(" QUOTATION / PROPOSAL GENERATOR")
    print("=" * 60)
    print("Press Enter to skip any optional field.")
    print("For lists (Scope of work, Terms & conditions, etc.) type one")
    print("point per line, then press Enter on a blank line to move on.")
    print("Wrap text in **double asterisks** to make it bold, e.g.")
    print("  Total cost is **Rs. 9,20,000**.")
    print("=" * 60)

    # ---- 1. Date -----------------------------------------------------
    date = ask("\nDate (DD/MM/YYYY)", default=datetime.now().strftime("%d/%m/%Y"))
    dept = ask("Department / Dept (optional, e.g. C-SMPM)")

    # ---- 2. Email / Cc -------------------------------------------------
    section("EMAIL")
    email_to_raw = ask("Email - To (comma-separated if more than one)")
    email_cc_raw = ask("Email - Cc (comma-separated, optional)")
    email_to = [e.strip() for e in email_to_raw.split(",") if e.strip()]
    email_cc = [e.strip() for e in email_cc_raw.split(",") if e.strip()]

    # ---- 3. Quotation Information --------------------------------------
    section("QUOTATION INFORMATION")
    customer_lines = ask_list(
        "Customer (name & address)",
        hint="one line at a time — name, street, city/pin — blank line to finish",
    )
    kind_attention = ask("Kind Attention")
    reference = ask("Reference")
    subject = ask("Subject")
    sac_code = ask("SAC Code")

    print("\n  Scope of work")
    scope_intro = ask("   Intro paragraph (optional)")
    scope_items = ask_list("   Bullet points")

    print("\n  Terms and conditions")
    terms_items = ask_list("   Bullet points")

    # ---- 4. Other headings: signature / approval block -----------------
    section("OTHER SECTIONS")
    tables = []
    if ask_yes_no("Add any table(s)? (e.g. cost break-up, pricing, item list)"):
        while True:
            tables.append(collect_table())
            if not ask_yes_no("   Add another table?"):
                break

    signatories = []
    if ask_yes_no("Add signatory / approval block(s) at the end?"):
        while True:
            s_name = ask("   Signatory name")
            s_lines = ask_list(
                "   Designation line(s)",
                hint="one per line, e.g. 'GH-SMC, Scientist-D' — blank line to finish",
            )
            signatories.append({"name": s_name, "lines": s_lines})
            if not ask_yes_no("   Add another signatory?"):
                break

    data = {
        "date": date,
        "dept": dept,
        "email_to": email_to,
        "email_cc": email_cc,
        "customer_lines": customer_lines,
        "kind_attention": kind_attention,
        "reference": reference,
        "subject": subject,
        "sac_code": sac_code,
        "scope_intro": scope_intro,
        "scope_items": scope_items,
        "terms_items": terms_items,
        "tables": tables,
        "signatories": signatories,
    }

    # ====================================================================
    # BUILD THE DOCUMENT
    # ====================================================================
    document = build_quotation_document(data)

    # ---- Save -----------------------------------------------------------
    default_name = default_filename(customer_lines, date)
    filename = ask("\nSave as (filename)", default=default_name)
    if not filename.lower().endswith(".docx"):
        filename += ".docx"
    document.save(filename)
    print(f"\nSaved: {filename}")





def build_quotation_document(data: dict) -> Document:
    if not isinstance(data, dict) and hasattr(data, "dict"):
        data = data.dict()
    elif not isinstance(data, dict) and hasattr(data, "model_dump"):
        data = data.model_dump()

    header_code = data.get("header_code") or "ISO 9001-2015 CMTI/PPBD/001/Rev-00"
    ref_no = data.get("ref_no") or data.get("reference") or ""
    date = data.get("date") or datetime.now().strftime("%d.%m.%Y")
    dept = data.get("dept") or ""
    email_to = data.get("email_to") or []
    email_cc = data.get("email_cc") or []

    customer_lines = data.get("customer_lines") or []
    kind_attention = data.get("kind_attention") or ""
    salutation = data.get("salutation") or "Dear Sir,"
    subject = data.get("subject") or ""
    email_ref = data.get("email_ref") or data.get("email_reference") or data.get("reference") or (email_to[0] if isinstance(email_to, list) and email_to else (email_to if isinstance(email_to, str) else ""))

    item_description = data.get("item_description") or ""
    quote_amount = data.get("quote_amount") or ""

    scope_intro = data.get("scope_intro") or ""
    scope_items = data.get("scope_items") or data.get("scope_of_work") or []

    raw_terms = (
        data.get("terms_items")
        or data.get("terms_and_conditions")
        or data.get("payment_terms_and_condition")
        or data.get("payment_terms_and_conditions")
        or []
    )
    if isinstance(raw_terms, str):
        terms_items = [t.strip() for line in raw_terms.split("\n") for t in line.split(";") if t.strip()]
    elif isinstance(raw_terms, list):
        terms_items = [str(t).strip() for t in raw_terms if str(t).strip()]
    else:
        terms_items = []

    validity = data.get("validity") or ""
    payment_terms = (
        data.get("payment_terms")
        or data.get("payment_terms_and_condition")
        or data.get("payment_terms_and_conditions")
        or ""
    )
    delivery = data.get("delivery") or ""

    contact_details = data.get("contact_details") or ""
    commercial_contact = data.get("commercial_contact") or ""

    signatory_name = data.get("signatory_name") or ""
    signatory_lines = data.get("signatory_lines") or []

    document = Document()
    set_a4_page(document)

    # 1. Header Code in Section Header (Appears on every page's top-right corner)
    section = document.sections[0]
    header = section.header
    p_hdr = header.paragraphs[0]
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_hdr.paragraph_format.space_after = Pt(0)
    r_hdr = p_hdr.add_run(header_code)
    r_hdr.font.name = "Calibri"
    r_hdr.font.size = Pt(9)
    r_hdr.font.bold = True

    if dept:
        p_dept = document.add_paragraph()
        p_dept.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_dept = p_dept.add_run(f"Dept: {dept}")
        r_dept.font.size = Pt(9)

    # 2. Document Title
    add_section_heading(document, "Quotation", center=True)

    if email_to:
        add_email_field(document, "Email", email_to)
    if email_cc:
        add_email_field(document, "Cc", email_cc)

    # 3. Ref No & Date
    p_ref = document.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_ref = p_ref.add_run(f"Ref No.: {ref_no}\nDate: {date}")
    r_ref.font.size = Pt(11)
    r_ref.font.bold = True

    document.add_paragraph()

    # 4. Customer Address Block
    for cline in customer_lines:
        p_c = document.add_paragraph()
        p_c.paragraph_format.space_after = Pt(2)
        r_c = p_c.add_run(cline)
        r_c.font.size = Pt(11)

    if kind_attention:
        p_ka = document.add_paragraph()
        p_ka.paragraph_format.space_after = Pt(4)
        r_ka = p_ka.add_run(f"Kind Attn: {kind_attention}")
        r_ka.font.bold = True
        r_ka.font.size = Pt(11)

    p_sal = document.add_paragraph()
    p_sal.paragraph_format.space_after = Pt(6)
    r_sal = p_sal.add_run(salutation)
    r_sal.font.size = Pt(11)

    # 5. Subject & Reference
    if subject:
        p_sub = document.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(4)
        r_sub = p_sub.add_run(f"Sub: {subject}")
        r_sub.font.bold = True
        r_sub.font.size = Pt(11)

    if email_ref:
        p_eref = document.add_paragraph()
        p_eref.paragraph_format.space_after = Pt(8)
        r_eref = p_eref.add_run(f"Ref: {email_ref}")
        r_eref.font.size = Pt(11)

    # 6. Intro Paragraph
    if item_description:
        p_intro = document.add_paragraph()
        p_intro.paragraph_format.space_after = Pt(8)
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_intro = p_intro.add_run(f"With reference to your above enquiry, we are pleased to submit our quote for {item_description}, with the following Terms and Conditions.")
        r_intro.font.size = Pt(11)

    # 7. Terms and Conditions
    # Point 1: Quotation Amount (if provided)
    if quote_amount:
        p_q1 = document.add_paragraph()
        p_q1.paragraph_format.space_after = Pt(6)
        r_q1_t = p_q1.add_run("1. Quotation:\n")
        r_q1_t.font.bold = True
        r_q1_t.font.size = Pt(11)
        amt_str = quote_amount if "Rs." in quote_amount or "Rupees" in quote_amount else f"Rs. {quote_amount}/-"
        r_q1_b = p_q1.add_run(f"Our Charges towards {item_description or 'the project'} is {amt_str}.")
        r_q1_b.font.size = Pt(11)

    # Point 2: Scope of work
    if scope_items:
        p_s2 = document.add_paragraph()
        p_s2.paragraph_format.space_after = Pt(4)
        r_s2 = p_s2.add_run("2. Scope of work:")
        r_s2.font.bold = True
        r_s2.font.size = Pt(11)
        add_bullets(document, scope_items)

    # Point 3: Validity
    if validity:
        p_v3 = document.add_paragraph()
        p_v3.paragraph_format.space_after = Pt(6)
        r_v3_t = p_v3.add_run("3. Quotation Validity: ")
        r_v3_t.font.bold = True
        r_v3_t.font.size = Pt(11)
        val_str = validity if "valid till" in validity.lower() else f"This quotation is valid till {validity}."
        r_v3_b = p_v3.add_run(val_str)
        r_v3_b.font.size = Pt(11)

    # Point 4: Payment Terms / Terms & Conditions
    p_pt4 = document.add_paragraph()
    p_pt4.paragraph_format.space_after = Pt(6)
    r_pt4_t = p_pt4.add_run("4. Payment Terms / Terms & Conditions:\n")
    r_pt4_t.font.bold = True
    r_pt4_t.font.size = Pt(11)
    if terms_items:
        add_bullets(document, terms_items)
    else:
        r_pt4_b = p_pt4.add_run(payment_terms)
        r_pt4_b.font.size = Pt(11)

    # Point 5: Delivery
    p_d5 = document.add_paragraph()
    p_d5.paragraph_format.space_after = Pt(6)
    r_d5_t = p_d5.add_run("5. Delivery:\n")
    r_d5_t.font.bold = True
    r_d5_t.font.size = Pt(11)
    r_d5_b = p_d5.add_run(delivery)
    r_d5_b.font.size = Pt(11)

    # Point 6: Bank Details
    p_b6 = document.add_paragraph()
    p_b6.paragraph_format.space_after = Pt(6)
    r_b6_t = p_b6.add_run("6. Bank Details:\n")
    r_b6_t.font.bold = True
    r_b6_t.font.size = Pt(11)
    r_b6_b = p_b6.add_run(
        "Name of the Bank: State Bank of India\n"
        "Address : Yeshwanthpur Branch, Tumkur Road, Bangalore-560022.\n"
        "Type of Account: Current Account\n"
        "Account No. : 10521862015"
    )
    r_b6_b.font.size = Pt(11)

    # Point 7: GST
    p_gst = document.add_paragraph()
    p_gst.paragraph_format.space_after = Pt(6)
    r_gst_t = p_gst.add_run("7. GST Charges: ")
    r_gst_t.font.bold = True
    r_gst_t.font.size = Pt(11)
    r_gst_b = p_gst.add_run("In addition to the above charges, GST will be charged as applicable at the time of billing. The present rate of GST is @ 18%.")
    r_gst_b.font.size = Pt(11)

    raw_tables = data.get("tables") or []
    for t in raw_tables:
        if isinstance(t, dict):
            add_table_block(document, t)
        elif hasattr(t, "dict"):
            add_table_block(document, t.dict())
        elif hasattr(t, "model_dump"):
            add_table_block(document, t.model_dump())

    # Point 8: Governing Law
    p_gl_head = document.add_paragraph()
    p_gl_head.paragraph_format.space_after = Pt(2)
    r_gl_t = p_gl_head.add_run("8. Governing Law and Dispute Resolution:")
    r_gl_t.font.bold = True
    r_gl_t.font.size = Pt(11)

    p_gl_body = document.add_paragraph()
    p_gl_body.paragraph_format.space_after = Pt(6)
    p_gl_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_gl_b = p_gl_body.add_run(
        "This contract/order shall be governed by and construed in accordance with the laws of India, without reference to its conflict of law's provisions. The Parties agree that any disputes arising out of or in relation to this contract shall be first attempted to be resolved mutually between the Parties, failing which, such dispute shall be finally referred to arbitration to be conducted in accordance with Arbitration and Conciliation Act, 1996 as amended till date ('Rules'). The arbitration shall be held in Bengaluru, India and shall be conducted in English language by one arbitrator, appointed by both the Parties in accordance with said Rules. If parties fail to appoint a single arbitrator, then a panel of 3 is constituted, whereby each of the parties will appoint the third arbitrator, as a contingency measure. The decision of such arbitrator/s shall be final and binding on the Parties and judgment thereon may be entered in any court of competent jurisdiction."
    )
    r_gl_b.font.size = Pt(11)

    # Point 9: Jurisdiction
    p_jur_head = document.add_paragraph()
    p_jur_head.paragraph_format.space_after = Pt(2)
    r_jur_t = p_jur_head.add_run("9. Jurisdiction:")
    r_jur_t.font.bold = True
    r_jur_t.font.size = Pt(11)

    p_jur_body = document.add_paragraph()
    p_jur_body.paragraph_format.space_after = Pt(6)
    p_jur_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_jur_b = p_jur_body.add_run("This contract shall be deemed to have been concluded in Bangalore for all purposes and therefore only courts of Bangalore shall have jurisdiction for the purpose of any adjudication in case of disputes and differences remain unsolved inspite of arbitration.")
    r_jur_b.font.size = Pt(11)

    # Point 10: PAN
    p_pan = document.add_paragraph()
    p_pan.paragraph_format.space_after = Pt(6)
    r_pan_t = p_pan.add_run("10. Our PAN No. AAATC 2085 K\n")
    r_pan_t.font.bold = True
    r_pan_t.font.size = Pt(11)
    r_pan_b = p_pan.add_run("CMTI reserves the rights to rectify the errors of typographical and clerical nature and arithmetical inaccuracies in this quotation.")
    r_pan_b.font.size = Pt(11)

    # Closing Trust Sentence
    p_trust = document.add_paragraph()
    p_trust.paragraph_format.space_after = Pt(8)
    p_trust.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_trust = p_trust.add_run("We Trust our offer meets with your requirement and we look forward to receive your order, which will receive our prompt attention and execution.\nPlease ensure to quote the reference No. and date of our letter in your Purchase order for immediate action.")
    r_trust.font.size = Pt(11)

    # Contact Info
    if contact_details or commercial_contact:
        p_cont = document.add_paragraph()
        p_cont.paragraph_format.space_after = Pt(12)
        p_cont.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cont_str = f"Should you need any information/clarification please feel free to contact {contact_details}" if contact_details else ""
        comm_str = f" For commercial clarification: {commercial_contact}" if commercial_contact else ""
        r_cont = p_cont.add_run(f"{cont_str}{comm_str}")
        r_cont.font.size = Pt(11)

    # Signatory Block
    p_sign = document.add_paragraph()
    p_sign.paragraph_format.space_after = Pt(4)
    r_sign = p_sign.add_run("Thanking you,\nYours sincerely,\nFor Central Manufacturing Technology Institute.\n\n")
    r_sign.font.size = Pt(11)

    if signatory_name:
        p_name = document.add_paragraph()
        p_name.paragraph_format.space_after = Pt(2)
        r_name = p_name.add_run(signatory_name)
        r_name.font.bold = True
        r_name.font.size = Pt(11)

    for sline in signatory_lines:
        p_sline = document.add_paragraph()
        p_sline.paragraph_format.space_after = Pt(2)
        r_sline = p_sline.add_run(sline)
        r_sline.font.size = Pt(11)

    return document


def generate_quotation_docx(data: dict) -> io.BytesIO:
    doc = build_quotation_document(data)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\nCancelled — no file was saved.")