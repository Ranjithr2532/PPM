import os
import tempfile
import uuid
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy import func
from sqlalchemy.orm import Session

from db import get_db
from models.model import DynamicTable
from services.dynamic_table_headers import compute_rows_for_header, set_cell_background

router = APIRouter(prefix="/dynamic-tables", tags=["Dynamic Tables"])

# ------------------------------------------------------------------
# Design tokens for the compact/simple document style
# ------------------------------------------------------------------
INK = RGBColor(0x22, 0x22, 0x22)
MUTE = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0x7A, 0x2E, 0x2E)
RULE_HEX = "CFCFCF"
LIGHT_HEX = "F7F7F7"
DARK_HEX = "222222"
FONT = "Calibri"


def format_indian_currency(value: Any, include_decimals: bool = True) -> str:
    if value is None or value == "":
        return ""
    try:
        if isinstance(value, str):
            clean_str = value.replace(",", "").replace(" ", "").strip()
            if not clean_str:
                return value
            num = float(clean_str)
        else:
            num = float(value)
    except (ValueError, TypeError):
        return str(value)

    if include_decimals:
        s = f"{num:.2f}"
        parts = s.split(".")
        integer_part = parts[0]
        decimal_part = parts[1]
    else:
        integer_part = str(int(round(num)))
        decimal_part = ""

    is_negative = integer_part.startswith("-")
    if is_negative:
        integer_part = integer_part[1:]

    if len(integer_part) <= 3:
        formatted_int = integer_part
    else:
        last_three = integer_part[-3:]
        remaining = integer_part[:-3]
        groups = []
        for i in range(len(remaining), 0, -2):
            start = max(0, i - 2)
            groups.insert(0, remaining[start:i])
        formatted_int = ",".join(groups) + "," + last_three

    if is_negative:
        formatted_int = "-" + formatted_int

    return f"{formatted_int}.{decimal_part}" if include_decimals else formatted_int


class DynamicTableItem(BaseModel):
    header_name: str
    columns: List[str]
    rows: List[Dict[str, Any]]
    category: Optional[str] = "recurring"


class GenerateWordPayload(BaseModel):
    title: Optional[str] = "Cost Breakdown"
    created_by: Optional[str] = None
    tables: List[DynamicTableItem]


# ------------------------------------------------------------------
# GET /dynamic-tables/{project_id}
# ------------------------------------------------------------------
@router.get("/{project_id}")
def get_saved_tables(project_id: int, db: Session = Depends(get_db)):
    """
    Returns the latest version's (raw, editable) tables for a project.
    Frontend uses this to pre-populate the Cost Estimation modal.
    """
    max_version = (
        db.query(func.max(DynamicTable.version))
        .filter(DynamicTable.project_id == project_id)
        .scalar()
    )
    if max_version is None:
        return []

    rows = (
        db.query(DynamicTable)
        .filter(
            DynamicTable.project_id == project_id,
            DynamicTable.version == max_version,
        )
        .order_by(DynamicTable.id)
        .all()
    )
    return [
        {
            "header_name": r.header_name,
            "columns": r.columns,
            "rows": r.rows,
            "category": r.category or "recurring",
        }
        for r in rows
    ]


# ------------------------------------------------------------------
# GET /dynamic-tables/{project_id}/versions
# ------------------------------------------------------------------
@router.get("/{project_id}/versions")
def list_versions(project_id: int, db: Session = Depends(get_db)):
    """
    Returns all distinct versions for a project along with metadata
    (who created it, when). One entry per version number.
    """
    subq = (
        db.query(
            DynamicTable.version,
            func.min(DynamicTable.created_at).label("created_at"),
            func.min(DynamicTable.created_by).label("created_by"),
        )
        .filter(DynamicTable.project_id == project_id)
        .group_by(DynamicTable.version)
        .order_by(DynamicTable.version)
        .all()
    )

    return [
        {
            "version": row.version,
            "created_at": row.created_at.isoformat() if row.created_at else None,
            "created_by": row.created_by,
        }
        for row in subq
    ]


# ------------------------------------------------------------------
# GET /dynamic-tables/{project_id}/version/{version}
# ------------------------------------------------------------------
@router.get("/{project_id}/version/{version}")
def get_version_tables(project_id: int, version: int, db: Session = Depends(get_db)):
    """
    Returns the (raw, editable) tables for a specific version of a project.
    Used when the user picks a version from the History drawer.
    """
    rows = (
        db.query(DynamicTable)
        .filter(
            DynamicTable.project_id == project_id,
            DynamicTable.version == version,
        )
        .order_by(DynamicTable.id)
        .all()
    )
    if not rows:
        raise HTTPException(
            status_code=404,
            detail=f"No data found for project {project_id} version {version}",
        )
    return [
        {
            "header_name": r.header_name,
            "columns": r.columns,
            "rows": r.rows,
            "category": r.category or "recurring",
        }
        for r in rows
    ]


# ------------------------------------------------------------------
# DELETE /dynamic-tables/{project_id}/version/{version}
# ------------------------------------------------------------------
@router.delete("/{project_id}/version/{version}")
def delete_version(project_id: int, version: int, db: Session = Depends(get_db)):
    """
    Permanently deletes all rows for a specific version of a project.
    Other versions are left untouched.
    """
    deleted = (
        db.query(DynamicTable)
        .filter(
            DynamicTable.project_id == project_id,
            DynamicTable.version == version,
        )
        .delete()
    )
    if deleted == 0:
        raise HTTPException(
            status_code=404,
            detail=f"No data found for project {project_id} version {version}",
        )
    db.commit()
    return {"detail": f"Version {version} deleted successfully"}


# ------------------------------------------------------------------
# GET /dynamic-tables/{project_id}/latest-costs
# ------------------------------------------------------------------
@router.get("/{project_id}/latest-costs")
def get_latest_costs_summary(project_id: int, db: Session = Depends(get_db)):
    """
    Returns a categorized cost summary (recurring vs non-recurring)
    for the latest version of a project.
    """
    max_version = (
        db.query(func.max(DynamicTable.version))
        .filter(DynamicTable.project_id == project_id)
        .scalar()
    )
    if max_version is None:
        return {
            "recurring": [],
            "non_recurring": [],
            "recurring_subtotal": 0.0,
            "non_recurring_subtotal": 0.0,
            "grand_total": 0.0
        }

    rows = (
        db.query(DynamicTable)
        .filter(
            DynamicTable.project_id == project_id,
            DynamicTable.version == max_version,
        )
        .order_by(DynamicTable.id)
        .all()
    )

    recurring_list = []
    non_recurring_list = []

    for r in rows:
        # Determine items list (roles or descriptions)
        first_col = r.columns[0] if r.columns else "Description"
        if r.header_name == "Manpower":
            first_col = "Role"
            
        items = [row.get(first_col, "") for row in r.rows if row.get(first_col)]

        # Compute subtotal
        computed_rows, total_amount = compute_rows_for_header(r.header_name, r.rows, r.columns)
        subtotal = total_amount if total_amount is not None else 0.0

        table_info = {
            "table_name": r.header_name,
            "subtotal": subtotal,
            "items": items
        }

        category = r.category or "recurring"
        if category == "recurring":
            recurring_list.append(table_info)
        else:
            non_recurring_list.append(table_info)

    rec_sub = sum(t["subtotal"] for t in recurring_list)
    non_rec_sub = sum(t["subtotal"] for t in non_recurring_list)
    grand = rec_sub + non_rec_sub

    return {
        "recurring": recurring_list,
        "non_recurring": non_recurring_list,
        "recurring_subtotal": round(rec_sub, 2),
        "non_recurring_subtotal": round(non_rec_sub, 2),
        "grand_total": round(grand, 2)
    }


# ------------------------------------------------------------------
# POST /dynamic-tables/{project_id}/generate-word
# Saves tables as a NEW version (never overwrites old versions),
# then returns the formatted .docx file in the compact/simple style.
# ------------------------------------------------------------------
@router.post("/{project_id}/generate-word")
def save_and_generate_word_document(
    project_id: int, payload: GenerateWordPayload, db: Session = Depends(get_db)
):
    """
    Saves all tables under a version number and generates a compact,
    single-accent-color corporate Word proposal.
    """
    if not payload.tables:
        raise HTTPException(status_code=400, detail="At least one table is required")

    # 1. Check if this exact table configuration matches any existing version to prevent duplicates
    matched_version = None
    versions = [v[0] for v in db.query(DynamicTable.version).filter(DynamicTable.project_id == project_id).distinct().all()]

    for ver in versions:
        db_tables = db.query(DynamicTable).filter(
            DynamicTable.project_id == project_id,
            DynamicTable.version == ver
        ).order_by(DynamicTable.header_name).all()

        p_tables = sorted(payload.tables, key=lambda x: x.header_name)
        if len(db_tables) != len(p_tables):
            continue

        match = True
        for db_tab, p_tab in zip(db_tables, p_tables):
            p_cat = getattr(p_tab, 'category', 'recurring') or 'recurring'
            db_cat = db_tab.category or 'recurring'
            if db_tab.header_name != p_tab.header_name or db_tab.columns != p_tab.columns or db_tab.rows != p_tab.rows or db_cat != p_cat:
                match = False
                break
        if match:
            matched_version = ver
            break

    if matched_version is not None:
        new_version = matched_version
    else:
        max_version = (
            db.query(func.max(DynamicTable.version))
            .filter(DynamicTable.project_id == project_id)
            .scalar()
        ) or 0
        new_version = max_version + 1

        for item in payload.tables:
            db.add(
                DynamicTable(
                    project_id=project_id,
                    version=new_version,
                    header_name=item.header_name,
                    columns=item.columns,
                    rows=item.rows,
                    category=item.category or "recurring",
                    created_by=payload.created_by,
                )
            )
        db.commit()

    # 2. Build the compact/simple corporate Word document (Targeted for 1-page layout)
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # ---- low-level OOXML helpers -----------------------------------
    def set_cell_padding(cell, top=45, bottom=45, left=100, right=100):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_cell_border(cell, top=None, left=None, bottom=None, right=None):
        """Each side, if given, is a dict: {"sz": int, "color": "RRGGBB", "val": "single"}"""
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side, spec in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
            el = OxmlElement(f'w:{side}')
            if spec is None:
                el.set(qn('w:val'), 'none')
            else:
                el.set(qn('w:val'), spec.get('val', 'single'))
                el.set(qn('w:sz'), str(spec.get('sz', 4)))
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), spec.get('color', '000000'))
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def apply_table_no_borders(table):
        tblPr = table._tbl.tblPr
        borders_el = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="none"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="none"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
        tbl_look = tblPr.find(qn('w:tblLook'))
        if tbl_look is not None:
            tbl_look.addprevious(borders_el)
        else:
            tblPr.append(borders_el)

    def set_row_cant_split(row):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))

    def set_row_repeat_header(row):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:tblHeader'))

    def run_size(run, pt):
        run.font.size = Pt(pt)

    def add_section_heading(text):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '2')
        bottom.set(qn('w:color'), '7A2E2E')
        pBdr.append(bottom)
        pPr.append(pBdr)

        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.bold = True
        run_size(r, 9.5)
        r.font.name = FONT
        r.font.color.rgb = INK
        return p

    def add_subsection_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text)
        r.font.bold = True
        run_size(r, 9)
        r.font.name = FONT
        r.font.color.rgb = ACCENT
        return p

    def add_subsubsection_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text)
        r.font.bold = True
        run_size(r, 8.5)
        r.font.name = FONT
        r.font.color.rgb = INK
        return p

    def build_data_table(headers, rows, is_total_row_fn):
        n_cols = len(headers)
        table = doc.add_table(rows=1, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        apply_table_no_borders(table)

        header_cells = table.rows[0].cells
        set_row_repeat_header(table.rows[0])
        set_row_cant_split(table.rows[0])
        for c_idx, htext in enumerate(headers):
            cell = header_cells[c_idx]
            set_cell_border(cell, bottom={"sz": 6, "color": DARK_HEX})
            set_cell_background(cell, LIGHT_HEX)
            set_cell_padding(cell, top=30, bottom=30, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            h_str = str(htext).strip()
            if h_str in ["Basis", "People"]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "Total" in h_str or "Amount" in h_str or "Subtotal" in h_str or c_idx == n_cols - 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            r = p.add_run(h_str)
            r.font.bold = True
            run_size(r, 8.5)
            r.font.name = FONT
            r.font.color.rgb = INK

        for row_data in rows:
            row_cells = table.add_row().cells
            set_row_cant_split(table.rows[-1])
            is_total = is_total_row_fn(row_data)
            bg = LIGHT_HEX if is_total else None

            for c_idx, htext in enumerate(headers):
                cell = row_cells[c_idx]
                if is_total:
                    set_cell_border(cell, top={"sz": 5, "color": DARK_HEX})
                else:
                    set_cell_border(cell, bottom={"sz": 2, "color": RULE_HEX})
                if bg:
                    set_cell_background(cell, bg)
                set_cell_padding(cell, top=24, bottom=24, left=80, right=80)

                raw_val = row_data.get(htext, "")
                h_str = str(htext).strip()

                if isinstance(raw_val, (int, float)):
                    inc_dec = ("Total" in h_str or "Amount" in h_str or "Subtotal" in h_str)
                    cell_value = format_indian_currency(raw_val, include_decimals=inc_dec)
                elif isinstance(raw_val, str) and ("Rate" in h_str or "Total" in h_str or "Amount" in h_str or "Subtotal" in h_str):
                    inc_dec = ("Total" in h_str or "Amount" in h_str or "Subtotal" in h_str)
                    cell_value = format_indian_currency(raw_val, include_decimals=inc_dec)
                else:
                    cell_value = str(raw_val if raw_val is not None else "")

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

                h_str = str(htext).strip()
                if h_str in ["Basis", "People"]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif "Total" in h_str or "Amount" in h_str or "Subtotal" in h_str or c_idx == n_cols - 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                r = p.add_run(cell_value)
                run_size(r, 9 if is_total else 8.5)
                r.font.name = FONT
                r.font.bold = bool(is_total)
                r.font.color.rgb = INK
        return table

    # ---- Header / Title ----
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(2)

    title_run = header_para.add_run(payload.title or "Industry 4.0 Pilot Project")
    run_size(title_run, 15)
    title_run.font.bold = True
    title_run.font.name = FONT
    title_run.font.color.rgb = INK

    # ---- Metadata block (compact, thin rule between rows) ----
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_no_borders(meta_table)

    meta_items = [
        ("Project Name", payload.title or "Industry 4.0 Pilot Project"),
        ("Prepared By", payload.created_by or "Project Management Team"),
        ("Date Created", datetime.now().strftime("%d-%m-%Y"))
    ]

    for idx, (lbl, v) in enumerate(meta_items):
        row_cells = meta_table.rows[idx].cells
        for cell in row_cells:
            set_cell_border(cell, bottom={"sz": 2, "color": RULE_HEX})
            set_cell_padding(cell, top=18, bottom=18, left=0, right=60)

        p0 = row_cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(lbl)
        r0.font.bold = True
        run_size(r0, 8.5)
        r0.font.name = FONT
        r0.font.color.rgb = MUTE

        p1 = row_cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(v)
        run_size(r1, 8.5)
        r1.font.name = FONT
        r1.font.color.rgb = INK

    p_space1 = doc.add_paragraph()
    p_space1.paragraph_format.space_before = Pt(5)
    p_space1.paragraph_format.space_after = Pt(0)
    p_space1.paragraph_format.line_spacing = 1.0
    p_space1.add_run().font.size = Pt(1)

    # ---- Pre-compute totals ----
    grand_total = 0.0
    recurring_tables = []
    non_recurring_tables = []

    for idx, item in enumerate(payload.tables):
        if not item.columns:
            raise HTTPException(status_code=400, detail=f"'{item.header_name}' requires at least one column")

        cols = item.columns
        if item.header_name == "Manpower":
            cols = ["Role", "Rate (₹)", "Basis", "Duration", "People", "Total (₹)"]

        rows, total_amount = compute_rows_for_header(item.header_name, item.rows, item.columns)
        cat = getattr(item, 'category', 'recurring') or 'recurring'
        
        table_info = {
            "header_name": item.header_name,
            "columns": cols,
            "rows": rows,
            "subtotal": total_amount or 0.0,
            "has_subtotal": total_amount is not None
        }

        if cat == "recurring":
            recurring_tables.append(table_info)
        else:
            non_recurring_tables.append(table_info)

        if total_amount is not None:
            grand_total += total_amount

    rec_sum = sum(t["subtotal"] for t in recurring_tables)
    non_rec_sum = sum(t["subtotal"] for t in non_recurring_tables)

    # ---- 1. Cost Summary ----
    add_section_heading("1. Cost Summary")

    exec_headers = ["Ref", "Cost Section / Table Name", "Subtotal (₹)"]
    exec_rows = []

    if recurring_tables:
        exec_rows.append({
            "Ref": "Section A",
            "Cost Section / Table Name": "Recurring Expenses (Total)",
            "Subtotal (₹)": rec_sum
        })
        for idx, t in enumerate(recurring_tables):
            exec_rows.append({
                "Ref": f"A{idx + 1}",
                "Cost Section / Table Name": f"  • {t['header_name']}",
                "Subtotal (₹)": t["subtotal"]
            })

    if non_recurring_tables:
        exec_rows.append({
            "Ref": "Section B",
            "Cost Section / Table Name": "Non-Recurring Expenses (Total)",
            "Subtotal (₹)": non_rec_sum
        })
        for idx, t in enumerate(non_recurring_tables):
            exec_rows.append({
                "Ref": f"B{idx + 1}",
                "Cost Section / Table Name": f"  • {t['header_name']}",
                "Subtotal (₹)": t["subtotal"]
            })

    exec_rows.append({
        "Ref": "",
        "Cost Section / Table Name": "Grand Total",
        "Subtotal (₹)": grand_total
    })

    build_data_table(
        exec_headers,
        exec_rows,
        is_total_row_fn=lambda r: r.get("Cost Section / Table Name") in ["Grand Total", "Recurring Expenses (Total)", "Non-Recurring Expenses (Total)"],
    )

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(1)
    p_space2.paragraph_format.space_after = Pt(0)
    p_space2.paragraph_format.line_spacing = 1.0
    p_space2.add_run().font.size = Pt(1)

    # ---- 2. Detailed Cost Breakdown ----
    add_section_heading("2. Detailed Cost Breakdown")

    if recurring_tables:
        add_subsection_heading("Section A — Recurring Expenses")
        for idx, t in enumerate(recurring_tables):
            add_subsubsection_heading(f"A{idx + 1}. {t['header_name']}")
            build_data_table(
                t["columns"],
                t["rows"],
                is_total_row_fn=lambda r, cols=t["columns"]: str(r.get(cols[0], "")).strip().lower() == "total" or str(r.get("People", "")).strip().lower() == "total",
            )

    if non_recurring_tables:
        add_subsection_heading("Section B — Non-Recurring Expenses")
        for idx, t in enumerate(non_recurring_tables):
            add_subsubsection_heading(f"B{idx + 1}. {t['header_name']}")
            build_data_table(
                t["columns"],
                t["rows"],
                is_total_row_fn=lambda r, cols=t["columns"]: str(r.get(cols[0], "")).strip().lower() == "total" or str(r.get("People", "")).strip().lower() == "total",
            )

    # ---- Grand Total banner ----
    p_space3 = doc.add_paragraph()
    p_space3.paragraph_format.space_before = Pt(1)
    p_space3.paragraph_format.space_after = Pt(0)
    p_space3.paragraph_format.line_spacing = 1.0
    p_space3.add_run().font.size = Pt(1)
    grand_box = doc.add_table(rows=1, cols=1)
    grand_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = grand_box.rows[0].cells[0]
    set_cell_border(c, top={"sz": 5, "color": DARK_HEX}, bottom={"sz": 5, "color": DARK_HEX})
    set_cell_background(c, LIGHT_HEX)
    set_cell_padding(c, top=45, bottom=45, left=100, right=100)

    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    formula_suffix = " (Section A + Section B)" if (recurring_tables and non_recurring_tables) else ""
    r1 = p.add_run(f"Grand Total Estimated Cost{formula_suffix}:  ")
    r1.font.bold = True
    run_size(r1, 9.5)
    r1.font.name = FONT
    r1.font.color.rgb = INK

    r2 = p.add_run(f"₹ {format_indian_currency(grand_total)}")
    r2.font.bold = True
    run_size(r2, 10)
    r2.font.name = FONT
    r2.font.color.rgb = ACCENT

    p_space4 = doc.add_paragraph()
    p_space4.paragraph_format.space_before = Pt(1)
    p_space4.paragraph_format.space_after = Pt(0)
    p_space4.paragraph_format.line_spacing = 1.0
    p_space4.add_run().font.size = Pt(1)

    # ---- Sign-off matrix ----
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_no_borders(sig_table)

    signatories = [
        ("Prepared By", payload.created_by or "Project Engineer"),
        # ("Reviewed By", "Head of Division"),
        # ("Approved By", "Director / General Manager"),
    ]

    for col_idx, (title, role) in enumerate(signatories):
        cell_top = sig_table.rows[0].cells[col_idx]
        set_cell_border(cell_top, top={"sz": 4, "color": RULE_HEX})
        set_cell_padding(cell_top, top=100, bottom=20, left=0, right=100)

        cell_bot = sig_table.rows[1].cells[col_idx]
        set_cell_padding(cell_bot, top=20, bottom=100, left=0, right=100)
        p_bot = cell_bot.paragraphs[0]
        r_t = p_bot.add_run(f"{title}\n")
        r_t.font.bold = True
        run_size(r_t, 9.5)
        r_t.font.name = FONT
        r_t.font.color.rgb = INK
        r_r = p_bot.add_run(f"({role})")
        run_size(r_r, 9)
        r_r.font.name = FONT
        r_r.font.color.rgb = MUTE

    tmp_dir = tempfile.gettempdir()
    filename = f"cost_breakdown_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(tmp_dir, filename)
    doc.save(filepath)

    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        headers={"X-Version": str(new_version)},
    )