"""
FastAPI Router & Document Generator for ISO Document 063: Bill of Materials (BOM).
Matches the exact specification from 063-BOM _Final_BEL.docx.
"""

from typing import List, Dict, Any, Optional
import io
from fastapi import APIRouter, HTTPException, status, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from iso.header import add_header_table
from iso.finalfooter import add_footer_table

# Helper functions for Word table manipulation
def set_cell_shading(cell, color_hex: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "color", "sz", "space"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))

def set_cell_margins(cell, top=35, start=35, bottom=35, end=35):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('left', start), ('bottom', bottom), ('right', end)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_width(cell, width_inches: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")

def add_text(cell_or_paragraph, text: str, font_name: str = "Arial", font_size: float = 9.5, bold: bool = False, italic: bool = False, color: RGBColor = RGBColor(0, 0, 0), alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, space_after: int = 0):
    if hasattr(cell_or_paragraph, 'paragraphs'):
        p = cell_or_paragraph.paragraphs[0]
    else:
        p = cell_or_paragraph

    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15

    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run

router = APIRouter(prefix="/iso", tags=["ISO Bill of Materials (Doc 063)"])


class BOMItemRequest(BaseModel):
    part_name: str = ""
    specification: str = ""
    make: str = ""
    quantity: str = ""
    function_criticality: str = "NC"


class BOMSectionRequest(BaseModel):
    title: str = ""
    content: str = ""
    headers: Optional[List[str]] = None
    rows: Optional[List[List[str]]] = None


class BOMRequest(BaseModel):
    project_title: str = ""
    project_no: str = ""
    customer_name: str = ""
    assembly_name: str = ""
    bom_rev: str = "Rev00"
    items: Optional[Any] = None
    sections: Optional[List[BOMSectionRequest]] = None
    total_estimated_cost: str = ""
    prepared_by: str = ""
    approved_by: str = ""
    group_name: str = ""
    centre_dept: str = ""
    doc_no: str = "063"
    doc_date: str = ""
    doc_code: str = ""
    filename: str = "ISO_Bill_of_Materials.docx"


DEFAULT_BOM_HEADERS = [
    "Part name/Part Number",
    "Specification",
    "Make",
    "Quantity",
    "Function Criticality"
]


def create_bom_document(
    project_title: str = "",
    project_no: str = "",
    customer_name: str = "",
    assembly_name: str = "",
    bom_rev: str = "Rev00",
    items: Optional[Any] = None,
    sections: Optional[List[Dict[str, Any]]] = None,
    total_estimated_cost: str = "",
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "063",
    doc_date: str = "",
    doc_code: str = ""
) -> Document:
    doc = Document()

    # Setup margins (A4 Portrait to match 063-BOM _Final_BEL.docx)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    header_group = (group_name or centre_dept or "SMC").strip().upper()
    if header_group.startswith("G-"):
        header_group = header_group[2:]
    elif header_group.startswith("C-"):
        header_group = header_group[2:]

    final_doc_no = str(doc_no or "063").strip()
    if final_doc_no.isdigit():
        final_doc_no = final_doc_no.zfill(3)

    # 1. Header Table
    add_header_table(
        doc.sections[0],
        title=f"BILL OF MATERIALS-{header_group}",
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=final_doc_no,
        date_str=doc_date
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # 2. Summary Paragraph
    proj_desc = project_title or "Project"
    if customer_name:
        summary_line = f"SUMMARY: This document details the bill of materials for {proj_desc} ({customer_name})"
    else:
        summary_line = f"SUMMARY: This document details the bill of materials for {proj_desc}"

    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(0)
    p_sum.paragraph_format.space_after = Pt(2)
    add_text(p_sum, summary_line, font_size=10, bold=True, color=RGBColor(15, 23, 42))

    p_conf = doc.add_paragraph()
    p_conf.paragraph_format.space_before = Pt(0)
    p_conf.paragraph_format.space_after = Pt(12)
    add_text(p_conf, "CONFIDENTIAL", font_size=9, bold=True, color=RGBColor(100, 116, 139))

    # 3. BOM Table (5 Columns)
    custom_headers = list(DEFAULT_BOM_HEADERS)
    custom_rows = []

    if isinstance(items, dict):
        custom_headers = items.get("headers") or custom_headers
        custom_rows = items.get("rows") or []
    elif isinstance(items, list):
        for r_idx, item in enumerate(items):
            if isinstance(item, dict):
                part = item.get("part_name") or item.get("part_name_part_number") or item.get("part_no_spec") or item.get("item_description") or ""
                spec = item.get("specification") or item.get("spec") or ""
                make = item.get("make") or item.get("make_supplier") or ""
                qty = str(item.get("quantity") or item.get("qty") or "")
                crit = item.get("function_criticality") or item.get("criticality") or item.get("remarks") or item.get("unit") or "NC"
                custom_rows.append([part, spec, make, qty, crit])
            elif isinstance(item, list):
                custom_rows.append([str(v or "") for v in item])

    total_rows = 1 + max(1, len(custom_rows))
    table_bom = doc.add_table(rows=total_rows, cols=len(custom_headers))
    table_bom.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_bom.autofit = False

    # Exact column widths matching 063-BOM _Final_BEL.docx (Total ~7.0 inches)
    if len(custom_headers) == 5:
        col_widths = [2.2, 1.4, 1.1, 0.85, 1.45]
    else:
        col_widths = [7.0 / len(custom_headers)] * len(custom_headers)

    border_fmt = {"val": "single", "sz": "4", "color": "000000"}

    # Format Header Row
    hdr_row = table_bom.rows[0]
    for c_idx, h_text in enumerate(custom_headers):
        cell = hdr_row.cells[c_idx]
        set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
        set_cell_margins(cell, top=40, start=35, bottom=40, end=35)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if c_idx < len(col_widths):
            set_cell_width(cell, col_widths[c_idx])
        add_text(cell, str(h_text), font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Format Data Rows
    for r_idx, r_data in enumerate(custom_rows):
        data_row = table_bom.rows[r_idx + 1]
        for c_idx, val in enumerate(r_data):
            if c_idx < len(data_row.cells):
                cell = data_row.cells[c_idx]
                set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
                set_cell_margins(cell, top=30, start=35, bottom=30, end=35)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if c_idx < len(col_widths):
                    set_cell_width(cell, col_widths[c_idx])
                
                # Center align Quantity and Function Criticality
                align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (3, 4) else WD_ALIGN_PARAGRAPH.LEFT
                add_text(cell, str(val or ""), font_size=9.5, alignment=align)

    if not custom_rows:
        data_row = table_bom.rows[1]
        for c_idx in range(len(custom_headers)):
            cell = data_row.cells[c_idx]
            set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            set_cell_margins(cell, top=30, start=35, bottom=30, end=35)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if c_idx < len(col_widths):
                set_cell_width(cell, col_widths[c_idx])
            add_text(cell, "-", font_size=9.5, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 4. Footer Section
    # Add Criticality legend and Document Revision code to Word Footer
    sec_footer = doc.sections[0].footer
    
    # Criticality Legend
    p_leg = sec_footer.paragraphs[0] if sec_footer.paragraphs else sec_footer.add_paragraph()
    p_leg.text = ""
    p_leg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_leg.paragraph_format.space_before = Pt(0)
    p_leg.paragraph_format.space_after = Pt(2)
    p_leg.paragraph_format.line_spacing = 1.0
    r_leg = p_leg.add_run("SC- Safety Critical, FC- Function Critical, NC- Not Critical")
    r_leg.font.name = "Arial"
    r_leg.font.size = Pt(8.5)
    r_leg.font.italic = True
    r_leg.font.color.rgb = RGBColor(80, 80, 80)

    # ISO Revision Code
    doc_code_str = doc_code or (f"CMTI-{header_group}-QMS-{final_doc_no}/Rev00" if header_group else f"CMTI-QMS-{final_doc_no}/Rev00")
    p_rev = sec_footer.add_paragraph()
    p_rev.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_rev.paragraph_format.space_before = Pt(0)
    p_rev.paragraph_format.space_after = Pt(0)
    p_rev.paragraph_format.line_spacing = 1.0
    r_rev = p_rev.add_run(doc_code_str)
    r_rev.font.name = "Arial"
    r_rev.font.size = Pt(8.5)
    r_rev.font.color.rgb = RGBColor(80, 80, 80)

    return doc


def generate_bom_bytes(
    project_title: str = "",
    project_no: str = "",
    customer_name: str = "",
    assembly_name: str = "",
    bom_rev: str = "Rev00",
    items: Optional[Any] = None,
    sections: Optional[List[Dict[str, Any]]] = None,
    total_estimated_cost: str = "",
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "063",
    doc_date: str = "",
    doc_code: str = ""
) -> io.BytesIO:
    doc = create_bom_document(
        project_title=project_title,
        project_no=project_no,
        customer_name=customer_name,
        assembly_name=assembly_name,
        bom_rev=bom_rev,
        items=items,
        sections=sections,
        total_estimated_cost=total_estimated_cost,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date,
        doc_code=doc_code
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# API ROUTE HANDLERS
# ============================================================

@router.get(
    "/bom/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO BOM Document (.docx) via GET"
)
async def generate_bom_doc_get(
    project_title: str = Query(""),
    project_no: str = Query(""),
    customer_name: str = Query(""),
    assembly_name: str = Query(""),
    bom_rev: str = Query("Rev00"),
    prepared_by: str = Query(""),
    approved_by: str = Query(""),
    group_name: str = Query(""),
    centre_dept: str = Query(""),
    doc_no: str = Query("063"),
    doc_date: str = Query(""),
    doc_code: str = Query(""),
    filename: str = Query("ISO_Bill_of_Materials.docx")
):
    buffer = generate_bom_bytes(
        project_title=project_title,
        project_no=project_no,
        customer_name=customer_name,
        assembly_name=assembly_name,
        bom_rev=bom_rev,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date,
        doc_code=doc_code
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/bom/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO BOM Document (.docx) via POST"
)
async def generate_bom_doc_post(
    payload: BOMRequest
):
    filename = payload.filename or "ISO_Bill_of_Materials.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_bom_bytes(
        project_title=payload.project_title,
        project_no=payload.project_no,
        customer_name=payload.customer_name,
        assembly_name=payload.assembly_name,
        bom_rev=payload.bom_rev or "Rev00",
        items=payload.items,
        sections=payload.sections,
        total_estimated_cost=payload.total_estimated_cost,
        prepared_by=payload.prepared_by,
        approved_by=payload.approved_by,
        group_name=payload.group_name,
        centre_dept=payload.centre_dept,
        doc_no=payload.doc_no or "063",
        doc_date=payload.doc_date,
        doc_code=payload.doc_code
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )
