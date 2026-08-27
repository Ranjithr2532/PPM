"""
FastAPI Router & Document Generator for ISO Document 064: Drawing Issue Register.
Generates Word (.docx) document matching CMTI-QMS-064/Rev00 specification.
"""

from typing import List, Dict, Any, Optional
import io
from fastapi import APIRouter, HTTPException, status, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from iso.header import add_header_table
from iso.finalfooter import add_footer_table
from iso.sqap import set_cell_shading, set_cell_border, set_cell_margins, add_text

router = APIRouter(prefix="/iso", tags=["ISO Drawing Issue Register (Doc 064)"])


class DrawingItemRequest(BaseModel):
    sl_no: str = ""
    drawing_no: str = ""
    title_description: str = ""
    rev_no: str = ""
    issue_date: str = ""
    issued_to: str = ""
    no_of_copies: str = ""
    remarks: str = ""

class DrawingSectionRequest(BaseModel):
    title: str = ""
    content: str = ""
    headers: Optional[List[str]] = None
    rows: Optional[List[List[str]]] = None

class DrawingRegisterRequest(BaseModel):
    project_title: str = ""
    project_no: str = ""
    customer_name: str = ""
    sub_system: str = ""
    register_rev: str = "Rev00"
    items: Optional[Any] = None
    sections: Optional[List[DrawingSectionRequest]] = None
    prepared_by: str = ""
    approved_by: str = ""
    group_name: str = ""
    centre_dept: str = ""
    doc_no: str = "064"
    doc_date: str = ""
    filename: str = "ISO_Drawing_Issue_Register.docx"


def create_drawing_register_document(
    project_title: str = "",
    project_no: str = "",
    customer_name: str = "",
    sub_system: str = "",
    register_rev: str = "Rev00",
    items: Optional[Any] = None,
    sections: Optional[List[Dict[str, Any]]] = None,
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "064",
    doc_date: str = ""
) -> Document:
    doc = Document()

    # Setup margins (A4 Landscape for wide table view)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)

    header_group = (group_name or centre_dept or "SMPM").strip().upper()
    if header_group.startswith("G-"):
        header_group = header_group[2:]
    elif header_group.startswith("C-"):
        header_group = header_group[2:]

    # Add Standard ISO Header Table
    add_header_table(
        doc.sections[0],
        title=f"DRAWING ISSUE REGISTER-{header_group}",
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=doc_no or "064",
        date_str=doc_date
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Document Main Title Header
    main_p = doc.add_paragraph()
    add_text(main_p, "DRAWING ISSUE REGISTER", font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # Metadata Card Table
    meta_table = doc.add_table(rows=2, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    border_fmt = {"val": "single", "sz": "4", "color": "CCCCCC"}

    meta_cells_data = [
        ("Project Title:", project_title or "--", "Project No:", project_no or "--", "Revision:", register_rev or "Rev00"),
        ("Customer Name:", customer_name or "--", "Sub-system / Module:", sub_system or "--", "Date:", doc_date or "--")
    ]

    for r_idx, row_tuples in enumerate(meta_cells_data):
        row = meta_table.rows[r_idx]
        lbl1, val1, lbl2, val2, lbl3, val3 = row_tuples

        for c_idx, (l, v) in enumerate([(lbl1, val1), (lbl2, val2), (lbl3, val3)]):
            cell = row.cells[c_idx]
            set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            set_cell_margins(cell, top=25, start=25, bottom=25, end=25)
            set_cell_shading(cell, "F4F6F9")

            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_text(p, f"{l} ", font_size=8.5, bold=True)
            add_text(p, str(v), font_size=8.5, bold=False)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Main Drawing Register Table
    custom_headers = ["Sl. No.", "Drawing No.", "Title / Description of Drawing", "Rev No.", "Date of Issue", "Issued To / Department", "No. of Copies", "Remarks"]
    custom_rows = []

    if isinstance(items, dict):
        custom_headers = items.get("headers") or custom_headers
        custom_rows = items.get("rows") or []
    elif isinstance(items, list):
        for r_idx, item in enumerate(items):
            if isinstance(item, dict):
                sl = item.get("sl_no") or str(r_idx + 1)
                dno = item.get("drawing_no") or ""
                tdesc = item.get("title_description") or ""
                rno = item.get("rev_no") or ""
                idate = item.get("issue_date") or ""
                ito = item.get("issued_to") or ""
                ncopies = item.get("no_of_copies") or ""
                rem = item.get("remarks") or ""
                custom_rows.append([sl, dno, tdesc, rno, idate, ito, ncopies, rem])
            elif isinstance(item, list):
                custom_rows.append([str(v or "") for v in item])

    if custom_headers:
        tbl_dwg = doc.add_table(rows=1 + max(1, len(custom_rows)), cols=len(custom_headers))
        tbl_dwg.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_dwg.autofit = False

        # Header Row
        hdr_row = tbl_dwg.rows[0]
        for c_idx, h_text in enumerate(custom_headers):
            cell = hdr_row.cells[c_idx]
            set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            set_cell_shading(cell, "D9E2EC")
            set_cell_margins(cell, top=30, start=25, bottom=30, end=25)
            add_text(cell, str(h_text), font_size=8.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Data Rows
        for r_idx, r_data in enumerate(custom_rows):
            data_row = tbl_dwg.rows[r_idx + 1]
            for c_idx, val in enumerate(r_data):
                if c_idx < len(data_row.cells):
                    cell = data_row.cells[c_idx]
                    set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
                    set_cell_margins(cell, top=25, start=25, bottom=25, end=25)
                    add_text(cell, str(val or ""), font_size=8.5)

        if not custom_rows:
            data_row = tbl_dwg.rows[1]
            for c_idx in range(len(custom_headers)):
                cell = data_row.cells[c_idx]
                set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
                set_cell_margins(cell, top=25, start=25, bottom=25, end=25)
                add_text(cell, "-", font_size=8.5, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Render Optional Extra Sections
    sec_list = sections or []
    for s_idx, sec in enumerate(sec_list):
        if isinstance(sec, dict):
            s_title = sec.get("title") or ""
            s_content = sec.get("content") or ""
            tbl_headers = sec.get("headers") or []
            tbl_rows = sec.get("rows") or []
        else:
            s_title = getattr(sec, "title", "") or ""
            s_content = getattr(sec, "content", "") or ""
            tbl_headers = getattr(sec, "headers", []) or []
            tbl_rows = getattr(sec, "rows", []) or []

        if s_title:
            sec_p = doc.add_paragraph()
            add_text(sec_p, s_title, font_size=11, bold=True, color=RGBColor(15, 23, 42), space_after=3)

        if s_content:
            cnt_p = doc.add_paragraph()
            add_text(cnt_p, s_content, font_size=9.5, space_after=6)

        if tbl_headers:
            tbl = doc.add_table(rows=1 + len(tbl_rows), cols=len(tbl_headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False

            hdr_row = tbl.rows[0]
            for c_idx, h_text in enumerate(tbl_headers):
                cell = hdr_row.cells[c_idx]
                set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
                set_cell_shading(cell, "E2E8F0")
                set_cell_margins(cell, top=25, start=25, bottom=25, end=25)
                add_text(cell, h_text, font_size=8.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

            for r_idx, r_data in enumerate(tbl_rows):
                data_row = tbl.rows[r_idx + 1]
                for c_idx, val_text in enumerate(r_data):
                    if c_idx < len(data_row.cells):
                        cell = data_row.cells[c_idx]
                        set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
                        set_cell_margins(cell, top=25, start=25, bottom=25, end=25)
                        add_text(cell, str(val_text or ""), font_size=8.5)

            doc.add_paragraph().paragraph_format.space_after = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Add ISO Footer Signature Block Table
    doc_code_footer = f"CMTI-{header_group}-QMS-{doc_no}/Rev00"
    add_footer_table(
        doc,
        prepared_name=prepared_by,
        approved_name=approved_by,
        group_name=header_group,
        doc_code=doc_code_footer,
        in_body=True
    )

    return doc


def generate_drawing_register_bytes(
    project_title: str = "",
    project_no: str = "",
    customer_name: str = "",
    sub_system: str = "",
    register_rev: str = "Rev00",
    items: Optional[Any] = None,
    sections: Optional[List[Dict[str, Any]]] = None,
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "064",
    doc_date: str = ""
) -> io.BytesIO:
    doc = create_drawing_register_document(
        project_title=project_title,
        project_no=project_no,
        customer_name=customer_name,
        sub_system=sub_system,
        register_rev=register_rev,
        items=items,
        sections=sections,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# API ROUTE HANDLERS
# ============================================================

@router.get(
    "/drawing-register/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Drawing Register Document (.docx) via GET"
)
async def generate_drawing_register_doc_get(
    project_title: str = Query(""),
    project_no: str = Query(""),
    customer_name: str = Query(""),
    sub_system: str = Query(""),
    register_rev: str = Query("Rev00"),
    prepared_by: str = Query(""),
    approved_by: str = Query(""),
    group_name: str = Query(""),
    centre_dept: str = Query(""),
    doc_no: str = Query("064"),
    doc_date: str = Query(""),
    filename: str = Query("ISO_Drawing_Issue_Register.docx")
):
    buffer = generate_drawing_register_bytes(
        project_title=project_title,
        project_no=project_no,
        customer_name=customer_name,
        sub_system=sub_system,
        register_rev=register_rev,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/drawing-register/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Drawing Register Document (.docx) via POST"
)
async def generate_drawing_register_doc_post(
    payload: DrawingRegisterRequest
):
    filename = payload.filename or "ISO_Drawing_Issue_Register.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    sec_dicts = [s.dict() if hasattr(s, 'dict') else s for s in payload.sections] if payload.sections else None

    buffer = generate_drawing_register_bytes(
        project_title=payload.project_title,
        project_no=payload.project_no,
        customer_name=payload.customer_name,
        sub_system=payload.sub_system,
        register_rev=payload.register_rev,
        items=payload.items,
        sections=sec_dicts,
        prepared_by=payload.prepared_by,
        approved_by=payload.approved_by,
        group_name=payload.group_name,
        centre_dept=payload.centre_dept,
        doc_no=payload.doc_no,
        doc_date=payload.doc_date
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )
