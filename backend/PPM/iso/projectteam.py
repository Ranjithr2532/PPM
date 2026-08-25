import os
import io
import re
from typing import Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel
from fastapi import APIRouter, status, Query, HTTPException, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from db import get_db
from models.model import Proposal
from iso.header import add_header_table, normalize_centre_dept
from iso.finalfooter import add_footer_table

# ============================================================
# FASTAPI ROUTER
# ============================================================

router = APIRouter(prefix="/iso", tags=["ISO Project Team Form Generator"])


class TeamMemberRequest(BaseModel):
    sl_no: int
    name: str = ""
    designation: str = ""
    member_type: str = ""  # Mech/Elect/software
    roles: str = ""
    signature: str = ""


class ProjectTeamRequest(BaseModel):
    project_id: Optional[int] = None
    project_no: str = "GST2502201"
    po_reference: str = "22TNGAEH0023, 31.05.2022"
    proposal_ref: str = ""
    subject: str = 'Concerning formation of a team for the project "CMTI – Central Manufacturing Facility Digitalization"'
    prepared_by: str = ""
    approved_by: str = ""
    group_name: str = ""  # Logged-in user's group name
    centre_dept: str = ""  # Logged-in user's centre
    doc_no: Optional[str] = ""
    doc_date: Optional[str] = ""
    team_members: Optional[List[TeamMemberRequest]] = None
    review_members: Optional[List[TeamMemberRequest]] = None
    filename: Optional[str] = "CMTI_Project_Team_Letter.docx"


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_margins(cell, top=50, start=50, bottom=50, end=50):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")

    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")

    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_width(cell, width_inches):
    """Set fixed cell width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_row_height(row, height_inches, rule="exact"):
    """Set row height with fixed or min rule."""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_inches * 1440)))
    trHeight.set(qn("w:hRule"), rule)
    trPr.append(trHeight)


def format_db_date(val, date_format="%d.%m.%Y"):
    if not val:
        return ""
    if hasattr(val, "strftime"):
        return val.strftime(date_format)
    val_str = str(val).strip()
    from datetime import datetime
    if re.match(r'^\d{4}-\d{2}-\d{2}$', val_str):
        try:
            dt = datetime.strptime(val_str, "%Y-%m-%d")
            return dt.strftime(date_format)
        except Exception:
            pass
    return val_str


def add_text(
    cell,
    text,
    font_size=9,
    bold=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    font_name="Arial",
    color=None
):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    lines = str(text).split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if isinstance(color, str):
            run.font.color.rgb = RGBColor.from_string(color)
        elif isinstance(color, RGBColor):
            run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)

    return paragraph.runs[0] if paragraph.runs else None


# ============================================================
# MAIN DOCUMENT GENERATION FUNCTION
# ============================================================

def create_project_team_document(
    project_no: str = "GST2502201",
    po_reference: str = "22TNGAEH0023, 31.05.2022",
    proposal_ref: str = "",
    subject: str = 'Concerning formation of a team for the project "CMTI – Central Manufacturing Facility Digitalization"',
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "",
    doc_date: str = "",
    team_members: Optional[List[TeamMemberRequest]] = None,
    review_members: Optional[List[TeamMemberRequest]] = None
) -> Document:
    doc = Document()
    section = doc.sections[0]

    # Margins
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Injections — use dynamic group name in header title and dynamic centre_dept
    group_upper = (group_name or "SMC").strip().upper()
    
    header_group = group_upper
    if "CMTI-QMS" in group_upper:
        parts = group_upper.split("-")
        if len(parts) >= 3:
            extracted = parts[2].strip()
            header_group = extracted if extracted else "SMC"
        else:
            header_group = "SMC"

    centre_upper = normalize_centre_dept(centre_dept)
    add_header_table(
        section,
        title=f"PROJECT TEAM LETTER-{header_group}",
        page_str="1 of 1",
        centre_dept=centre_upper,
        doc_no=doc_no,
        date_str=doc_date
    )
    add_footer_table(
        section,
        prepared_name=prepared_by,
        approved_name=approved_by,
        group_name=group_upper,
        doc_code="045"
    )

    # Body spacing & metadata list - reuse the default first paragraph to avoid empty space
    if doc.paragraphs:
        p_meta = doc.paragraphs[0]
    else:
        p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(12)
    p_meta.paragraph_format.line_spacing = 1.15

    # Project No
    r_pt1 = p_meta.add_run("Project No : ")
    r_pt1.bold = True
    r_pt1.font.name = "Arial"
    r_pt1.font.size = Pt(10)
    r_pt2 = p_meta.add_run(f"{project_no}\n")
    r_pt2.font.name = "Arial"
    r_pt2.font.size = Pt(10)

    # Customer PO Reference with date
    r_po1 = p_meta.add_run("Customer PO Reference with date: ")
    r_po1.bold = True
    r_po1.font.name = "Arial"
    r_po1.font.size = Pt(10)
    r_po2 = p_meta.add_run(f"{po_reference}\n")
    r_po2.font.name = "Arial"
    r_po2.font.size = Pt(10)

    # Ref Proposal / Quotation
    r_ref1 = p_meta.add_run("Ref Proposal / Quotation: ")
    r_ref1.bold = True
    r_ref1.font.name = "Arial"
    r_ref1.font.size = Pt(10)
    r_ref2 = p_meta.add_run(f"{proposal_ref}\n")
    r_ref2.font.name = "Arial"
    r_ref2.font.size = Pt(10)

    # Subject
    r_sub1 = p_meta.add_run("Subject: ")
    r_sub1.bold = True
    r_sub1.font.name = "Arial"
    r_sub1.font.size = Pt(10)
    r_sub2 = p_meta.add_run(f"{subject}")
    r_sub2.font.name = "Arial"
    r_sub2.font.size = Pt(10)

    # Description Paragraph
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_before = Pt(12)
    p_desc.paragraph_format.space_after = Pt(12)
    r_desc = p_desc.add_run(
        "With reference to the above subject and project title, following representations have been "
        "identified for assistance and timely execution of the project."
    )
    r_desc.font.name = "Arial"
    r_desc.font.size = Pt(10)

    # Default Project Team members if none provided (empty by default as requested)
    if not team_members:
        team_members = []

    # Default Review Team members if none provided (empty by default as requested)
    if not review_members:
        review_members = []

    border_format = {"val": "single", "sz": 4, "color": "000000"}
    col_widths = [0.50, 1.50, 1.10, 1.20, 1.17, 0.80]

    # ==========================================
    # 1. PROJECT TEAM TABLE
    # ==========================================
    team_table = doc.add_table(rows=len(team_members) + 1, cols=6)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    team_table.autofit = False

    # Apply widths & height constraints
    for row in team_table.rows:
        for c_idx, w in enumerate(col_widths):
            set_cell_width(row.cells[c_idx], w)
    set_row_height(team_table.rows[0], 0.30, rule="exact")
    for r_idx in range(1, len(team_members) + 1):
        set_row_height(team_table.rows[r_idx], 0.25, rule="atLeast")

    # Cell borders & padding
    for row in team_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=25, start=45, bottom=25, end=45)

    # Populate Headers
    headers = ['Sl. No.', 'Name', 'Designation', 'Type (Mech/ Elect/software)', 'Roles', 'Signature']
    for idx, h_text in enumerate(headers):
        add_text(team_table.cell(0, idx), h_text, font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Populate Rows
    for idx, member in enumerate(team_members):
        row_idx = idx + 1
        add_text(team_table.cell(row_idx, 0), f"{member.sl_no}.", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(team_table.cell(row_idx, 1), member.name, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(team_table.cell(row_idx, 2), member.designation, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(team_table.cell(row_idx, 3), member.member_type, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(team_table.cell(row_idx, 4), member.roles, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(team_table.cell(row_idx, 5), member.signature, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Spacing between tables
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(12)
    p_space.paragraph_format.space_after = Pt(6)

    # Review Team Header
    p_rev = doc.add_paragraph()
    p_rev.paragraph_format.space_before = Pt(6)
    p_rev.paragraph_format.space_after = Pt(6)
    r_rev = p_rev.add_run("Review Team")
    r_rev.bold = True
    r_rev.font.name = "Arial"
    r_rev.font.size = Pt(10)

    # ==========================================
    # 2. REVIEW TEAM TABLE
    # ==========================================
    rev_table = doc.add_table(rows=len(review_members) + 1, cols=6)
    rev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rev_table.autofit = False

    for row in rev_table.rows:
        for c_idx, w in enumerate(col_widths):
            set_cell_width(row.cells[c_idx], w)
    set_row_height(rev_table.rows[0], 0.30, rule="exact")
    for r_idx in range(1, len(review_members) + 1):
        set_row_height(rev_table.rows[r_idx], 0.25, rule="atLeast")

    for row in rev_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=25, start=45, bottom=25, end=45)

    for idx, h_text in enumerate(headers):
        add_text(rev_table.cell(0, idx), h_text, font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, member in enumerate(review_members):
        row_idx = idx + 1
        add_text(rev_table.cell(row_idx, 0), f"{member.sl_no}.", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(rev_table.cell(row_idx, 1), member.name, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(rev_table.cell(row_idx, 2), member.designation, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(rev_table.cell(row_idx, 3), member.member_type, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(rev_table.cell(row_idx, 4), member.roles, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(rev_table.cell(row_idx, 5), member.signature, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


def generate_project_team_bytes(
    project_no: str = "GST2502201",
    po_reference: str = "22TNGAEH0023, 31.05.2022",
    proposal_ref: str = "",
    subject: str = 'Concerning formation of a team for the project "CMTI – Central Manufacturing Facility Digitalization"',
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "",
    doc_date: str = "",
    team_members: Optional[List[TeamMemberRequest]] = None,
    review_members: Optional[List[TeamMemberRequest]] = None
) -> io.BytesIO:
    doc = create_project_team_document(
        project_no=project_no,
        po_reference=po_reference,
        proposal_ref=proposal_ref,
        subject=subject,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date,
        team_members=team_members,
        review_members=review_members
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# ENDPOINTS
# ============================================================

@router.get(
    "/project-team/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Project Team Letter (.docx)"
)
async def generate_project_team_doc_get(
    project_id: Optional[int] = Query(None, description="Proposal DB ID to fetch details"),
    project_no: str = Query("", description="Project Number"),
    po_reference: str = Query("", description="Customer PO Reference with date"),
    proposal_ref: str = Query("", description="Ref Proposal / Quotation"),
    subject: str = Query("", description="Subject line text"),
    prepared_by: str = Query("", description="Prepared By Signature block name"),
    approved_by: str = Query("", description="Approved By Signature block name"),
    group_name: str = Query("", description="Logged-in user's group name for header title"),
    centre_dept: str = Query("", description="Logged-in user's centre for header"),
    doc_no: str = Query("", description="Document Number"),
    doc_date: str = Query("", description="Document Date"),
    filename: str = Query("CMTI_Project_Team_Letter.docx", description="Output filename"),
    db: Session = Depends(get_db)
):
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    if project_id is not None:
        proposal = db.query(Proposal).filter(Proposal.id == project_id).first()
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        db_project_no = proposal.project_number or ""
        
        # Format PO Date and number
        db_po_date = format_db_date(proposal.order_date, "%d.%m.%Y")
        db_po_reference = f"{proposal.order_number or ''}"
        if db_po_reference and db_po_date:
            db_po_reference += f", {db_po_date}"
        elif db_po_date:
            db_po_reference = db_po_date

        db_proposal_ref = ""  # Keep empty as requested
        
        # Build subject from activity or fallback
        activity_title = proposal.activity or proposal.quote_description or ""
        db_subject = f'Concerning formation of team for the project "{activity_title}"'

        # Overwrite parameter values only if they are empty or default
        if not project_no or project_no == "GST2502201":
            project_no = db_project_no
        if not po_reference or po_reference == "22TNGAEH0023, 31.05.2022":
            po_reference = db_po_reference
        if not proposal_ref:
            proposal_ref = db_proposal_ref
        if not subject or subject in (
            'Concerning formation of team for the project ""',
            'Concerning formation of team for the project "CMTI – Central Manufacturing Facility Digitalization"',
            'Concerning formation of a team for the project "CMTI – Central Manufacturing Facility Digitalization"'
        ):
            subject = db_subject

    # Defaults if empty
    if not project_no: project_no = "GST2502201"
    if not po_reference: po_reference = "22TNGAEH0023, 31.05.2022"
    if not subject: subject = 'Concerning formation of team for the project "CMTI – Central Manufacturing Facility Digitalization"'

    buffer = generate_project_team_bytes(
        project_no=project_no,
        po_reference=po_reference,
        proposal_ref=proposal_ref,
        subject=subject,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/project-team/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Project Team Letter (.docx) via POST"
)
async def generate_project_team_doc_post(
    payload: ProjectTeamRequest,
    db: Session = Depends(get_db)
):
    filename = payload.filename or "CMTI_Project_Team_Letter.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    project_no = payload.project_no
    po_reference = payload.po_reference
    proposal_ref = payload.proposal_ref
    subject = payload.subject
    prepared_by = payload.prepared_by
    approved_by = payload.approved_by
    group_name = payload.group_name or ""
    centre_dept = payload.centre_dept or ""
    doc_no = payload.doc_no or ""
    doc_date = payload.doc_date or ""

    if payload.project_id is not None:
        proposal = db.query(Proposal).filter(Proposal.id == payload.project_id).first()
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        db_project_no = proposal.project_number or ""
        db_po_date = format_db_date(proposal.order_date, "%d.%m.%Y")
        db_po_reference = f"{proposal.order_number or ''}"
        if db_po_reference and db_po_date:
            db_po_reference += f", {db_po_date}"
        elif db_po_date:
            db_po_reference = db_po_date

        db_proposal_ref = ""  # Keep empty as requested
        
        activity_title = proposal.activity or proposal.quote_description or ""
        db_subject = f'Concerning formation of team for the project "{activity_title}"'

        # Overwrite parameter values only if they are empty or default
        if not project_no or project_no == "GST2502201":
            project_no = db_project_no
        if not po_reference or po_reference == "22TNGAEH0023, 31.05.2022":
            po_reference = db_po_reference
        if not proposal_ref:
            proposal_ref = db_proposal_ref
        if not subject or subject in (
            'Concerning formation of team for the project ""',
            'Concerning formation of team for the project "CMTI – Central Manufacturing Facility Digitalization"',
            'Concerning formation of a team for the project "CMTI – Central Manufacturing Facility Digitalization"'
        ):
            subject = db_subject

    buffer = generate_project_team_bytes(
        project_no=project_no,
        po_reference=po_reference,
        proposal_ref=proposal_ref,
        subject=subject,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date,
        team_members=payload.team_members,
        review_members=payload.review_members
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


# ============================================================
# STANDALONE TEST
# ============================================================

if __name__ == "__main__":
    doc = create_project_team_document()
    output_filename = "CMTI_Project_Team_Letter.docx"
    doc.save(output_filename)
    print(f"Document created successfully: {output_filename}")
