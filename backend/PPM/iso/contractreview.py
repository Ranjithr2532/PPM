import os
import io
from typing import Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel
from fastapi import APIRouter, status, Query, HTTPException, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from db import get_db
from models.model import Proposal, ISODocumentList
from iso.header import add_header_table
from iso.finalfooter import add_footer_table

# ============================================================
# FASTAPI ROUTER
# ============================================================

router = APIRouter(prefix="/iso", tags=["ISO Contract Review Checklist Generator"])


class ContractReviewItemRequest(BaseModel):
    sl_no: int
    checklist: str
    quotation_val: Optional[str] = None
    po_val: Optional[str] = None
    decision: Optional[str] = None


class ContractReviewRequest(BaseModel):
    project_id: Optional[int] = None
    quote_no: str = ""
    quote_date: str = ""
    po_number: str = ""
    po_date: str = ""
    customer_name: str = ""
    select_type: str = "Quotation"  # "Quotation" | "Tender" | "Proposal"
    centre_dept: str = ""  # Logged-in user's centre
    group_name: str = ""  # Logged-in user's group name
    doc_code: Optional[str] = ""
    doc_no: Optional[str] = ""
    doc_date: Optional[str] = ""
    prepared_by: Optional[str] = ""
    approved_by: Optional[str] = ""
    review_items: Optional[List[ContractReviewItemRequest]] = None
    filename: Optional[str] = "Customer_Contract_Review_Checklist.docx"


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_margins(cell, top=50, start=50, bottom=50, end=50):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")

    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")

    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_width(cell, width_inches):
    """Set fixed cell width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_row_height(row, height_inches, rule="exact"):
    """Set row height with fixed or min rule."""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_inches * 1440)))
    trHeight.set(qn("w:hRule"), rule)
    trPr.append(trHeight)


def format_db_date(val, date_format="%d.%m.%Y"):
    if not val:
        return ""
    if hasattr(val, "strftime"):
        return val.strftime(date_format)
    val_str = str(val).strip()
    import re
    from datetime import datetime
    if re.match(r'^\d{4}-\d{2}-\d{2}$', val_str):
        try:
            dt = datetime.strptime(val_str, "%Y-%m-%d")
            return dt.strftime(date_format)
        except Exception:
            pass
    return val_str


def add_text(
    cell,
    text,
    font_size=9,
    bold=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    font_name="Arial",
    color=None
):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if isinstance(color, str):
            run.font.color.rgb = RGBColor.from_string(color)
        elif isinstance(color, RGBColor):
            run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)

    return paragraph.runs[0] if paragraph.runs else None


def add_header_type_selection(cell, selected_type="Quotation"):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    sel = (selected_type or "Quotation").strip().lower()

    # Quotation No
    run_q = paragraph.add_run("Quotation No")
    run_q.bold = True
    run_q.font.name = "Arial"
    run_q.font.size = Pt(9)
    run_q._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run_q._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if sel not in ("quotation", "quotation no", "quote"):
        run_q.font.strike = True

    run_sep1 = paragraph.add_run(" / ")
    run_sep1.bold = True
    run_sep1.font.name = "Arial"
    run_sep1.font.size = Pt(9)

    # Tender
    run_t = paragraph.add_run("Tender")
    run_t.bold = True
    run_t.font.name = "Arial"
    run_t.font.size = Pt(9)
    run_t._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run_t._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if sel != "tender":
        run_t.font.strike = True

    run_sep2 = paragraph.add_run(" / \n")
    run_sep2.bold = True
    run_sep2.font.name = "Arial"
    run_sep2.font.size = Pt(9)

    # Proposal
    run_p = paragraph.add_run("Proposal")
    run_p.bold = True
    run_p.font.name = "Arial"
    run_p.font.size = Pt(9)
    run_p._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run_p._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if sel != "proposal":
        run_p.font.strike = True


def format_contract_review_doc_code(group_name: str = "", doc_no: str = "051") -> str:
    """
    Constructs document code following the pattern:
    CMTI-QMS-<group>-<doc_no>/Rev00
    e.g., if document_no is 051 and logged-in user group is SPMA -> CMTI-QMS-SPMA-051/Rev00
    """
    raw_no = str(doc_no or "051").strip()
    clean_no = raw_no.zfill(3) if raw_no.isdigit() else raw_no
    
    clean_group = str(group_name or "").strip().upper()
    if clean_group.startswith("C-") or clean_group.startswith("G-"):
        clean_group = clean_group[2:]
    
    # If a full legacy code string was passed, extract just the group token
    if "CMTI" in clean_group:
        parts = [
            p for p in clean_group.replace("/", "-").split("-")
            if p.upper() not in ("CMTI", "QMS", "REV", "REV00", "REV0", "051", "51", "")
        ]
        clean_group = parts[0].upper() if parts else ""
    
    group_str = clean_group if clean_group else "      "
    return f"CMTI-QMS-{group_str}-{clean_no}/Rev00"


# ============================================================
# MAIN DOCUMENT GENERATION FUNCTION
# ============================================================

def create_contract_review_document(
    quote_no: str = "",
    quote_date: str = "",
    po_number: str = "",
    po_date: str = "",
    customer_name: str = "",
    select_type: str = "Quotation",
    centre_dept: str = "",
    group_name: str = "",
    doc_no: str = "",
    doc_date: str = "",
    prepared_by: str = "",
    approved_by: str = "",
    doc_code: str = "",
    review_items: Optional[List[ContractReviewItemRequest]] = None
) -> Document:
    doc = Document()
    section = doc.sections[0]

    # Margins
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Format document number (e.g., 51 -> 051)
    clean_doc_no = str(doc_no or "051").strip()
    if clean_doc_no.isdigit():
        clean_doc_no = clean_doc_no.zfill(3)

    # Injections
    add_header_table(
        section,
        title="Customer Contract Review Checklist",
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=clean_doc_no,
        date_str=doc_date
    )
    
    # Construct document code using document_no and user group: CMTI-QMS-<group>-<doc_no>/Rev00
    final_doc_code = format_contract_review_doc_code(group_name=group_name, doc_no=clean_doc_no)
    add_footer_table(
        section,
        prepared_name=prepared_by,
        approved_name=approved_by,
        group_name=group_name,
        doc_code=final_doc_code
    )

    # Spacing before first table
    if not doc.paragraphs:
        p_top_space = doc.add_paragraph()
    else:
        p_top_space = doc.paragraphs[0]
    p_top_space.paragraph_format.space_before = Pt(0)
    p_top_space.paragraph_format.space_after = Pt(4)
    p_top_space.paragraph_format.line_spacing = Pt(1)
    run_space = p_top_space.add_run()
    run_space.font.size = Pt(1)

    # ==========================================
    # 1. DETAILS TABLE (5 Columns)
    # ==========================================
    details_table = doc.add_table(rows=2, cols=5)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    details_table.autofit = False

    # Layout column widths (Total = 6.80 inches)
    details_widths = [1.30, 1.10, 1.80, 1.10, 1.50]
    for row in details_table.rows:
        for c_idx, w in enumerate(details_widths):
            set_cell_width(row.cells[c_idx], w)

    # Heights (minimum heights)
    set_row_height(details_table.rows[0], 0.35, rule="atLeast")
    set_row_height(details_table.rows[1], 0.40, rule="atLeast")

    # Cell-level formatting
    border_format = {"val": "single", "sz": 4, "color": "000000"}
    for row in details_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=25, start=45, bottom=25, end=45)

    # Populate Details Header
    add_header_type_selection(details_table.cell(0, 0), selected_type=select_type)
    add_text(details_table.cell(0, 1), "Quotation\nDate / Tender", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(details_table.cell(0, 2), "Purchase Order\nNo", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(details_table.cell(0, 3), "Purchase Order\nDate", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(details_table.cell(0, 4), "Customer Name", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Populate Details Values
    add_text(details_table.cell(1, 0), quote_no, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(details_table.cell(1, 1), quote_date, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(details_table.cell(1, 2), po_number, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(details_table.cell(1, 3), po_date, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(details_table.cell(1, 4), customer_name, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ==========================================
    # 2. TITLE SPACING
    # ==========================================
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(16)
    p_title.paragraph_format.space_after = Pt(10)
    p_title.paragraph_format.line_spacing = 1.0
    run_title = p_title.add_run("Customer Order Review Checklist")
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(11)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ==========================================
    # 3. REVIEW TABLE (5 Columns)
    # ==========================================
    review_table = doc.add_table(rows=16, cols=5)
    review_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    review_table.autofit = False

    # Widths (Total = 6.80 inches)
    review_widths = [0.50, 1.50, 2.00, 2.00, 0.80]
    for row in review_table.rows:
        for c_idx, w in enumerate(review_widths):
            set_cell_width(row.cells[c_idx], w)

    # Set Heights (Header is exact, others are minimum)
    set_row_height(review_table.rows[0], 0.35, rule="exact")
    for r_idx in range(1, 16):
        set_row_height(review_table.rows[r_idx], 0.25, rule="atLeast")

    # Cell-level formatting
    for row in review_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=25, start=45, bottom=25, end=45)

    # Populate Review Table Header
    add_text(review_table.cell(0, 0), "Sl No", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(review_table.cell(0, 1), "Checklist", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(review_table.cell(0, 2), "As Per Quotation /\nTender/Proposal", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(review_table.cell(0, 3), "As Per PO", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(review_table.cell(0, 4), "Decision", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Checklist Points Defaults (Default to empty columns)
    default_items = [
        (1, "Correct Company Name", "", "", ""),
        (2, "Scope of Supply including Qty", "", "", ""),
        (3, "Any Technical Requirements", "No", "", ""),
        (4, "Billing Address", "", "", ""),
        (5, "Shipping Address", "", "", ""),
        (6, "Delivery Time/Date", "", "", ""),
        (7, "Mode of Delivery", "", "", ""),
        (8, "Supporting Documentation", "", "", ""),
        (9, "National & International Standards", "NIL", "", ""),
        (10, "Payment Terms", "", "", ""),
        (11, "Any Penalty clause", "", "", ""),
        (12, "Any Claims", "", "", ""),
        (13, "Any Specific Legal Requirements", "", "", ""),
        (14, "Warranty / Guarantee", "", "", ""),
        (15, "Any Other Requirements(Specify)", "", "", "")
    ]

    # Map overrides if provided
    items_map = {}
    if review_items:
        for item in review_items:
            items_map[item.sl_no] = (item.quotation_val, item.po_val, item.decision)

    # Populate Review Rows
    for idx, (sl, checklist_text, def_q, def_p, def_d) in enumerate(default_items):
        row_idx = idx + 1
        q_val = def_q
        p_val = def_p
        d_val = def_d

        if sl in items_map:
            q_override, p_override, d_override = items_map[sl]
            if q_override is not None:
                q_val = q_override
            if p_override is not None:
                p_val = p_override
            if d_override is not None:
                d_val = d_override

        add_text(review_table.cell(row_idx, 0), f"{sl}.", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(review_table.cell(row_idx, 1), checklist_text, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(review_table.cell(row_idx, 2), q_val, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(review_table.cell(row_idx, 3), p_val, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(review_table.cell(row_idx, 4), d_val, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


def generate_contract_review_bytes(
    quote_no: str = "",
    quote_date: str = "",
    po_number: str = "",
    po_date: str = "",
    customer_name: str = "",
    select_type: str = "Quotation",
    centre_dept: str = "",
    group_name: str = "",
    doc_no: str = "",
    doc_date: str = "",
    prepared_by: str = "",
    approved_by: str = "",
    doc_code: str = "",
    review_items: Optional[List[ContractReviewItemRequest]] = None
) -> io.BytesIO:
    doc = create_contract_review_document(
        quote_no=quote_no,
        quote_date=quote_date,
        po_number=po_number,
        po_date=po_date,
        customer_name=customer_name,
        select_type=select_type,
        centre_dept=centre_dept,
        group_name=group_name,
        doc_no=doc_no,
        doc_date=doc_date,
        prepared_by=prepared_by,
        approved_by=approved_by,
        doc_code=doc_code,
        review_items=review_items
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# ENDPOINTS
# ============================================================

@router.get(
    "/contract-review/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Customer Contract Review Checklist (.docx)"
)
async def generate_contract_review_doc_get(
    project_id: Optional[int] = Query(None, description="Proposal DB ID to fetch details"),
    quote_no: str = Query("", description="Quotation Number"),
    quote_date: str = Query("", description="Quotation Date"),
    po_number: str = Query("", description="Purchase Order Number"),
    po_date: str = Query("", description="Purchase Order Date"),
    customer_name: str = Query("", description="Customer Name"),
    select_type: str = Query("Quotation", description="Selected document type: 'Quotation', 'Tender', or 'Proposal'"),
    centre_dept: str = Query("", description="Centre / Dept name for header (e.g. C-SVT)"),
    group_name: str = Query("", description="Group ISO code name for footer"),
    doc_no: str = Query("", description="Document Number for header"),
    doc_date: str = Query("", description="Document Date for header"),
    prepared_by: str = Query("", description="Prepared by signature name"),
    approved_by: str = Query("", description="Approved by signature name"),
    
    # Dynamic checklist values (Sl 1 to 15)
    q1_val: Optional[str] = Query(None), p1_val: Optional[str] = Query(None), d1_val: Optional[str] = Query(None),
    q2_val: Optional[str] = Query(None), p2_val: Optional[str] = Query(None), d2_val: Optional[str] = Query(None),
    q3_val: Optional[str] = Query("No"), p3_val: Optional[str] = Query(None), d3_val: Optional[str] = Query(None),
    q4_val: Optional[str] = Query(None), p4_val: Optional[str] = Query(None), d4_val: Optional[str] = Query(None),
    q5_val: Optional[str] = Query(None), p5_val: Optional[str] = Query(None), d5_val: Optional[str] = Query(None),
    q6_val: Optional[str] = Query(None), p6_val: Optional[str] = Query(None), d6_val: Optional[str] = Query(None),
    q7_val: Optional[str] = Query(None), p7_val: Optional[str] = Query(None), d7_val: Optional[str] = Query(None),
    q8_val: Optional[str] = Query(None), p8_val: Optional[str] = Query(None), d8_val: Optional[str] = Query(None),
    q9_val: Optional[str] = Query("NIL"), p9_val: Optional[str] = Query(None), d9_val: Optional[str] = Query(None),
    q10_val: Optional[str] = Query(None), p10_val: Optional[str] = Query(None), d10_val: Optional[str] = Query(None),
    q11_val: Optional[str] = Query(None), p11_val: Optional[str] = Query(None), d11_val: Optional[str] = Query(None),
    q12_val: Optional[str] = Query(None), p12_val: Optional[str] = Query(None), d12_val: Optional[str] = Query(None),
    q13_val: Optional[str] = Query(None), p13_val: Optional[str] = Query(None), d13_val: Optional[str] = Query(None),
    q14_val: Optional[str] = Query(None), p14_val: Optional[str] = Query(None), d14_val: Optional[str] = Query(None),
    q15_val: Optional[str] = Query(None), p15_val: Optional[str] = Query(None), d15_val: Optional[str] = Query(None),
    
    filename: str = Query("Customer_Contract_Review_Checklist.docx", description="Output filename"),
    db: Session = Depends(get_db)
):
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    if project_id is not None:
        proposal = db.query(Proposal).filter(Proposal.id == project_id).first()
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")
        
        # Pull values from database
        db_quote_no = proposal.quote_reference or ""
        db_quote_date = format_db_date(proposal.quote_date, "%d.%m.%Y")
        db_po_number = proposal.order_number or ""
        db_po_date = format_db_date(proposal.order_date, "%d-%m-%Y")
        db_customer_name = proposal.customer_name or ""

        # Auto-detect select_type from request_type
        db_request_type = (proposal.request_type or "").strip().lower()
        if "tender" in db_request_type:
            db_select_type = "Tender"
        elif "proposal" in db_request_type:
            db_select_type = "Proposal"
        else:
            db_select_type = "Quotation"

        # Overwrite parameter values if empty
        if not quote_no:
            quote_no = db_quote_no
        if not quote_date:
            quote_date = db_quote_date
        if not po_number:
            po_number = db_po_number
        if not po_date:
            po_date = db_po_date
        if not customer_name:
            customer_name = db_customer_name
        if not select_type:
            select_type = db_select_type

    # Construct review items list from query params
    review_items = [
        ContractReviewItemRequest(sl_no=1, checklist="Correct Company Name", quotation_val=q1_val, po_val=p1_val, decision=d1_val),
        ContractReviewItemRequest(sl_no=2, checklist="Scope of Supply including Qty", quotation_val=q2_val, po_val=p2_val, decision=d2_val),
        ContractReviewItemRequest(sl_no=3, checklist="Any Technical Requirements", quotation_val=q3_val, po_val=p3_val, decision=d3_val),
        ContractReviewItemRequest(sl_no=4, checklist="Billing Address", quotation_val=q4_val, po_val=p4_val, decision=d4_val),
        ContractReviewItemRequest(sl_no=5, checklist="Shipping Address", quotation_val=q5_val, po_val=p5_val, decision=d5_val),
        ContractReviewItemRequest(sl_no=6, checklist="Delivery Time/Date", quotation_val=q6_val, po_val=p6_val, decision=d6_val),
        ContractReviewItemRequest(sl_no=7, checklist="Mode of Delivery", quotation_val=q7_val, po_val=p7_val, decision=d7_val),
        ContractReviewItemRequest(sl_no=8, checklist="Supporting Documentation", quotation_val=q8_val, po_val=p8_val, decision=d8_val),
        ContractReviewItemRequest(sl_no=9, checklist="National & International Standards", quotation_val=q9_val, po_val=p9_val, decision=d9_val),
        ContractReviewItemRequest(sl_no=10, checklist="Payment Terms", quotation_val=q10_val, po_val=p10_val, decision=d10_val),
        ContractReviewItemRequest(sl_no=11, checklist="Any Penalty clause", quotation_val=q11_val, po_val=p11_val, decision=d11_val),
        ContractReviewItemRequest(sl_no=12, checklist="Any Claims", quotation_val=q12_val, po_val=p12_val, decision=d12_val),
        ContractReviewItemRequest(sl_no=13, checklist="Any Specific Legal Requirements", quotation_val=q13_val, po_val=p13_val, decision=d13_val),
        ContractReviewItemRequest(sl_no=14, checklist="Warranty / Guarantee", quotation_val=q14_val, po_val=p14_val, decision=d14_val),
        ContractReviewItemRequest(sl_no=15, checklist="Any Other Requirements(Specify)", quotation_val=q15_val, po_val=p15_val, decision=d15_val),
    ]

    # Fetch document number from ISODocumentList database table if not provided
    iso_doc = db.query(ISODocumentList).filter(
        (ISODocumentList.document_no == "051") | 
        (ISODocumentList.document_no == "51") |
        (ISODocumentList.name.ilike("%contract%review%")) |
        (ISODocumentList.name.ilike("%order%review%"))
    ).first()

    db_doc_no = iso_doc.document_no if (iso_doc and iso_doc.document_no) else "051"
    resolved_doc_no = doc_no or db_doc_no
    if resolved_doc_no and resolved_doc_no.strip().isdigit():
        resolved_doc_no = resolved_doc_no.strip().zfill(3)

    resolved_doc_code = format_contract_review_doc_code(group_name=group_name, doc_no=resolved_doc_no)

    buffer = generate_contract_review_bytes(
        quote_no=quote_no,
        quote_date=quote_date,
        po_number=po_number,
        po_date=po_date,
        customer_name=customer_name,
        select_type=select_type,
        centre_dept=centre_dept,
        group_name=group_name,
        doc_no=resolved_doc_no,
        doc_date=doc_date,
        prepared_by=prepared_by,
        approved_by=approved_by,
        doc_code=resolved_doc_code,
        review_items=review_items
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/contract-review/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Customer Contract Review Checklist (.docx) via POST"
)
async def generate_contract_review_doc_post(
    payload: ContractReviewRequest,
    db: Session = Depends(get_db)
):
    filename = payload.filename or "Customer_Contract_Review_Checklist.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    quote_no = payload.quote_no
    quote_date = payload.quote_date
    po_number = payload.po_number
    po_date = payload.po_date
    customer_name = payload.customer_name
    select_type = payload.select_type
    centre_dept = payload.centre_dept or ""
    doc_no = payload.doc_no or ""
    doc_date = payload.doc_date or ""
    prepared_by = payload.prepared_by or ""
    approved_by = payload.approved_by or ""

    if payload.project_id is not None:
        proposal = db.query(Proposal).filter(Proposal.id == payload.project_id).first()
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")
        
        # Pull values from database
        db_quote_no = proposal.quote_reference or ""
        db_quote_date = format_db_date(proposal.quote_date, "%d.%m.%Y")
        db_po_number = proposal.order_number or ""
        db_po_date = format_db_date(proposal.order_date, "%d-%m-%Y")
        db_customer_name = proposal.customer_name or ""

        # Auto-detect select_type from request_type
        db_request_type = (proposal.request_type or "").strip().lower()
        if "tender" in db_request_type:
            db_select_type = "Tender"
        elif "proposal" in db_request_type:
            db_select_type = "Proposal"
        else:
            db_select_type = "Quotation"

        # Overwrite parameter values if empty
        if not quote_no:
            quote_no = db_quote_no
        if not quote_date:
            quote_date = db_quote_date
        if not po_number:
            po_number = db_po_number
        if not po_date:
            po_date = db_po_date
        if not customer_name:
            customer_name = db_customer_name
        if not select_type:
            select_type = db_select_type

    # Fetch document number from ISODocumentList database table if not provided
    iso_doc = db.query(ISODocumentList).filter(
        (ISODocumentList.document_no == "051") | 
        (ISODocumentList.document_no == "51") |
        (ISODocumentList.name.ilike("%contract%review%")) |
        (ISODocumentList.name.ilike("%order%review%"))
    ).first()

    db_doc_no = iso_doc.document_no if (iso_doc and iso_doc.document_no) else "051"
    resolved_doc_no = doc_no or db_doc_no
    if resolved_doc_no and resolved_doc_no.strip().isdigit():
        resolved_doc_no = resolved_doc_no.strip().zfill(3)

    resolved_doc_code = format_contract_review_doc_code(group_name=payload.group_name, doc_no=resolved_doc_no)

    buffer = generate_contract_review_bytes(
        quote_no=quote_no,
        quote_date=quote_date,
        po_number=po_number,
        po_date=po_date,
        customer_name=customer_name,
        select_type=select_type,
        centre_dept=centre_dept,
        group_name=payload.group_name,
        doc_no=resolved_doc_no,
        doc_date=doc_date,
        prepared_by=prepared_by,
        approved_by=approved_by,
        doc_code=resolved_doc_code,
        review_items=payload.review_items
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


# ============================================================
# STANDALONE TEST
# ============================================================

if __name__ == "__main__":
    doc = create_contract_review_document()
    output_filename = "Customer_Contract_Review_Checklist.docx"
    doc.save(output_filename)
    print(f"Document created successfully: {output_filename}")
