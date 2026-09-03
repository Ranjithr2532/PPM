"""
FastAPI Router & Document Generator for ISO Document 053: Project Plan (Schedule Table).
Generates Word (.docx) document matching CMTI-QMS-053/Rev00 specification.
"""

from typing import List, Dict, Any, Optional
import io
from fastapi import APIRouter, HTTPException, status, Query, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from db import get_db

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from iso.header import add_header_table
from iso.finalfooter import add_footer_table

router = APIRouter(prefix="/iso", tags=["ISO Project Plan (Doc 053)"])

# Helper function to set cell background color (shading)
def set_cell_shading(cell, color_hex: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Helper function to set cell borders
def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "color", "sz", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_margins(cell, top=40, start=40, bottom=40, end=40):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('left', start), ('bottom', bottom), ('right', end)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_width(cell, width_inches: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")

def add_text(cell_or_paragraph, text: str, font_name: str = "Arial", font_size: int = 9, bold: bool = False, italic: bool = False, color: RGBColor = RGBColor(0, 0, 0), alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, space_after: int = 0):
    if hasattr(cell_or_paragraph, 'paragraphs'):
        p = cell_or_paragraph.paragraphs[0]
    else:
        p = cell_or_paragraph

    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


class PlanTaskRequest(BaseModel):
    sl_no: str = ""
    sub_no: str = ""
    task_name: str = ""
    active_weeks: Optional[List[int]] = None


class ProjectPlanRequest(BaseModel):
    project_title: str = ""
    schedule_title: str = ""
    project_no: str = ""
    customer_name: str = ""
    commencement_date: str = ""
    completion_date: str = ""
    total_months: int = 6
    tasks: Optional[List[PlanTaskRequest]] = None
    task_active_weeks: Optional[Dict[str, Any]] = None
    prepared_by: str = ""
    approved_by: str = ""
    group_name: str = ""
    centre_dept: str = ""
    doc_no: str = "053"
    doc_date: str = ""
    filename: str = "ISO_Project_Plan.docx"
    plan_type: str = "PLANNED"  # "PLANNED" or "ACTUAL"
    planned_submission_id: Optional[int] = None
    proposal_id: Optional[int] = None


DEFAULT_TASKS = []


def format_project_plan_doc_code(group_name: str = "", doc_no: str = "053") -> str:
    raw_no = str(doc_no or "053").strip()
    clean_no = raw_no.zfill(3) if raw_no.isdigit() else raw_no
    
    clean_group = str(group_name or "").strip().upper()
    if clean_group.startswith("C-") or clean_group.startswith("G-"):
        clean_group = clean_group[2:]
    
    if "CMTI" in clean_group:
        parts = [
            p for p in clean_group.replace("/", "-").split("-")
            if p.upper() not in ("CMTI", "QMS", "REV", "REV00", "REV0", "053", "53", "")
        ]
        clean_group = parts[0].upper() if parts else ""
    
    group_str = clean_group if clean_group else "      "
    return f"CMTI-QMS-{group_str}-{clean_no}/Rev00"


def create_project_plan_document(
    project_title: str = "",
    schedule_title: str = "",
    project_no: str = "",
    customer_name: str = "",
    commencement_date: str = "",
    completion_date: str = "",
    total_months: int = 6,
    tasks: Optional[List[Dict[str, Any]]] = None,
    task_active_weeks: Optional[Dict[str, Any]] = None,
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "053",
    doc_date: str = "",
    doc_code: str = "",
    plan_type: str = "PLANNED"
) -> Document:
    doc = Document()

    # Set page setup to Landscape orientation for Gantt chart clarity
    for section in doc.sections:
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.page_width = Inches(11.69) # A4 Landscape width
        section.page_height = Inches(8.27) # A4 Landscape height

    header_group = (group_name or centre_dept or "").upper()
    if header_group.startswith("G-"):
        header_group = header_group[2:]
    elif header_group.startswith("C-"):
        header_group = header_group[2:]

    final_doc_no = str(doc_no or "053").strip()
    if final_doc_no.isdigit():
        final_doc_no = final_doc_no.zfill(3)

    final_doc_code = doc_code or format_project_plan_doc_code(group_name=header_group, doc_no=final_doc_no)

    # Add Standard ISO Header Table
    is_actual = (plan_type or "").upper() == "ACTUAL"
    base_title = "ACTUAL PROJECT PLAN" if is_actual else "PROJECT PLAN"
    title_suffix = f"-{header_group}" if header_group else ""
    add_header_table(
        doc.sections[0],
        title=f"{base_title}{title_suffix}",
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=final_doc_no,
        date_str=doc_date
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Document Header Title Info
    title_p = doc.add_paragraph()
    title_str = project_title or customer_name or ("Actual Project Execution Schedule" if is_actual else "Project Plan & Execution Schedule")
    add_text(title_p, f"Project: {title_str}", font_size=11, bold=True, space_after=2)

    if schedule_title:
        sched_p = doc.add_paragraph()
        add_text(sched_p, f"Schedule: {schedule_title}", font_size=10, italic=True, space_after=4)

    months_cnt = max(1, min(12, total_months or 6))
    weeks_cols = months_cnt * 4

    # Table columns: Sl. No (0), Sub No (1), Task Name (2), Weeks 1..N (3..3+weeks_cols-1)
    total_cols = 3 + weeks_cols

    task_list = tasks or DEFAULT_TASKS

    # Calculate total table rows: 4 header rows + len(task_list)
    table_rows = 4 + len(task_list)

    table = doc.add_table(rows=table_rows, cols=total_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    border_fmt = {"val": "single", "sz": "4", "color": "D3D3D3"}

    # Set explicit fitted column widths (total width = ~10.5 in <= 11.09 in printable width)
    week_col_width = Inches(0.28) if weeks_cols <= 24 else Inches(0.22)
    col_widths = [Inches(0.35), Inches(0.35), Inches(3.1)] + [week_col_width] * weeks_cols

    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=20, start=20, bottom=20, end=20)
            if c_idx < len(col_widths):
                set_cell_width(cell, col_widths[c_idx] / Inches(1))

    # Header Row 0: Merged Project Title
    c0 = table.cell(0, 0).merge(table.cell(0, total_cols - 1))
    add_text(c0, f"PROJECT PLAN: {title_str}", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(c0, "F0F4F8")

    # Header Row 1: Merged Schedule Subtitle
    sub_title_text = schedule_title or "Software Development & Implementation Schedule"
    c1 = table.cell(1, 0).merge(table.cell(1, total_cols - 1))
    add_text(c1, sub_title_text, font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(c1, "F7F9FC")

    # Header Row 2 & 3: Merge Sl. No, Sub, Task Name vertically across Row 2 & Row 3
    c_sl = table.cell(2, 0).merge(table.cell(3, 0))
    add_text(c_sl, "Sl. No.", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(c_sl, "F0F4F8")

    c_sub = table.cell(2, 1).merge(table.cell(3, 1))
    add_text(c_sub, "Sub", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(c_sub, "F0F4F8")

    c_task = table.cell(2, 2).merge(table.cell(3, 2))
    add_text(c_task, "TASK NAME", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_shading(c_task, "F0F4F8")

    # Header Row 2 & 3: Months & Weeks
    for m in range(months_cnt):
        col_start = 3 + (m * 4)
        col_end = col_start + 3

        # Merge 4 week columns horizontally for Month header in Row 2
        c_month = table.cell(2, col_start).merge(table.cell(2, col_end))
        add_text(c_month, f"MONTH {m + 1}", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(c_month, "E6EEF8")

        # Row 3: Week Numbers (1, 2, 3, 4 per month)
        for w in range(4):
            c_idx = col_start + w
            add_text(table.cell(3, c_idx), str(w + 1), font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(table.cell(3, c_idx), "F0F4F8")

    # Populate Tasks Rows (Row 4 onwards)
    active_map = task_active_weeks or {}

    for idx, t_item in enumerate(task_list):
        r_idx = 4 + idx

        if isinstance(t_item, dict):
            sl = str(t_item.get("sl_no") or "")
            sub = str(t_item.get("sub_no") or "")
            name = str(t_item.get("task_name") or "")
            weeks = t_item.get("active_weeks") or []
        else:
            sl = str(getattr(t_item, "sl_no", ""))
            sub = str(getattr(t_item, "sub_no", ""))
            name = str(getattr(t_item, "task_name", ""))
            weeks = getattr(t_item, "active_weeks", []) or []

        # Check map overrides
        str_idx = str(idx)
        val = active_map.get(str_idx) if isinstance(active_map, dict) else None
        if val is None and isinstance(active_map, dict):
            val = active_map.get(idx) or active_map.get(f"task_{idx}")

        if isinstance(val, list):
            weeks = [int(w) for w in val if str(w).isdigit()]
        elif isinstance(val, str):
            weeks = [int(w.strip()) for w in val.split(",") if w.strip().isdigit()]

        is_main_hdr = bool(sl and not sub)

        # Set Sl. No. and Sub. No.
        add_text(table.cell(r_idx, 0), sl, font_size=8, bold=is_main_hdr, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(table.cell(r_idx, 1), sub, font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Set Task Name
        add_text(table.cell(r_idx, 2), name, font_size=8, bold=is_main_hdr)
        if is_main_hdr:
            set_cell_shading(table.cell(r_idx, 2), "F7F9FC")

        # Fill active week cells
        for m in range(months_cnt):
            for w in range(4):
                wk_num = (m * 4) + w + 1
                c_idx = 3 + (m * 4) + w
                if wk_num in weeks:
                    fill_color = "C6E0B4" if is_actual else "C6D9F1"
                    set_cell_shading(table.cell(r_idx, c_idx), fill_color)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Add ISO Footer Block Table
    add_footer_table(
        doc,
        prepared_name=prepared_by,
        approved_name=approved_by,
        group_name=header_group,
        doc_code=final_doc_code,
        in_body=True
    )

    return doc


def generate_project_plan_bytes(
    project_title: str = "",
    schedule_title: str = "",
    project_no: str = "",
    customer_name: str = "",
    commencement_date: str = "",
    completion_date: str = "",
    total_months: int = 6,
    tasks: Optional[List[PlanTaskRequest]] = None,
    task_active_weeks: Optional[Dict[str, Any]] = None,
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "053",
    doc_date: str = "",
    plan_type: str = "PLANNED"
) -> io.BytesIO:
    task_dicts = [t.dict() if hasattr(t, 'dict') else t for t in tasks] if tasks else None

    doc = create_project_plan_document(
        project_title=project_title,
        schedule_title=schedule_title,
        project_no=project_no,
        customer_name=customer_name,
        commencement_date=commencement_date,
        completion_date=completion_date,
        total_months=total_months,
        tasks=task_dicts,
        task_active_weeks=task_active_weeks,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date,
        plan_type=plan_type
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# API ROUTE HANDLERS
# ============================================================

@router.get(
    "/project-plan/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Project Plan (.docx) via GET"
)
@router.get(
    "/projectplan/generate",
    status_code=status.HTTP_200_OK,
    include_in_schema=False
)
async def generate_project_plan_doc_get(
    project_title: str = Query(""),
    schedule_title: str = Query(""),
    project_no: str = Query(""),
    customer_name: str = Query(""),
    commencement_date: str = Query(""),
    completion_date: str = Query(""),
    total_months: int = Query(6),
    prepared_by: str = Query(""),
    approved_by: str = Query(""),
    group_name: str = Query(""),
    centre_dept: str = Query(""),
    doc_no: str = Query("053"),
    doc_date: str = Query(""),
    plan_type: str = Query("PLANNED"),
    filename: str = Query("ISO_Project_Plan.docx")
):
    buffer = generate_project_plan_bytes(
        project_title=project_title,
        schedule_title=schedule_title,
        project_no=project_no,
        customer_name=customer_name,
        commencement_date=commencement_date,
        completion_date=completion_date,
        total_months=total_months,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date,
        plan_type=plan_type
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/project-plan/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Project Plan (.docx) via POST"
)
@router.post(
    "/projectplan/generate",
    status_code=status.HTTP_200_OK,
    include_in_schema=False
)
async def generate_project_plan_doc_post(
    payload: ProjectPlanRequest
):
    is_actual = (payload.plan_type or "").upper() == "ACTUAL"
    default_fn = "ISO_Actual_Project_Plan.docx" if is_actual else "ISO_Project_Plan.docx"
    filename = payload.filename or default_fn
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_project_plan_bytes(
        project_title=payload.project_title,
        schedule_title=payload.schedule_title,
        project_no=payload.project_no,
        customer_name=payload.customer_name,
        commencement_date=payload.commencement_date,
        completion_date=payload.completion_date,
        total_months=payload.total_months,
        tasks=payload.tasks,
        task_active_weeks=payload.task_active_weeks,
        prepared_by=payload.prepared_by,
        approved_by=payload.approved_by,
        group_name=payload.group_name,
        centre_dept=payload.centre_dept,
        doc_no=payload.doc_no,
        doc_date=payload.doc_date,
        plan_type=payload.plan_type
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/project-plan/create-actual",
    status_code=status.HTTP_200_OK,
    summary="Create Actual Project Plan payload from Planned Plan"
)
def create_actual_project_plan_endpoint(payload: ProjectPlanRequest):
    """
    Derives an Actual Project Plan from a Planned Plan.
    Copies: project info, task structure (names, sl, sub).
    Clears: execution dates (commencement_date, completion_date) and active_weeks.
    """
    derived_tasks = []
    if payload.tasks:
        for t in payload.tasks:
            t_data = t.dict() if hasattr(t, "dict") else dict(t)
            # Clear Gantt execution weeks
            t_data["active_weeks"] = []
            derived_tasks.append(t_data)

    return {
        "project_title": payload.project_title,
        "schedule_title": payload.schedule_title or "Actual Project Execution Schedule",
        "project_no": payload.project_no,
        "customer_name": payload.customer_name,
        "commencement_date": "",  # Empty for actual execution
        "completion_date": "",    # Empty for actual execution
        "total_months": payload.total_months or 6,
        "tasks": derived_tasks,
        "task_active_weeks": {},
        "prepared_by": payload.prepared_by,
        "approved_by": "",
        "group_name": payload.group_name,
        "centre_dept": payload.centre_dept,
        "doc_no": payload.doc_no or "053",
        "doc_date": payload.doc_date,
        "plan_type": "ACTUAL",
        "planned_submission_id": payload.planned_submission_id
    }


class ProjectPlanComparisonRequest(BaseModel):
    project_title: str = ""
    schedule_title: str = ""
    project_no: str = ""
    customer_name: str = ""
    commencement_date: str = ""
    completion_date: str = ""
    total_months: int = 6
    planned_tasks: Optional[List[PlanTaskRequest]] = None
    actual_tasks: Optional[List[PlanTaskRequest]] = None
    prepared_by: str = ""
    approved_by: str = ""
    group_name: str = ""
    centre_dept: str = ""
    doc_no: str = "053"
    doc_date: str = ""
    filename: str = "ISO_Project_Plan_Comparison.docx"


def create_project_plan_comparison_document(
    project_title: str = "",
    schedule_title: str = "",
    project_no: str = "",
    customer_name: str = "",
    commencement_date: str = "",
    completion_date: str = "",
    total_months: int = 6,
    planned_tasks: Optional[List[Dict[str, Any]]] = None,
    actual_tasks: Optional[List[Dict[str, Any]]] = None,
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "053",
    doc_date: str = "",
    doc_code: str = ""
) -> Document:
    doc = Document()

    # Set page setup to Landscape orientation for Gantt chart clarity
    for section in doc.sections:
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.page_width = Inches(11.69) # A4 Landscape width
        section.page_height = Inches(8.27) # A4 Landscape height

    header_group = (group_name or centre_dept or "").upper()
    if header_group.startswith("G-"):
        header_group = header_group[2:]
    elif header_group.startswith("C-"):
        header_group = header_group[2:]

    final_doc_no = str(doc_no or "053").strip()
    if final_doc_no.isdigit():
        final_doc_no = final_doc_no.zfill(3)

    final_doc_code = doc_code or format_project_plan_doc_code(group_name=header_group, doc_no=final_doc_no)

    # Add Standard ISO Header Table
    title_suffix = f"-{header_group}" if header_group else ""
    add_header_table(
        doc.sections[0],
        title=f"PROJECT PLAN - COMPARISON SHEET{title_suffix}",
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=final_doc_no,
        date_str=doc_date
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Document Header Title Info
    title_p = doc.add_paragraph()
    title_str = project_title or customer_name or "Project Plan: Planned vs Actual Comparison Sheet"
    add_text(title_p, f"Project: {title_str}", font_size=11, bold=True, space_after=2)

    sched_str = schedule_title or "Planned vs Actual Schedule Execution & Variance Matrix"
    sched_p = doc.add_paragraph()
    add_text(sched_p, f"Schedule: {sched_str}", font_size=10, italic=True, space_after=4)

    months_cnt = max(1, min(12, total_months or 6))
    weeks_cols = months_cnt * 4

    # Table columns: Sl. No (0), Sub (1), Task Name (2), Plan Type (3), Weeks 1..N (4..4+weeks_cols-1)
    total_cols = 4 + weeks_cols

    p_list = planned_tasks or []
    a_list = actual_tasks or []
    max_len = max(len(p_list), len(a_list))
    task_rows_count = max_len * 2

    table_rows = 4 + task_rows_count
    table = doc.add_table(rows=table_rows, cols=total_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    border_fmt = {"val": "single", "sz": "4", "color": "D3D3D3"}
    week_col_width = Inches(0.26) if weeks_cols <= 24 else Inches(0.20)
    col_widths = [Inches(0.35), Inches(0.35), Inches(2.8), Inches(0.45)] + [week_col_width] * weeks_cols

    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=15, start=15, bottom=15, end=15)
            if c_idx < len(col_widths):
                set_cell_width(cell, col_widths[c_idx] / Inches(1))

    # Header Row 0: Merged Project Title
    c0 = table.cell(0, 0).merge(table.cell(0, total_cols - 1))
    add_text(c0, f"PROJECT PLAN: PLANNED VS ACTUAL COMPARISON - {title_str}", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(c0, "F0F4F8")

    # Header Row 1: Merged Subtitle
    c1 = table.cell(1, 0).merge(table.cell(1, total_cols - 1))
    add_text(c1, sched_str, font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(c1, "F7F9FC")

    # Header Row 2 & 3: Merge Sl. No, Sub, Task Name, Type vertically
    c_sl = table.cell(2, 0).merge(table.cell(3, 0))
    add_text(c_sl, "Sl. No.", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(c_sl, "F0F4F8")

    c_sub = table.cell(2, 1).merge(table.cell(3, 1))
    add_text(c_sub, "Sub", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(c_sub, "F0F4F8")

    c_task = table.cell(2, 2).merge(table.cell(3, 2))
    add_text(c_task, "TASK NAME", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_shading(c_task, "F0F4F8")

    c_type = table.cell(2, 3).merge(table.cell(3, 3))
    add_text(c_type, "TYPE", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(c_type, "F0F4F8")

    # Header Row 2 & 3: Months & Weeks
    for m in range(months_cnt):
        col_start = 4 + (m * 4)
        col_end = col_start + 3

        c_month = table.cell(2, col_start).merge(table.cell(2, col_end))
        add_text(c_month, f"MONTH {m + 1}", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(c_month, "E6EEF8")

        for w in range(4):
            c_idx = col_start + w
            add_text(table.cell(3, c_idx), str(w + 1), font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(table.cell(3, c_idx), "F0F4F8")

    # Populate Tasks Rows
    for t_idx in range(max_len):
        p_item = p_list[t_idx] if t_idx < len(p_list) else {}
        a_item = a_list[t_idx] if t_idx < len(a_list) else {}

        sl = str(p_item.get("sl_no") or a_item.get("sl_no") or "")
        sub = str(p_item.get("sub_no") or a_item.get("sub_no") or "")
        name = str(p_item.get("task_name") or a_item.get("task_name") or "")

        p_weeks = p_item.get("active_weeks") or []
        a_weeks = a_item.get("active_weeks") or []

        r_planned = 4 + (t_idx * 2)
        r_actual = r_planned + 1

        is_main_hdr = bool(sl and not sub)

        # Merge Sl. No, Sub, Task Name across Planned and Actual rows for this task
        c_t_sl = table.cell(r_planned, 0).merge(table.cell(r_actual, 0))
        add_text(c_t_sl, sl, font_size=8, bold=is_main_hdr, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        c_t_sub = table.cell(r_planned, 1).merge(table.cell(r_actual, 1))
        add_text(c_t_sub, sub, font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        c_t_name = table.cell(r_planned, 2).merge(table.cell(r_actual, 2))
        add_text(c_t_name, name, font_size=8, bold=is_main_hdr)
        if is_main_hdr:
            set_cell_shading(c_t_name, "F7F9FC")

        # Planned Row indicator
        add_text(table.cell(r_planned, 3), "Plan", font_size=7, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(r_planned, 3), "EBF2FA")

        # Actual Row indicator
        add_text(table.cell(r_actual, 3), "Act", font_size=7, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(r_actual, 3), "EAF5EA")

        # Populate week cells for Planned and Actual
        for m in range(months_cnt):
            for w in range(4):
                wk_num = (m * 4) + w + 1
                c_idx = 4 + (m * 4) + w

                # Planned cell
                if wk_num in p_weeks:
                    add_text(table.cell(r_planned, c_idx), "P", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    set_cell_shading(table.cell(r_planned, c_idx), "C6D9F1")

                # Actual cell
                if wk_num in a_weeks:
                    add_text(table.cell(r_actual, c_idx), "A", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    set_cell_shading(table.cell(r_actual, c_idx), "C6E0B4")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Add ISO Footer Block Table
    add_footer_table(
        doc,
        prepared_name=prepared_by,
        approved_name=approved_by,
        group_name=header_group,
        doc_code=final_doc_code,
        in_body=True
    )

    return doc


def generate_project_plan_comparison_bytes(
    project_title: str = "",
    schedule_title: str = "",
    project_no: str = "",
    customer_name: str = "",
    commencement_date: str = "",
    completion_date: str = "",
    total_months: int = 6,
    planned_tasks: Optional[List[PlanTaskRequest]] = None,
    actual_tasks: Optional[List[PlanTaskRequest]] = None,
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "053",
    doc_date: str = ""
) -> io.BytesIO:
    p_dicts = [t.dict() if hasattr(t, 'dict') else t for t in planned_tasks] if planned_tasks else None
    a_dicts = [t.dict() if hasattr(t, 'dict') else t for t in actual_tasks] if actual_tasks else None

    doc = create_project_plan_comparison_document(
        project_title=project_title,
        schedule_title=schedule_title,
        project_no=project_no,
        customer_name=customer_name,
        commencement_date=commencement_date,
        completion_date=completion_date,
        total_months=total_months,
        planned_tasks=p_dicts,
        actual_tasks=a_dicts,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@router.post(
    "/project-plan/generate-comparison",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Project Plan Comparison Sheet (.docx)"
)
async def generate_project_plan_comparison_doc_post(
    payload: ProjectPlanComparisonRequest
):
    filename = payload.filename or "ISO_Project_Plan_Comparison.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_project_plan_comparison_bytes(
        project_title=payload.project_title,
        schedule_title=payload.schedule_title,
        project_no=payload.project_no,
        customer_name=payload.customer_name,
        commencement_date=payload.commencement_date,
        completion_date=payload.completion_date,
        total_months=payload.total_months,
        planned_tasks=payload.planned_tasks,
        actual_tasks=payload.actual_tasks,
        prepared_by=payload.prepared_by,
        approved_by=payload.approved_by,
        group_name=payload.group_name,
        centre_dept=payload.centre_dept,
        doc_no=payload.doc_no,
        doc_date=payload.doc_date
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )
