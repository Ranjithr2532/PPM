import os
import io
from typing import Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel
from fastapi import APIRouter, status, Query, HTTPException, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from db import get_db
from models.model import Proposal
from iso.header import add_header_table
from iso.finalfooter import add_footer_table

# ============================================================
# FASTAPI ROUTER
# ============================================================

router = APIRouter(prefix="/iso", tags=["ISO Feasibility Review Generator"])


class ReviewPointRequest(BaseModel):
    sl_no: int
    point: str
    response: str = ""
    details: str = ""


class FeasibilityRequest(BaseModel):
    project_id: Optional[int] = None
    party_details: str = "Bharat Electronics Limited"
    enquiry_ref: str = "Email dated 19-06-2023"
    description: str = "Digitization of Fabrication & Components Group & IIOT based software OEE solution related to Machine to Machine Connectivity"
    review_points: Optional[List[ReviewPointRequest]] = None
    conclusion: str = "Feasible"  # "Feasible" or "Not Feasible"
    centre_dept: str = ""  # Logged-in user's centre
    group_name: str = ""  # Logged-in user's group ISO code
    doc_no: Optional[str] = ""
    doc_date: Optional[str] = ""
    prepared_by: Optional[str] = ""
    approved_by: Optional[str] = ""
    filename: Optional[str] = "CMTI_Feasibility_Report.docx"


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_margins(cell, top=50, start=50, bottom=50, end=50):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")

    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")

    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_width(cell, width_inches):
    """Set fixed cell width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_row_height(row, height_inches, rule="exact"):
    """Set row height with fixed or min rule."""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_inches * 1440)))
    trHeight.set(qn("w:hRule"), rule)
    trPr.append(trHeight)


def add_text(
    cell,
    text,
    font_size=10,
    bold=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    font_name="Arial",
    color=None
):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    # Handle multiline text in input
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if isinstance(color, str):
            run.font.color.rgb = RGBColor.from_string(color)
        elif isinstance(color, RGBColor):
            run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)

    return paragraph.runs[0] if paragraph.runs else None




# ============================================================
# MAIN DOCUMENT GENERATION FUNCTION
# ============================================================

def create_feasibility_document(
    party_details: str = "",
    enquiry_ref: str = "",
    description: str = "",
    review_points: Optional[List[ReviewPointRequest]] = None,
    conclusion: str = "Feasible",
    centre_dept: str = "",
    group_name: str = "",
    doc_no: str = "",
    doc_date: str = "",
    prepared_by: str = "",
    approved_by: str = "",
    doc_code: str = ""
) -> Document:
    doc = Document()
    section = doc.sections[0]

    # Margins
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Injections
    add_header_table(
        section,
        title="FEASIBILITY REVIEW FORM",
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=doc_no,
        date_str=doc_date
    )
    
    final_doc_code = doc_code or group_name or "049"
    add_footer_table(
        section,
        prepared_name=prepared_by,
        approved_name=approved_by,
        group_name=group_name,
        doc_code=final_doc_code
    )

    # Spacing between header and first body table
    if not doc.paragraphs:
        p_top_space = doc.add_paragraph()
    else:
        p_top_space = doc.paragraphs[0]
    p_top_space.paragraph_format.space_before = Pt(0)
    p_top_space.paragraph_format.space_after = Pt(4)
    p_top_space.paragraph_format.line_spacing = Pt(1)
    run_space = p_top_space.add_run()
    run_space.font.size = Pt(1)

    # ==========================================
    # 1. DETAILS TABLE
    # ==========================================
    details_table = doc.add_table(rows=2, cols=4)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    details_table.autofit = False

    # Layout column widths
    details_widths = [1.20, 1.93, 1.40, 1.73]
    for row in details_table.rows:
        for c_idx, w in enumerate(details_widths):
            set_cell_width(row.cells[c_idx], w)

    # Heights
    set_row_height(details_table.rows[0], 0.35)
    set_row_height(details_table.rows[1], 0.40)

    # Cell-level formatting
    border_format = {"val": "single", "sz": 4, "color": "000000"}
    for row in details_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=40, start=60, bottom=40, end=60)

    # Merge Row 1 Col 1-3
    desc_val_cell = details_table.cell(1, 1).merge(details_table.cell(1, 3))

    # Add text to details table
    add_text(details_table.cell(0, 0), "Party details", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_text(details_table.cell(0, 1), party_details, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_text(details_table.cell(0, 2), "Enquiry ref. No.\n(Mail dated/\ntender no.)", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_text(details_table.cell(0, 3), enquiry_ref, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_text(details_table.cell(1, 0), "Description of\nthe enquiry", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_text(desc_val_cell, description, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Space between tables - shrunken to fit cleanly
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(4)
    p_space.paragraph_format.line_spacing = Pt(1)
    run_space = p_space.add_run()
    run_space.font.size = Pt(1)

    # ==========================================
    # 2. REVIEW POINTS GRID TABLE
    # ==========================================
    review_table = doc.add_table(rows=7, cols=4)
    review_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    review_table.autofit = False

    review_widths = [0.50, 3.20, 0.90, 1.66]
    for row in review_table.rows:
        for c_idx, w in enumerate(review_widths):
            set_cell_width(row.cells[c_idx], w)

    # Heights
    set_row_height(review_table.rows[0], 0.25, rule="atLeast")
    for r in range(1, 7):
        set_row_height(review_table.rows[r], 0.25, rule="atLeast")

    # Formatting
    for row in review_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=40, start=60, bottom=40, end=60)

    # Add header text
    add_text(review_table.cell(0, 0), "Sl\nNo", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(review_table.cell(0, 1), "Review Points", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(review_table.cell(0, 2), "Yes/No/Na", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(review_table.cell(0, 3), "Details", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Default Review Points details (Only review points are filled; responses and details are blank by default)
    default_points = [
        (
            1,
            "Compliance of all technical requirements?",
            "",
            ""
        ),
        (
            2,
            "Delivery and Post Delivery activity compliance",
            "",
            ""
        ),
        (
            3,
            "Any other requirements not stated in the enquiry, but necessary in the intended use.\nEg : Item has be flame proof,\nItem has to be used in different sites etc.\nPlease mention in details",
            "",
            ""
        ),
        (
            4,
            "Any Critical / Special Characteristic identified in drawing /specifications?",
            "",
            ""
        ),
        (
            5,
            "All Statutory & Regulatory requirement applicable?\neg : Fire safety certification, etc",
            "",
            ""
        ),
        (
            6,
            "Any Operation Risk related to following is identified (if yes give details)\n1. New Technology\n2. Ability and capacity to provide product or service\n3. Short delivery time frame",
            "",
            ""
        )
    ]

    # Map requested review points overrides
    points_map = {}
    if review_points:
        for rp in review_points:
            points_map[rp.sl_no] = (rp.response, rp.details)

    # Populate Bottom Table
    for idx, (sl, pt, def_resp, def_det) in enumerate(default_points):
        row_idx = idx + 1
        resp = def_resp
        det = def_det

        if sl in points_map:
            resp, det = points_map[sl]

        add_text(review_table.cell(row_idx, 0), f"{sl}.", font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(review_table.cell(row_idx, 1), pt, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(review_table.cell(row_idx, 2), resp, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(review_table.cell(row_idx, 3), det, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # ==========================================
    # 3. CONCLUSION BLOCK
    # ==========================================
    p_concl = doc.add_paragraph()
    p_concl.paragraph_format.space_before = Pt(16)
    p_concl.paragraph_format.space_after = Pt(4)
    run_concl = p_concl.add_run("Conclusion")
    run_concl.bold = True
    run_concl.font.name = "Arial"
    run_concl.font.size = Pt(11)
    run_concl.underline = True

    # Feasible check
    p_feasible = doc.add_paragraph()
    p_feasible.paragraph_format.space_before = Pt(0)
    p_feasible.paragraph_format.space_after = Pt(4)
    run_feas_txt = p_feasible.add_run("Feasible       ")
    run_feas_txt.font.name = "Arial"
    run_feas_txt.font.size = Pt(10)
    
    feas_box = "☑" if conclusion == "Feasible" else "☐"
    run_feas_box = p_feasible.add_run(feas_box)
    run_feas_box.bold = True
    run_feas_box.font.name = "Arial"
    run_feas_box.font.size = Pt(12)

    # Not Feasible check + Note
    p_not_feasible = doc.add_paragraph()
    p_not_feasible.paragraph_format.space_before = Pt(0)
    p_not_feasible.paragraph_format.space_after = Pt(0)
    run_not_feas_txt = p_not_feasible.add_run("Not Feasible ")
    run_not_feas_txt.font.name = "Arial"
    run_not_feas_txt.font.size = Pt(10)

    not_feas_box = "☑" if conclusion == "Not Feasible" else "☐"
    run_not_feas_box = p_not_feasible.add_run(not_feas_box)
    run_not_feas_box.bold = True
    run_not_feas_box.font.name = "Arial"
    run_not_feas_box.font.size = Pt(12)

    # Add spacing and note on the same line
    p_not_feasible.add_run("      ")
    run_note_lbl = p_not_feasible.add_run("Note: ")
    run_note_lbl.bold = True
    run_note_lbl.font.name = "Arial"
    run_note_lbl.font.size = Pt(10)
    run_note_val = p_not_feasible.add_run("If not feasible, a negotiation for the terms to be made and feasibility form to be refilled for the same.")
    run_note_val.font.name = "Arial"
    run_note_val.font.size = Pt(10)

    return doc


def generate_feasibility_bytes(
    party_details: str = "",
    enquiry_ref: str = "",
    description: str = "",
    review_points: Optional[List[ReviewPointRequest]] = None,
    conclusion: str = "Feasible",
    centre_dept: str = "",
    group_name: str = "",
    doc_no: str = "",
    doc_date: str = "",
    prepared_by: str = "",
    approved_by: str = "",
    doc_code: str = ""
) -> io.BytesIO:
    doc = create_feasibility_document(
        party_details=party_details,
        enquiry_ref=enquiry_ref,
        description=description,
        review_points=review_points,
        conclusion=conclusion,
        centre_dept=centre_dept,
        group_name=group_name,
        doc_no=doc_no,
        doc_date=doc_date,
        prepared_by=prepared_by,
        approved_by=approved_by,
        doc_code=doc_code or group_name
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# API ENDPOINTS
# ============================================================

@router.get(
    "/feasibility/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Feasibility Review Word Document (.docx)"
)
async def generate_feasibility_doc_get(
    project_id: Optional[int] = Query(None, description="Project ID / Proposal ID"),
    party_details: str = Query("", description="Party details (overridden by project_id)"),
    enquiry_ref: str = Query("", description="Enquiry reference (overridden by project_id)"),
    description: str = Query("", description="Enquiry Description (overridden by project_id)"),
    conclusion: str = Query("Feasible", description="Feasible or Not Feasible"),
    centre_dept: str = Query("", description="Centre / Dept name for header (e.g. C-SVT)"),
    group_name: str = Query("", description="Group ISO code name for footer"),
    doc_code: str = Query("", description="Full Document Code e.g. CMTI-QMS-VT-049/Rev00"),
    r1_response: str = Query("", description="Compliance of all technical requirements? (Yes/No/Na)"),
    r1_details: str = Query("", description="Compliance of all technical requirements details"),
    r2_response: str = Query("", description="Delivery and Post Delivery activity compliance (Yes/No/Na)"),
    r2_details: str = Query("", description="Delivery and Post Delivery activity compliance details"),
    r3_response: str = Query("", description="Any other requirements not stated... (Yes/No/Na)"),
    r3_details: str = Query("", description="Any other requirements details"),
    r4_response: str = Query("", description="Any Critical / Special Characteristic... (Yes/No/Na)"),
    r4_details: str = Query("", description="Any Critical / Special Characteristic details"),
    r5_response: str = Query("", description="All Statutory & Regulatory requirement... (Yes/No/Na)"),
    r5_details: str = Query("", description="All Statutory & Regulatory details"),
    r6_response: str = Query("", description="Any Operation Risk... (Yes/No/Na)"),
    r6_details: str = Query("", description="Any Operation Risk details"),
    doc_no: str = Query("", description="Document Number for header"),
    doc_date: str = Query("", description="Document Date for header"),
    prepared_by: str = Query("", description="Prepared by signature name"),
    approved_by: str = Query("", description="Approved by signature name"),
    filename: str = Query("CMTI_Feasibility_Report.docx", description="Output filename"),
    db: Session = Depends(get_db)
):
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    if project_id is not None:
        proposal = db.query(Proposal).filter(Proposal.id == project_id).first()
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")
        if not party_details:
            party_details = proposal.customer_name or ""
        if not enquiry_ref:
            enquiry_ref = proposal.email_reference or ""
        if not description:
            description = proposal.quote_description or ""

    # Construct review points list from query params
    review_points = [
        ReviewPointRequest(sl_no=1, point="Compliance of all technical requirements?", response=r1_response, details=r1_details),
        ReviewPointRequest(sl_no=2, point="Delivery and Post Delivery activity compliance", response=r2_response, details=r2_details),
        ReviewPointRequest(sl_no=3, point="Any other requirements not stated in the enquiry...", response=r3_response, details=r3_details),
        ReviewPointRequest(sl_no=4, point="Any Critical / Special Characteristic...", response=r4_response, details=r4_details),
        ReviewPointRequest(sl_no=5, point="All Statutory & Regulatory requirement...", response=r5_response, details=r5_details),
        ReviewPointRequest(sl_no=6, point="Any Operation Risk related to following...", response=r6_response, details=r6_details),
    ]

    buffer = generate_feasibility_bytes(
        party_details=party_details,
        enquiry_ref=enquiry_ref,
        description=description,
        review_points=review_points,
        conclusion=conclusion,
        centre_dept=centre_dept,
        group_name=group_name,
        doc_no=doc_no,
        doc_date=doc_date,
        prepared_by=prepared_by,
        approved_by=approved_by,
        doc_code=doc_code or group_name
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/feasibility/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Feasibility Review Word Document (.docx) via POST"
)
async def generate_feasibility_doc_post(
    payload: FeasibilityRequest,
    db: Session = Depends(get_db)
):
    filename = payload.filename or "CMTI_Feasibility_Report.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    party_details = payload.party_details
    enquiry_ref = payload.enquiry_ref
    description = payload.description
    doc_no = payload.doc_no or ""
    doc_date = payload.doc_date or ""
    prepared_by = payload.prepared_by or ""
    approved_by = payload.approved_by or ""

    if payload.project_id is not None:
        proposal = db.query(Proposal).filter(Proposal.id == payload.project_id).first()
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")
        if not party_details:
            party_details = proposal.customer_name or ""
        if not enquiry_ref:
            enquiry_ref = proposal.email_reference or ""
        if not description:
            description = proposal.quote_description or ""

    buffer = generate_feasibility_bytes(
        party_details=party_details,
        enquiry_ref=enquiry_ref,
        description=description,
        review_points=payload.review_points,
        conclusion=payload.conclusion,
        centre_dept=payload.centre_dept or "",
        group_name=payload.group_name or "",
        doc_no=doc_no,
        doc_date=doc_date,
        prepared_by=prepared_by,
        approved_by=approved_by,
        doc_code=payload.group_name or ""
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


# ============================================================
# MAIN SCRIPT EXECUTION (Standalone Test)
# ============================================================

if __name__ == "__main__":
    doc = create_feasibility_document()
    output_filename = "CMTI_Feasibility_Report.docx"
    doc.save(output_filename)
    print(f"Document created successfully: {output_filename}")
