import os
import io
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel
from fastapi import APIRouter, status, Query
from fastapi.responses import StreamingResponse

# ============================================================
# FASTAPI ROUTER
# ============================================================

router = APIRouter(prefix="/iso", tags=["ISO Header Generator"])


class HeaderInfoRequest(BaseModel):
    centre_dept: str = ""
    doc_no: str = ""
    date: str = "13-01-2026"
    page: str = "1 of 2"
    title: str = "PROJECT TEAM LETTER-SMC"
    iso_spec: str = "ISO 9001-2015"
    filename: Optional[str] = "CMTI_Project_Team_Header.docx"


def normalize_centre_dept(centre_dept: Optional[str]) -> str:
    """Normalize a logged-in user centre to the header format, without forcing SMPM."""
    if centre_dept is None:
        return ""
    value = str(centre_dept).strip()
    if not value:
        return ""
    upper_value = value.upper()
    if upper_value.startswith("C-"):
        return upper_value
    return f"C-{value}".upper()


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_margins(cell, top=50, start=50, bottom=50, end=50):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")

    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")

    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_shading(cell, fill):
    """Set background color of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches):
    """Set fixed cell width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_row_height(row, height_inches):
    """Set fixed row height."""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_inches * 1440)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def add_text(
    cell,
    text,
    font_size=10,
    bold=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    font_name="Arial",
    color=None
):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if isinstance(color, str):
        run.font.color.rgb = RGBColor.from_string(color)
    elif isinstance(color, RGBColor):
        run.font.color.rgb = color

    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    return run


# ============================================================
# DOCUMENT GENERATION FUNCTION
# ============================================================

def add_header_table(
    section,
    centre_dept: str = "",
    doc_no: str = "",
    date_str: str = "13-01-2026",
    page_str: str = "1 of 1",
    title: str = "PROJECT TEAM LETTER-SMC",
    iso_spec: str = "ISO 9001-2015",
    logo_path: Optional[str] = None
):
    centre_dept = normalize_centre_dept(centre_dept)
    
    import datetime
    if not date_str:
        date_str = datetime.date.today().strftime("%d-%m-%Y")
        
    header = section.header

    # Main Header Table (5 rows x 4 cols)
    table = header.add_table(rows=5, cols=4, width=Inches(6.26))

    # Remove the default empty paragraph in the header to avoid extra top spacing
    if header.paragraphs:
        p = header.paragraphs[0]._element
        p.getparent().remove(p)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Column widths adjusted to prevent wrapping of "CENTRE / DEPT"
    col_widths = [1.025, 2.95, 1.15, 1.135]

    # Set widths and heights on raw cells
    for r_idx in range(5):
        row = table.rows[r_idx]
        h = 0.40 if r_idx == 0 else 0.22
        set_row_height(row, h)
        for c_idx in range(4):
            set_cell_width(row.cells[c_idx], col_widths[c_idx])

    # Apply borders & vertical alignment to all cells
    border_format = {"val": "single", "sz": 4, "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=border_format,
                bottom=border_format,
                left=border_format,
                right=border_format
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Merge cells
    logo_cell = table.cell(0, 0).merge(table.cell(4, 0))
    title_cell = table.cell(0, 1).merge(table.cell(0, 3))
    doc_title_cell = table.cell(1, 1).merge(table.cell(4, 1))

    # Resolve Logo Path
    if not logo_path:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        logo_path = os.path.join(os.path.dirname(base_dir), "images", "cmti.png")

    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.paragraph_format.space_before = Pt(0)
    logo_paragraph.paragraph_format.space_after = Pt(0)

    if os.path.exists(logo_path):
        run = logo_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(0.87))
    else:
        run = logo_paragraph.add_run("[CMTI LOGO]")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)

    # Add Main Title (Row 0)
    p1 = title_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("CENTRAL MANUFACTURING TECHNOLOGY INSTITUTE")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(20)

    p2 = title_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(iso_spec)
    r2.bold = True
    r2.font.name = "Arial"
    r2.font.size = Pt(20)

    # Add Document Title
    p_doc = doc_title_cell.paragraphs[0]
    p_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_doc.paragraph_format.space_before = Pt(0)
    p_doc.paragraph_format.space_after = Pt(0)
    r_doc = p_doc.add_run(title)
    r_doc.bold = True
    r_doc.font.name = "Arial"
    r_doc.font.size = Pt(16)

    # Info Key-Values
    info = [
        ("CENTRE / DEPT", centre_dept),
        ("Doc  No.", doc_no),
        ("Date", date_str),
        ("Page", page_str),
    ]

    for i, (label, val) in enumerate(info):
        cell_lbl = table.cell(i + 1, 2)
        cell_val = table.cell(i + 1, 3)

        # Set cell margins (padding) to prevent wrapping/clipping of text
        set_cell_margins(cell_lbl, top=25, start=45, bottom=25, end=45)
        set_cell_margins(cell_val, top=25, start=45, bottom=25, end=45)

        add_text(
            cell_lbl,
            label,
            font_size=10,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT
        )

        if label == "Page":
            paragraph = cell_val.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0

            # Page number run
            r_page = paragraph.add_run()
            r_page.bold = True
            r_page.font.name = "Arial"
            r_page.font.size = Pt(10)
            r_page._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
            r_page._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText1 = OxmlElement('w:instrText')
            instrText1.set(qn('xml:space'), 'preserve')
            instrText1.text = "PAGE"
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            t1 = OxmlElement('w:t')
            t1.text = "1"
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            r_page._r.append(fldChar1)
            r_page._r.append(instrText1)
            r_page._r.append(fldChar2)
            r_page._r.append(t1)
            r_page._r.append(fldChar3)

            # Separator
            run_of = paragraph.add_run(" of ")
            run_of.bold = True
            run_of.font.name = "Arial"
            run_of.font.size = Pt(10)
            run_of._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
            run_of._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

            # NumPages run
            r_num = paragraph.add_run()
            r_num.bold = True
            r_num.font.name = "Arial"
            r_num.font.size = Pt(10)
            r_num._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
            r_num._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'begin')
            instrText2 = OxmlElement('w:instrText')
            instrText2.set(qn('xml:space'), 'preserve')
            instrText2.text = "NUMPAGES"
            fldChar5 = OxmlElement('w:fldChar')
            fldChar5.set(qn('w:fldCharType'), 'separate')
            t2 = OxmlElement('w:t')
            t2.text = "1"
            fldChar6 = OxmlElement('w:fldChar')
            fldChar6.set(qn('w:fldCharType'), 'end')

            r_num._r.append(fldChar4)
            r_num._r.append(instrText2)
            r_num._r.append(fldChar5)
            r_num._r.append(t2)
            r_num._r.append(fldChar6)
        else:
            val_bold = True if label == "CENTRE/ DEPT" else False
            val_size = 9 if label == "CENTRE / DEPT" else 10
            add_text(
                cell_val,
                val,
                font_size=val_size,
                bold=val_bold,
                alignment=WD_ALIGN_PARAGRAPH.LEFT
            )

    # Add a spacing paragraph after the header table to prevent body table
    # from touching the bottom of the header table on page 2 onwards.
    p_spacing = header.add_paragraph()
    p_spacing.paragraph_format.space_before = Pt(0)
    p_spacing.paragraph_format.space_after = Pt(8)
    p_spacing.paragraph_format.line_spacing = Pt(1)
    run_sp = p_spacing.add_run()
    run_sp.font.size = Pt(1)


def create_iso_header_document(
    centre_dept: str = "",
    doc_no: str = "",
    date_str: str = "13-01-2026",
    page_str: str = "1 of 2",
    title: str = "PROJECT TEAM LETTER-SMC",
    iso_spec: str = "ISO 9001-2015",
    logo_path: Optional[str] = None
) -> Document:
    doc = Document()
    section = doc.sections[0]

    # A4 page dimensions
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # Margins
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    add_header_table(
        section,
        centre_dept=centre_dept,
        doc_no=doc_no,
        date_str=date_str,
        page_str=page_str,
        title=title,
        iso_spec=iso_spec,
        logo_path=logo_path
    )

    return doc


def generate_iso_header_bytes(
    centre_dept: str = "",
    doc_no: str = "",
    date_str: str = "13-01-2026",
    page_str: str = "1 of 2",
    title: str = "PROJECT TEAM LETTER-SMC",
    iso_spec: str = "ISO 9001-2015"
) -> io.BytesIO:
    doc = create_iso_header_document(
        centre_dept=centre_dept,
        doc_no=doc_no,
        date_str=date_str,
        page_str=page_str,
        title=title,
        iso_spec=iso_spec
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# API ENDPOINTS
# ============================================================

@router.get(
    "/header/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Header Word Document (.docx)"
)
async def generate_header_doc_get(
    centre_dept: str = Query("", description="Centre / Dept name"),
    doc_no: str = Query("", description="Document Number"),
    date_str: str = Query("13-01-2026", alias="date", description="Date string"),
    page_str: str = Query("1 of 2", alias="page", description="Page info"),
    title: str = Query("PROJECT TEAM LETTER-SMC", description="Document Title"),
    iso_spec: str = Query("ISO 9001-2015", description="ISO Spec"),
    filename: str = Query("CMTI_Project_Team_Header.docx", description="Output filename")
):
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_iso_header_bytes(
        centre_dept=centre_dept,
        doc_no=doc_no,
        date_str=date_str,
        page_str=page_str,
        title=title,
        iso_spec=iso_spec
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/header/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Header Word Document (.docx) via POST"
)
async def generate_header_doc_post(payload: HeaderInfoRequest):
    filename = payload.filename or "CMTI_Project_Team_Header.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_iso_header_bytes(
        centre_dept=payload.centre_dept,
        doc_no=payload.doc_no,
        date_str=payload.date,
        page_str=payload.page,
        title=payload.title,
        iso_spec=payload.iso_spec
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


# ============================================================
# MAIN SCRIPT EXECUTION (Standalone Test)
# ============================================================

if __name__ == "__main__":
    doc = create_iso_header_document()
    output_filename = "CMTI_Project_Team_Header.docx"
    doc.save(output_filename)
    print(f"Document created successfully: {output_filename}")