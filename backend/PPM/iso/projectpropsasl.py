import os
import io
import re
from datetime import datetime
from typing import Optional, List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel
from fastapi import APIRouter, status, Query, HTTPException
from fastapi.responses import StreamingResponse
from iso.header import add_header_table, normalize_centre_dept
from iso.finalfooter import add_footer_table

# ============================================================
# FASTAPI ROUTER
# ============================================================

router = APIRouter(prefix="/iso", tags=["ISO Project Proposal Generator"])


class BudgetItemRequest(BaseModel):
    sl_no: str = ""
    item_type: str = "Recurring"  # "Recurring" or "Non-Recurring"
    items: str = ""
    budget_amount: str = ""
    remarks: str = ""


class TimelineTaskRequest(BaseModel):
    sl_no: int
    task_name: str = ""
    active_months: Optional[List[int]] = None  # e.g. [1, 2] for month 1 and 2


class EquipmentItemRequest(BaseModel):
    sl_no: int
    technical_name: str = ""
    key_specifications: str = ""
    estimated_cost: str = ""
    amc_required: str = ""
    utilization_plan: str = ""


class ProjectProposalRequest(BaseModel):
    title_of_project: str = ""
    project_no: str = ""
    project_category: str = ""
    sponsoring_agency: str = ""
    sanction_order: str = ""
    total_cost: str = ""
    
    project_leader: str = ""
    co_leaders: str = ""
    core_st_members: Optional[List[str]] = None
    
    dev_partners_name: str = ""
    dev_partners_roles: str = ""
    
    commencement_date: str = ""
    completion_date: str = ""
    
    proposed_objectives: Optional[List[str]] = None
    current_status: str = ""
    
    research_tasks: Optional[List[str]] = None
    timeline_tasks: Optional[List[TimelineTaskRequest]] = None
    task_active_months: Optional[Dict[str, Any]] = None
    taskActiveMonths: Optional[Dict[str, Any]] = None
    
    salient_achievements: str = ""
    expected_trl: str = ""
    ipr_details: str = ""
    human_resources: Optional[List[str]] = None
    revenue_generated: str = ""
    
    recurring_budget: Optional[List[BudgetItemRequest]] = None
    non_recurring_budget: Optional[List[BudgetItemRequest]] = None
    equipment_details: Optional[List[EquipmentItemRequest]] = None
    infrastructure_details: str = ""
    
    prepared_by: Optional[str] = ""
    approved_by: Optional[str] = ""
    group_name: Optional[str] = ""
    centre_dept: Optional[str] = ""
    doc_no: Optional[str] = ""
    doc_date: Optional[str] = ""
    filename: Optional[str] = "CMTI_Project_Proposal.docx"


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def format_db_date(date_obj, date_format="%d-%m-%Y") -> str:
    """Safely format database date fields."""
    if not date_obj:
        return ""
    if isinstance(date_obj, str):
        return date_obj
    try:
        return date_obj.strftime(date_format)
    except Exception:
        return str(date_obj)


def set_cell_margins(cell, top=50, start=50, bottom=50, end=50):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")

    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")

    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_shading(cell, fill):
    """Set background color of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches):
    """Set fixed cell width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_row_height(row, height_inches, rule="atLeast"):
    """Set row height with rule."""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_inches * 1440)))
    trHeight.set(qn("w:hRule"), rule)
    trPr.append(trHeight)


def add_text(
    cell_or_paragraph,
    text,
    font_size=10,
    bold=False,
    italic=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    font_name="Arial",
    color=None,
    space_before=0,
    space_after=0
):
    """Utility helper to append formatted text into cell or paragraph with bold tag support (**text**, <b>text</b>, [b]text[/b])."""
    if hasattr(cell_or_paragraph, "paragraphs"):
        paragraph = cell_or_paragraph.paragraphs[0]
    else:
        paragraph = cell_or_paragraph

    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.15

    lines = str(text or "").split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            paragraph.add_run("\n")

        # Parse bold formatting tags: **bold**, <b>bold</b>, or [b]bold[/b]
        tokens = re.split(r"(\*\*.*?\*\*|<b>.*?</b>|\[b\].*?\[/b\])", line)
        for token in tokens:
            if not token:
                continue
            token_bold = bold
            clean_token = token
            if token.startswith("**") and token.endswith("**") and len(token) >= 4:
                token_bold = True
                clean_token = token[2:-2]
            elif token.startswith("<b>") and token.endswith("</b>") and len(token) >= 7:
                token_bold = True
                clean_token = token[3:-4]
            elif token.startswith("[b]") and token.endswith("[/b]") and len(token) >= 7:
                token_bold = True
                clean_token = token[3:-4]

            run = paragraph.add_run(clean_token)
            run.bold = token_bold
            run.italic = italic
            run.font.name = font_name
            run.font.size = Pt(font_size)
            if isinstance(color, str):
                run.font.color.rgb = RGBColor.from_string(color)
            elif isinstance(color, RGBColor):
                run.font.color.rgb = color
            run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)

def calculate_duration_months(commencement_date_str, completion_date_str):
    """Calculates total duration in months between commencement and completion dates."""
    if not commencement_date_str or not completion_date_str:
        return 6

    start_str = str(commencement_date_str).strip()
    end_str = str(completion_date_str).strip()

    formats = [
        "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y",
        "%d-%b-%Y", "%b %Y", "%m-%Y", "%Y-%m"
    ]

    start_dt = None
    end_dt = None

    for fmt in formats:
        try:
            start_dt = datetime.strptime(start_str, fmt)
            break
        except Exception:
            pass

    for fmt in formats:
        try:
            end_dt = datetime.strptime(end_str, fmt)
            break
        except Exception:
            pass

    if start_dt and end_dt and end_dt >= start_dt:
        num_months = (end_dt.year - start_dt.year) * 12 + (end_dt.month - start_dt.month) + 1
        return max(num_months, 1)

    try:
        match = re.search(r'(\d+)\s*months?', end_str, re.IGNORECASE)
        if match:
            return int(match.group(1))
    except Exception:
        pass

    return 6


# ============================================================
# MAIN DOCUMENT GENERATION FUNCTION
# ============================================================

def create_project_proposal_document(
    title_of_project: str = "",
    project_no: str = "",
    project_category: str = "",
    sponsoring_agency: str = "",
    sanction_order: str = "",
    total_cost: str = "",
    project_leader: str = "",
    co_leaders: str = "",
    core_st_members: Optional[List[str]] = None,
    dev_partners_name: str = "",
    dev_partners_roles: str = "",
    commencement_date: str = "",
    completion_date: str = "",
    proposed_objectives: Optional[List[str]] = None,
    current_status: str = "",
    research_tasks: Optional[List[str]] = None,
    task_active_months: Optional[Dict[str, Any]] = None,
    timeline_tasks: Optional[List[TimelineTaskRequest]] = None,
    salient_achievements: str = "",
    expected_trl: str = "",
    ipr_details: str = "",
    human_resources: Optional[List[str]] = None,
    revenue_generated: str = "",
    recurring_budget: Optional[List[BudgetItemRequest]] = None,
    non_recurring_budget: Optional[List[BudgetItemRequest]] = None,
    equipment_details: Optional[List[EquipmentItemRequest]] = None,
    infrastructure_details: str = "",
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "",
    doc_date: str = ""
) -> Document:
    doc = Document()
    section = doc.sections[0]

    # Margins
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Dynamic header & footer as required by project standard
    group_upper = (group_name or "SMC").strip().upper()
    header_group = group_upper
    if "CMTI-QMS" in group_upper:
        parts = group_upper.split("-")
        header_group = parts[-1] if len(parts) > 1 else group_upper

    add_header_table(
        section,
        title="PROJECT PROPOSAL",
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=doc_no,
        date_str=doc_date
    )

    core_st_members = core_st_members or []
    proposed_objectives = proposed_objectives or []
    research_tasks = research_tasks or []
    human_resources = human_resources or []
    recurring_budget = recurring_budget or []
    non_recurring_budget = non_recurring_budget or []
    equipment_details = equipment_details or []

    border_format = {"val": "single", "sz": 4, "color": "000000"}

    # Top Heading
    p_title = doc.add_paragraph()
    add_text(p_title, "FORMAT FOR PROJECT PROPOSAL", font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)

    p_sub = doc.add_paragraph()
    add_text(p_sub, "(Required to be submit for all categories of projects)", font_size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)

    # 1. Title of the Project
    p_proj_title = doc.add_paragraph()
    add_text(p_proj_title, f"1. Title of the Project: {title_of_project}", font_size=11, bold=True, space_after=8)

    # 2. Project Details
    p_proj_details = doc.add_paragraph()
    add_text(p_proj_details, "2. Project Details", font_size=11, bold=True, space_after=4)

    p_num = doc.add_paragraph()
    add_text(p_num, f"Project Number (for use of PPM): {project_no}", font_size=10, space_after=4)

    p_cat = doc.add_paragraph()
    add_text(p_cat, f"Project Category: DIP / GSP / ISP / GAP / CLP / ICP / AIP / LSP / ILP : {project_category or '-'}", font_size=10, space_after=4)

    p_agency = doc.add_paragraph()
    add_text(p_agency, f"Sponsoring Agency/Industry (in case other than ILP): {sponsoring_agency}", font_size=10, space_after=4)

    p_sanction = doc.add_paragraph()
    add_text(p_sanction, f"Original Sanction Order of the Sponsoring Agency: {sanction_order}", font_size=10, space_after=4)

    p_cost = doc.add_paragraph()
    add_text(p_cost, f"3. Total cost of the project (Rs. in Lakh). {total_cost}", font_size=10, bold=True, space_after=2)
    p_cost_note = doc.add_paragraph()
    add_text(p_cost_note, "\t(Please provide head wise details as per sanction order)", font_size=9, italic=True, space_after=6)

    p_fin = doc.add_paragraph()
    add_text(p_fin, "4. Financial proposal (Rs. in lakhs).", font_size=10, bold=True, space_after=2)
    p_fin_note = doc.add_paragraph()
    add_text(p_fin_note, "\t(Please provide headwise and yearwise as per the project cost estimation format)", font_size=9, italic=True, space_after=6)

    # FINANCIAL PROPOSAL TABLE (Table 0)
    total_a = 0
    for item in recurring_budget:
        try:
            total_a += float(re.sub(r"[^\d.]", "", item.budget_amount))
        except Exception:
            pass

    total_b = 0
    for item in non_recurring_budget:
        try:
            total_b += float(re.sub(r"[^\d.]", "", item.budget_amount))
        except Exception:
            pass

    grand_total = total_a + total_b

    # Construct rows count
    total_rows = 1 + 1 + 1 + max(len(recurring_budget), 1) + 1 + 1 + max(len(non_recurring_budget), 1) + 1 + 1
    budget_table = doc.add_table(rows=total_rows, cols=4)
    budget_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    budget_table.autofit = False

    budget_widths = [0.6, 2.8, 1.2, 2.2]
    for row in budget_table.rows:
        for c_idx, w in enumerate(budget_widths):
            set_cell_width(row.cells[c_idx], w)

    # Format cells
    for row in budget_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=40, start=60, bottom=40, end=60)

    # Header Rows
    add_text(budget_table.cell(0, 0), "Sl. No", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(budget_table.cell(0, 1), "Items", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(budget_table.cell(0, 2), "Budget (in Rs.)", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(budget_table.cell(0, 3), "Remarks/Justification", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    curr_row = 1

    # Section A Header
    cell_a = budget_table.cell(curr_row, 0).merge(budget_table.cell(curr_row, 3))
    add_text(cell_a, "A. Recurring", font_size=10, bold=True)
    set_cell_shading(cell_a, "E8E8E8")
    curr_row += 1

    # Recurring Items
    if recurring_budget:
        for item in recurring_budget:
            add_text(budget_table.cell(curr_row, 0), item.sl_no, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(budget_table.cell(curr_row, 1), item.items, font_size=9)
            add_text(budget_table.cell(curr_row, 2), item.budget_amount, font_size=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            add_text(budget_table.cell(curr_row, 3), item.remarks, font_size=9)
            curr_row += 1
    else:
        add_text(budget_table.cell(curr_row, 0), "-", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(budget_table.cell(curr_row, 1), "-", font_size=9)
        add_text(budget_table.cell(curr_row, 2), "-", font_size=9)
        add_text(budget_table.cell(curr_row, 3), "-", font_size=9)
        curr_row += 1

    # Total (A)
    budget_table.cell(curr_row, 0).merge(budget_table.cell(curr_row, 1))
    add_text(budget_table.cell(curr_row, 0), "Total (A)", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(budget_table.cell(curr_row, 2), str(int(total_a) if total_a == int(total_a) else f"{total_a:.2f}") if total_a > 0 else "-", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(budget_table.cell(curr_row, 3), "", font_size=9)
    curr_row += 1

    # Section B Header
    cell_b = budget_table.cell(curr_row, 0).merge(budget_table.cell(curr_row, 3))
    add_text(cell_b, "B. Non-Recurring", font_size=10, bold=True)
    set_cell_shading(cell_b, "E8E8E8")
    curr_row += 1

    # Non-Recurring Items
    if non_recurring_budget:
        for item in non_recurring_budget:
            add_text(budget_table.cell(curr_row, 0), item.sl_no, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(budget_table.cell(curr_row, 1), item.items, font_size=9)
            add_text(budget_table.cell(curr_row, 2), item.budget_amount, font_size=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            add_text(budget_table.cell(curr_row, 3), item.remarks, font_size=9)
            curr_row += 1
    else:
        add_text(budget_table.cell(curr_row, 0), "-", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(budget_table.cell(curr_row, 1), "-", font_size=9)
        add_text(budget_table.cell(curr_row, 2), "-", font_size=9)
        add_text(budget_table.cell(curr_row, 3), "-", font_size=9)
        curr_row += 1

    # Total (B)
    budget_table.cell(curr_row, 0).merge(budget_table.cell(curr_row, 1))
    add_text(budget_table.cell(curr_row, 0), "Total (B)", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(budget_table.cell(curr_row, 2), str(int(total_b) if total_b == int(total_b) else f"{total_b:.2f}") if total_b > 0 else "-", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(budget_table.cell(curr_row, 3), "", font_size=9)
    curr_row += 1

    # Grand Total (A+B)
    budget_table.cell(curr_row, 0).merge(budget_table.cell(curr_row, 1))
    add_text(budget_table.cell(curr_row, 0), "Grand Total (A+B)", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(budget_table.cell(curr_row, 2), str(int(grand_total) if grand_total == int(grand_total) else f"{grand_total:.2f}") if grand_total > 0 else "-", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(budget_table.cell(curr_row, 3), "", font_size=9)
    set_cell_shading(budget_table.cell(curr_row, 0), "F0F0F0")
    set_cell_shading(budget_table.cell(curr_row, 2), "F0F0F0")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 5. Project Leader and Co-leaders
    p_leaders_hdr = doc.add_paragraph()
    add_text(p_leaders_hdr, "5. Project Leader and Co-leaders (if any)", font_size=11, bold=True, space_after=4)

    p_pleader = doc.add_paragraph()
    add_text(p_pleader, f"Principal Coordinator/ Leader/ Investigator: {project_leader}", font_size=10, space_after=4)

    p_coleader = doc.add_paragraph()
    add_text(p_coleader, f"Co-leaders: {co_leaders}", font_size=10, space_after=6)

    p_core_st_hdr = doc.add_paragraph()
    add_text(p_core_st_hdr, "Core S&T members (Scientist-B and above): (it is mandatory to include the core team on project completion report)", font_size=10, bold=True, space_after=4)

    if core_st_members:
        for member in core_st_members:
            p_mem = doc.add_paragraph()
            add_text(p_mem, member, font_size=10, space_after=2)
    else:
        p_mem = doc.add_paragraph()
        add_text(p_mem, "-", font_size=10, space_after=2)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 6. Participating institute/collaborators
    p_part_hdr = doc.add_paragraph()
    add_text(p_part_hdr, "6. Participating (development partner if any) institute/collaborators/industry partners with their role of involvement/responsibility", font_size=11, bold=True, space_after=4)

    p_dev_name = doc.add_paragraph()
    add_text(p_dev_name, f"Name and address of development partners: {dev_partners_name}", font_size=10, space_after=4)

    p_dev_roles = doc.add_paragraph()
    add_text(p_dev_roles, f"Roles and responsibility (original as per MoU agreements/Revised if any): {dev_partners_roles}", font_size=10, space_after=8)

    # 7. Date of commencement and completion
    p_dates_hdr = doc.add_paragraph()
    add_text(p_dates_hdr, "7. Date of commencement and completion", font_size=11, bold=True, space_after=4)

    p_comm = doc.add_paragraph()
    add_text(p_comm, f"Date of commencement: {commencement_date}", font_size=10, space_after=4)

    p_comp = doc.add_paragraph()
    add_text(p_comp, f"Expected date of completion: {completion_date}", font_size=10, space_after=8)

    # 8. Proposed Objectives
    p_obj_hdr = doc.add_paragraph()
    add_text(p_obj_hdr, "8. Proposed Objectives", font_size=11, bold=True, space_after=4)
    if proposed_objectives:
        for obj in proposed_objectives:
            p_obj = doc.add_paragraph()
            add_text(p_obj, f"• {obj}", font_size=10, space_after=2)
    else:
        p_obj = doc.add_paragraph()
        add_text(p_obj, "-", font_size=10, space_after=4)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 9. Current status
    p_status_hdr = doc.add_paragraph()
    add_text(p_status_hdr, "9. Current domestic and international status (State of the art vis-a-vis knowledge gaps) (not more than a page", font_size=11, bold=True, space_after=4)

    p_status_body = doc.add_paragraph()
    add_text(p_status_body, current_status or "-", font_size=10, space_after=8)

    # 10. Research Methodology and Timeline
    p_res_hdr = doc.add_paragraph()
    add_text(p_res_hdr, "10. Research Methodology and Timeline (Research Tasks and timeline as planned) (include Gantt chart)", font_size=11, bold=True, space_after=6)

    # Research Tasks List / Table 1
    tasks_count = max(len(research_tasks), 1)
    tasks_table = doc.add_table(rows=tasks_count, cols=2)
    tasks_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tasks_table.autofit = False

    task_widths = [0.5, 6.3]
    for row in tasks_table.rows:
        for c_idx, w in enumerate(task_widths):
            set_cell_width(row.cells[c_idx], w)

    if research_tasks:
        for idx, task in enumerate(research_tasks):
            cell_num = tasks_table.cell(idx, 0)
            cell_text = tasks_table.cell(idx, 1)
            set_cell_border(cell_num, top=border_format, bottom=border_format, left=border_format, right=border_format)
            set_cell_border(cell_text, top=border_format, bottom=border_format, left=border_format, right=border_format)
            set_cell_margins(cell_num, top=30, start=40, bottom=30, end=40)
            set_cell_margins(cell_text, top=30, start=40, bottom=30, end=40)
            add_text(cell_num, str(idx + 1), font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(cell_text, task, font_size=9)
    else:
        cell_num = tasks_table.cell(0, 0)
        cell_text = tasks_table.cell(0, 1)
        set_cell_border(cell_num, top=border_format, bottom=border_format, left=border_format, right=border_format)
        set_cell_border(cell_text, top=border_format, bottom=border_format, left=border_format, right=border_format)
        set_cell_margins(cell_num, top=30, start=40, bottom=30, end=40)
        set_cell_margins(cell_text, top=30, start=40, bottom=30, end=40)
        add_text(cell_num, "1", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(cell_text, "-", font_size=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p_t1 = doc.add_paragraph()
    add_text(p_t1, "Table 1. Timeline for the proposed tasks.", font_size=10, bold=True, space_after=6)

    # Timeline Gantt Grid Table (Table 2: rows x 7 cols)
    duration_months = calculate_duration_months(commencement_date, completion_date)
    gantt_cols = 1 + duration_months
    gantt_rows = max(len(research_tasks), 1) + 1
    gantt_table = doc.add_table(rows=gantt_rows, cols=gantt_cols)
    gantt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    gantt_table.autofit = False

    # Calculate column widths
    first_col_w = max(2.2, 6.8 - (duration_months * 0.25))
    if duration_months > 12:
        first_col_w = 2.0
    rem_w = max(0.15, (6.8 - first_col_w) / duration_months)

    for row in gantt_table.rows:
        set_cell_width(row.cells[0], first_col_w)
        for m in range(1, gantt_cols):
            set_cell_width(row.cells[m], rem_w)

    for row in gantt_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=20, start=20, bottom=20, end=20)

    # Header Row
    add_text(gantt_table.cell(0, 0), "Activities / Months", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    for m in range(1, gantt_cols):
        add_text(gantt_table.cell(0, m), str(m), font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(gantt_table.cell(0, 0), "F0F0F0")

    active_map = task_active_months or {}

    if research_tasks:
        for idx, task in enumerate(research_tasks):
            r_idx = idx + 1
            add_text(gantt_table.cell(r_idx, 0), f"{idx + 1}. {task}", font_size=8)

            str_idx = str(idx)
            val = None
            if isinstance(active_map, dict):
                val = active_map.get(str_idx)
                if val is None:
                    val = active_map.get(idx)
                if val is None:
                    val = active_map.get(f"task_{idx}")
                if val is None:
                    val = active_map.get(f"task_{str_idx}")

            if isinstance(val, list):
                active_m = [int(m) for m in val if str(m).isdigit()]
            elif isinstance(val, (set, tuple)):
                active_m = [int(m) for m in val if str(m).isdigit()]
            elif isinstance(val, str):
                active_m = [int(m.strip()) for m in val.split(",") if m.strip().isdigit()]
            else:
                active_m = []

            for m in range(1, gantt_cols):
                if m in active_m:
                    add_text(gantt_table.cell(r_idx, m), "X", font_size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    set_cell_shading(gantt_table.cell(r_idx, m), "D9D9D9")
    else:
        add_text(gantt_table.cell(1, 0), "1. -", font_size=8)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 11. Technical Performance
    p_tech_hdr = doc.add_paragraph()
    add_text(p_tech_hdr, "11. Technical Performance: Salient achievements envisaged summarizing the contributions towards technology development and research outputs", font_size=11, bold=True, space_after=2)
    p_tech_note = doc.add_paragraph()
    add_text(p_tech_note, "\t(Please provide key highlights / novelty/techno-economic benefits of developments)", font_size=9, italic=True, space_after=4)
    p_tech_body = doc.add_paragraph()
    add_text(p_tech_body, salient_achievements or "-", font_size=10, space_after=6)

    # Sub-heading under 11: Important Technology/Products
    p_trl_hdr = doc.add_paragraph()
    add_text(p_trl_hdr, "Important Technology/Products/Machines/Software Tools to be developed and expected TRL", font_size=10, bold=True, space_after=4)
    p_trl_body = doc.add_paragraph()
    add_text(p_trl_body, expected_trl or "-", font_size=10, space_after=6)

    # Sub-heading under 11: IPR
    p_ipr_hdr = doc.add_paragraph()
    add_text(p_ipr_hdr, "Intellectual Property Rights (IPR) to be developed (patents/copyrights/trademarks)", font_size=10, bold=True, space_after=4)
    p_ipr_body = doc.add_paragraph()
    add_text(p_ipr_body, ipr_details or "-", font_size=10, space_after=6)

    # Sub-heading under 11: Human resources
    p_hr_hdr = doc.add_paragraph()
    add_text(p_hr_hdr, "Human resources to be trained under this project (JRF/SRF/Project Fellow Ph.D’s / M.Tech. students)", font_size=10, bold=True, space_after=4)
    if human_resources:
        for hr in human_resources:
            p_hr_item = doc.add_paragraph()
            add_text(p_hr_item, hr, font_size=10, space_after=2)
    else:
        p_hr_item = doc.add_paragraph()
        add_text(p_hr_item, "-", font_size=10, space_after=2)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 12. Revenue/income generated
    p_rev_hdr = doc.add_paragraph()
    add_text(p_rev_hdr, "12. Revenue/income generated for the Institute (Provide details)", font_size=11, bold=True, space_after=4)
    p_rev_body = doc.add_paragraph()
    add_text(p_rev_body, revenue_generated or "-", font_size=10, space_after=8)

    # 13. Details of Equipments and instruments (Table 3)
    p_eq_hdr = doc.add_paragraph()
    add_text(p_eq_hdr, "13. Details of Equipments and instruments (Project head C-3, C-4)", font_size=11, bold=True, space_after=6)

    eq_rows = max(len(equipment_details), 1) + 1
    eq_table = doc.add_table(rows=eq_rows, cols=6)
    eq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    eq_table.autofit = False

    eq_widths = [0.5, 1.8, 2.2, 0.9, 0.8, 0.6]
    for row in eq_table.rows:
        for c_idx, w in enumerate(eq_widths):
            set_cell_width(row.cells[c_idx], w)

    for row in eq_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=40, start=50, bottom=40, end=50)

    # Equipment Table Header
    add_text(eq_table.cell(0, 0), "Sl No", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(eq_table.cell(0, 1), "Technical Name of the facility", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(eq_table.cell(0, 2), "Key specifications", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(eq_table.cell(0, 3), "Estimated Cost", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(eq_table.cell(0, 4), "AMC required (yes/no) & period", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(eq_table.cell(0, 5), "Utilization plan", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    for c_idx in range(6):
        set_cell_shading(eq_table.cell(0, c_idx), "F0F0F0")

    if equipment_details:
        for idx, eq in enumerate(equipment_details):
            r_idx = idx + 1
            add_text(eq_table.cell(r_idx, 0), str(eq.sl_no), font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(eq_table.cell(r_idx, 1), eq.technical_name, font_size=8)
            add_text(eq_table.cell(r_idx, 2), eq.key_specifications, font_size=8)
            add_text(eq_table.cell(r_idx, 3), eq.estimated_cost, font_size=8, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            add_text(eq_table.cell(r_idx, 4), eq.amc_required, font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(eq_table.cell(r_idx, 5), eq.utilization_plan, font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_text(eq_table.cell(1, 0), "1", font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(eq_table.cell(1, 1), "-", font_size=8)
        add_text(eq_table.cell(1, 2), "-", font_size=8)
        add_text(eq_table.cell(1, 3), "-", font_size=8)
        add_text(eq_table.cell(1, 4), "-", font_size=8)
        add_text(eq_table.cell(1, 5), "-", font_size=8)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 14. Details of Infrastructure proposed to be created
    p_infra_hdr = doc.add_paragraph()
    add_text(p_infra_hdr, "14. Details of Infrastructure proposed to be created (Project head C1)", font_size=11, bold=True, space_after=4)

    p_infra_body = doc.add_paragraph()
    add_text(p_infra_body, infrastructure_details or "-", font_size=10, space_after=16)

    # 16. Signatures Block
    p_sig1 = doc.add_paragraph()
    add_text(p_sig1, "(Signature of Project Leader and Co-leaders)", font_size=10, bold=True, space_after=24)

    p_sig2 = doc.add_paragraph()
    add_text(p_sig2, "(Signature of Center Head)\nJD & CH, C-SMPM", font_size=10, bold=True, space_after=24)

    p_sig3 = doc.add_paragraph()
    add_text(p_sig3, "Head, PP&BD:", font_size=10, bold=True, space_after=24)

    p_sig4 = doc.add_paragraph()
    add_text(p_sig4, "FA & CAO", font_size=10, bold=True, space_after=24)

    p_sig5 = doc.add_paragraph()
    add_text(p_sig5, "Director – for kind approval please", font_size=10, bold=True, space_after=12)

    # 17. Prepared By / Approved By ISO Footer Block (Only on last page body)
    add_footer_table(
        doc,
        prepared_name=prepared_by,
        approved_name=approved_by,
        group_name=header_group,
        doc_code="CMTI/PPBD/009/Rev-00",
        in_body=True
    )

    return doc


def generate_project_proposal_bytes(
    title_of_project: str = "",
    project_no: str = "",
    project_category: str = "",
    sponsoring_agency: str = "",
    sanction_order: str = "",
    total_cost: str = "",
    project_leader: str = "",
    co_leaders: str = "",
    core_st_members: Optional[List[str]] = None,
    dev_partners_name: str = "",
    dev_partners_roles: str = "",
    commencement_date: str = "",
    completion_date: str = "",
    proposed_objectives: Optional[List[str]] = None,
    current_status: str = "",
    research_tasks: Optional[List[str]] = None,
    task_active_months: Optional[Dict[str, Any]] = None,
    timeline_tasks: Optional[List[TimelineTaskRequest]] = None,
    salient_achievements: str = "",
    expected_trl: str = "",
    ipr_details: str = "",
    human_resources: Optional[List[str]] = None,
    revenue_generated: str = "",
    recurring_budget: Optional[List[BudgetItemRequest]] = None,
    non_recurring_budget: Optional[List[BudgetItemRequest]] = None,
    equipment_details: Optional[List[EquipmentItemRequest]] = None,
    infrastructure_details: str = "",
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "",
    doc_date: str = ""
) -> io.BytesIO:
    doc = create_project_proposal_document(
        title_of_project=title_of_project,
        project_no=project_no,
        project_category=project_category,
        sponsoring_agency=sponsoring_agency,
        sanction_order=sanction_order,
        total_cost=total_cost,
        project_leader=project_leader,
        co_leaders=co_leaders,
        core_st_members=core_st_members,
        dev_partners_name=dev_partners_name,
        dev_partners_roles=dev_partners_roles,
        commencement_date=commencement_date,
        completion_date=completion_date,
        proposed_objectives=proposed_objectives,
        current_status=current_status,
        research_tasks=research_tasks,
        task_active_months=task_active_months,
        timeline_tasks=timeline_tasks,
        salient_achievements=salient_achievements,
        expected_trl=expected_trl,
        ipr_details=ipr_details,
        human_resources=human_resources,
        revenue_generated=revenue_generated,
        recurring_budget=recurring_budget,
        non_recurring_budget=non_recurring_budget,
        equipment_details=equipment_details,
        infrastructure_details=infrastructure_details,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# API ROUTE HANDLERS
# ============================================================

@router.get(
    "/project-proposal/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Project Proposal (.docx) via GET"
)
@router.get(
    "/projectpropsasl/generate",
    status_code=status.HTTP_200_OK,
    include_in_schema=False
)
async def generate_project_proposal_doc_get(
    title_of_project: str = Query(""),
    project_no: str = Query(""),
    project_category: str = Query(""),
    sponsoring_agency: str = Query(""),
    total_cost: str = Query(""),
    project_leader: str = Query(""),
    prepared_by: str = Query(""),
    approved_by: str = Query(""),
    group_name: str = Query(""),
    centre_dept: str = Query(""),
    doc_no: str = Query(""),
    doc_date: str = Query(""),
    filename: str = Query("CMTI_Project_Proposal.docx")
):
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_project_proposal_bytes(
        title_of_project=title_of_project,
        project_no=project_no,
        project_category=project_category,
        sponsoring_agency=sponsoring_agency,
        total_cost=total_cost,
        project_leader=project_leader,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/project-proposal/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Project Proposal (.docx) via POST"
)
@router.post(
    "/projectpropsasl/generate",
    status_code=status.HTTP_200_OK,
    include_in_schema=False
)
async def generate_project_proposal_doc_post(
    payload: ProjectProposalRequest
):
    filename = payload.filename or "CMTI_Project_Proposal.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_project_proposal_bytes(
        title_of_project=payload.title_of_project,
        project_no=payload.project_no,
        project_category=payload.project_category,
        sponsoring_agency=payload.sponsoring_agency,
        sanction_order=payload.sanction_order,
        total_cost=payload.total_cost,
        project_leader=payload.project_leader,
        co_leaders=payload.co_leaders,
        core_st_members=payload.core_st_members,
        dev_partners_name=payload.dev_partners_name,
        dev_partners_roles=payload.dev_partners_roles,
        commencement_date=payload.commencement_date,
        completion_date=payload.completion_date,
        proposed_objectives=payload.proposed_objectives,
        current_status=payload.current_status,
        research_tasks=payload.research_tasks,
        task_active_months=payload.task_active_months or payload.taskActiveMonths,
        timeline_tasks=payload.timeline_tasks,
        salient_achievements=payload.salient_achievements,
        expected_trl=payload.expected_trl,
        ipr_details=payload.ipr_details,
        human_resources=payload.human_resources,
        revenue_generated=payload.revenue_generated,
        recurring_budget=payload.recurring_budget,
        non_recurring_budget=payload.non_recurring_budget,
        equipment_details=payload.equipment_details,
        infrastructure_details=payload.infrastructure_details,
        prepared_by=payload.prepared_by or "",
        approved_by=payload.approved_by or "",
        group_name=payload.group_name or "",
        centre_dept=payload.centre_dept or "",
        doc_no=payload.doc_no or "",
        doc_date=payload.doc_date or ""
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


# ============================================================
# STANDALONE TEST
# ============================================================

if __name__ == "__main__":
    doc = create_project_proposal_document()
    output_filename = "CMTI_Project_Proposal_Test.docx"
    doc.save(output_filename)
    print(f"Document created successfully: {output_filename}")
