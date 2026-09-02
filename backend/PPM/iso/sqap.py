"""
FastAPI Router & Document Generator for ISO Document 055: Software Quality Assurance Plan (SQAP).
Generates Word (.docx) document matching CMTI-QMS-055/Rev00 specification with Release & Acceptance Table and Project Info Table.
"""

from typing import List, Dict, Any, Optional
import io
from fastapi import APIRouter, HTTPException, status, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from iso.header import add_header_table
from iso.finalfooter import add_footer_table

router = APIRouter(prefix="/iso", tags=["ISO Software Quality Assurance Plan (Doc 055)"])

# Helper function to set cell background color
def set_cell_shading(cell, color_hex: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Helper function to set cell borders
def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "color", "sz", "space"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))

def set_cell_margins(cell, top=60, start=60, bottom=60, end=60):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('left', start), ('bottom', bottom), ('right', end)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_width(cell, width_inches: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")

def add_text(cell_or_paragraph, text: str, font_name: str = "Arial", font_size: int = 10, bold: bool = False, italic: bool = False, color: RGBColor = RGBColor(0, 0, 0), alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, space_after: int = 0):
    if hasattr(cell_or_paragraph, 'paragraphs'):
        p = cell_or_paragraph.paragraphs[0]
    else:
        p = cell_or_paragraph

    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15

    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


class SQAPUserAgencyItem(BaseModel):
    role: str = "Co-ordinated by"
    name: str = ""
    organization: Optional[str] = ""
    signature: str = ""


class SQAPRequest(BaseModel):
    project_title: str = ""
    customer_name: str = ""
    sanction_letter_no: str = ""
    project_no: str = ""
    software_version: Optional[str] = "v1.0"
    
    # Table 1: Release & Acceptance
    released_by_org: Optional[str] = "CMTI"
    user_agency_org: Optional[str] = ""
    prepared_by_name: Optional[str] = ""
    prepared_by_sig: Optional[str] = ""
    checked_by_name: Optional[str] = ""
    checked_by_sig: Optional[str] = ""
    approved_by_name: Optional[str] = ""
    approved_by_sig: Optional[str] = ""
    user_agency_rows: Optional[List[Dict[str, Any]]] = None
    
    # Standard ISO metadata
    prepared_by: Optional[str] = ""
    approved_by: Optional[str] = ""
    group_name: Optional[str] = ""
    centre_dept: Optional[str] = ""
    doc_no: Optional[str] = "055"
    doc_date: Optional[str] = ""
    filename: Optional[str] = "ISO_Software_Quality_Assurance_Plan.docx"


def create_sqap_document(
    project_title: str = "",
    customer_name: str = "",
    sanction_letter_no: str = "",
    project_no: str = "",
    software_version: str = "v1.0",
    released_by_org: str = "CMTI",
    user_agency_org: str = "",
    prepared_by_name: str = "",
    prepared_by_sig: str = "",
    checked_by_name: str = "",
    checked_by_sig: str = "",
    approved_by_name: str = "",
    approved_by_sig: str = "",
    user_agency_rows: Optional[List[Any]] = None,
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "055",
    doc_date: str = "",
    sections: Optional[Any] = None
) -> Document:
    doc = Document()

    # Setup margins (A4 Portrait)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    header_group = (group_name or centre_dept or "SMC").strip().upper()
    if header_group.startswith("G-"):
        header_group = header_group[2:]
    elif header_group.startswith("C-"):
        header_group = header_group[2:]

    final_doc_no = str(doc_no or "055").strip()
    if final_doc_no.isdigit():
        final_doc_no = final_doc_no.zfill(3)

    # Format Document Code for Header and Footer
    doc_code_str = f"CMTI-{header_group}-QMS-{final_doc_no}/Rev00" if header_group else f"CMTI-QMS-{final_doc_no}/Rev00"

    # Add Standard ISO Header Table
    add_header_table(
        doc.sections[0],
        title=f"SOFTWARE QUALITY ASSURANCE PLAN-{header_group}",
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=final_doc_no,
        date_str=doc_date
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # -------------------------------------------------------------
    # TABLE 1: Release & Acceptance Table (5 columns)
    # -------------------------------------------------------------
    prep_name = prepared_by_name or prepared_by or ""
    prep_sig = prepared_by_sig or ""
    chk_name = checked_by_name or ""
    chk_sig = checked_by_sig or ""
    app_name = approved_by_name or approved_by or ""
    app_sig = approved_by_sig or ""
    org_name = released_by_org or "CMTI"

    # Agency rows
    raw_agency = user_agency_rows
    if not raw_agency:
        agency_items = [
            {"role": "Co-ordinated by", "name": "", "org": "", "sig": ""},
            {"role": "Co-ordinated by", "name": "", "org": "", "sig": ""}
        ]
    else:
        agency_items = []
        for item in raw_agency:
            if isinstance(item, dict):
                agency_items.append({
                    "role": item.get("role") or "Co-ordinated by",
                    "name": item.get("name") or "",
                    "org": item.get("organization") or item.get("org") or "",
                    "sig": item.get("signature") or item.get("sig") or ""
                })
            else:
                agency_items.append({
                    "role": getattr(item, "role", "Co-ordinated by"),
                    "name": getattr(item, "name", ""),
                    "org": getattr(item, "organization", getattr(item, "org", "")),
                    "sig": getattr(item, "signature", getattr(item, "sig", ""))
                })
        if not agency_items:
            agency_items = [
                {"role": "Co-ordinated by", "name": "", "org": "", "sig": ""},
                {"role": "Co-ordinated by", "name": "", "org": "", "sig": ""}
            ]

    agency_count = max(1, len(agency_items))
    total_t1_rows = 1 + 3 + agency_count  # 1 header + 3 released by + agency rows

    table1 = doc.add_table(rows=total_t1_rows, cols=5)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False

    t1_widths = [1.4, 1.2, 1.8, 1.2, 1.4]  # total = 7.0 inches
    border_fmt = {"val": "single", "sz": "4", "color": "000000"}

    # Apply base borders, margins, widths, vertical center to all cells
    for row in table1.rows:
        for c_idx, cell in enumerate(row.cells):
            set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            set_cell_margins(cell, top=60, start=60, bottom=60, end=60)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if c_idx < len(t1_widths):
                set_cell_width(cell, t1_widths[c_idx])

    # Header Row (Row 0)
    add_text(table1.cell(0, 0), "", font_size=10, bold=True)
    add_text(table1.cell(0, 1), "", font_size=10, bold=True)
    add_text(table1.cell(0, 2), "Name", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(table1.cell(0, 3), "Organization", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(table1.cell(0, 4), "Signature", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 1: Prepared by
    add_text(table1.cell(1, 0), "Released By", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_text(table1.cell(1, 1), "Prepared by", font_size=10)
    add_text(table1.cell(1, 2), prep_name, font_size=10)
    add_text(table1.cell(1, 3), org_name, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(table1.cell(1, 4), prep_sig, font_size=10)

    # Row 2: Checked by
    add_text(table1.cell(2, 1), "Checked by", font_size=10)
    add_text(table1.cell(2, 2), chk_name, font_size=10)
    add_text(table1.cell(2, 4), chk_sig, font_size=10)

    # Row 3: Approved by
    add_text(table1.cell(3, 1), "Approved by", font_size=10)
    add_text(table1.cell(3, 2), app_name, font_size=10)
    add_text(table1.cell(3, 4), app_sig, font_size=10)

    # Merge Released By in col 0 (rows 1..3)
    table1.cell(1, 0).merge(table1.cell(3, 0))
    # Merge Organization in col 3 (rows 1..3)
    table1.cell(1, 3).merge(table1.cell(3, 3))

    # User Agency Rows (Row 4 onwards)
    add_text(table1.cell(4, 0), "Accepted by\nuser agency", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    u_org = user_agency_org or customer_name or (agency_items[0].get("org") if agency_items else "")
    add_text(table1.cell(4, 3), u_org, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for a_idx, item in enumerate(agency_items):
        r_idx = 4 + a_idx
        add_text(table1.cell(r_idx, 1), item.get("role", "Co-ordinated by"), font_size=10)
        add_text(table1.cell(r_idx, 2), item.get("name", ""), font_size=10)
        add_text(table1.cell(r_idx, 4), item.get("sig", ""), font_size=10)

    # Merge Accepted by user agency in col 0 (rows 4..4+agency_count-1)
    if agency_count > 1:
        table1.cell(4, 0).merge(table1.cell(4 + agency_count - 1, 0))
        # Merge Organization in col 3 (rows 4..4+agency_count-1)
        table1.cell(4, 3).merge(table1.cell(4 + agency_count - 1, 3))

    # Spacing between tables
    doc.add_paragraph().paragraph_format.space_after = Pt(28)

    # -------------------------------------------------------------
    # TABLE 2: Project Details Table (2 columns)
    # -------------------------------------------------------------
    table2 = doc.add_table(rows=3, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.autofit = False

    t2_widths = [2.4, 4.6]  # total = 7.0 inches

    for row in table2.rows:
        for c_idx, cell in enumerate(row.cells):
            set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            set_cell_margins(cell, top=60, start=60, bottom=60, end=60)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if c_idx < len(t2_widths):
                set_cell_width(cell, t2_widths[c_idx])

    # Row 0: Project
    add_text(table2.cell(0, 0), "Project", font_size=10, bold=True)
    add_text(table2.cell(0, 1), project_title or "", font_size=10)

    # Row 1: Customer
    add_text(table2.cell(1, 0), "Customer", font_size=10, bold=True)
    add_text(table2.cell(1, 1), customer_name or "", font_size=10)

    # Row 2: Project Sanction letter No.
    sanction_no = sanction_letter_no or project_no or ""
    add_text(table2.cell(2, 0), "Project Sanction letter No.", font_size=10, bold=True)
    add_text(table2.cell(2, 1), sanction_no, font_size=10)

    # -------------------------------------------------------------
    # FOOTER: Document Number & Revision Line (Prints once in Page Footer)
    # -------------------------------------------------------------
    section_footer = doc.sections[0].footer
    footer_p = section_footer.paragraphs[0] if section_footer.paragraphs else section_footer.add_paragraph()
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    footer_p.paragraph_format.line_spacing = 1.0

    r_f = footer_p.add_run(doc_code_str)
    r_f.font.name = "Arial"
    r_f.font.size = Pt(9)
    r_f.font.color.rgb = RGBColor(60, 60, 60)

    return doc


def generate_sqap_bytes(
    project_title: str = "",
    customer_name: str = "",
    sanction_letter_no: str = "",
    project_no: str = "",
    software_version: str = "v1.0",
    released_by_org: str = "CMTI",
    user_agency_org: str = "",
    prepared_by_name: str = "",
    prepared_by_sig: str = "",
    checked_by_name: str = "",
    checked_by_sig: str = "",
    approved_by_name: str = "",
    approved_by_sig: str = "",
    user_agency_rows: Optional[List[Any]] = None,
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "055",
    doc_date: str = ""
) -> io.BytesIO:
    doc = create_sqap_document(
        project_title=project_title,
        customer_name=customer_name,
        sanction_letter_no=sanction_letter_no,
        project_no=project_no,
        software_version=software_version,
        released_by_org=released_by_org,
        user_agency_org=user_agency_org,
        prepared_by_name=prepared_by_name,
        prepared_by_sig=prepared_by_sig,
        checked_by_name=checked_by_name,
        checked_by_sig=checked_by_sig,
        approved_by_name=approved_by_name,
        approved_by_sig=approved_by_sig,
        user_agency_rows=user_agency_rows,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# API ROUTE HANDLERS
# ============================================================

@router.get(
    "/sqap/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO SQAP Document (.docx) via GET"
)
async def generate_sqap_doc_get(
    project_title: str = Query(""),
    customer_name: str = Query(""),
    sanction_letter_no: str = Query(""),
    project_no: str = Query(""),
    software_version: str = Query("v1.0"),
    released_by_org: str = Query("CMTI"),
    user_agency_org: str = Query(""),
    prepared_by_name: str = Query(""),
    prepared_by_sig: str = Query(""),
    checked_by_name: str = Query(""),
    checked_by_sig: str = Query(""),
    approved_by_name: str = Query(""),
    approved_by_sig: str = Query(""),
    prepared_by: str = Query(""),
    approved_by: str = Query(""),
    group_name: str = Query(""),
    centre_dept: str = Query(""),
    doc_no: str = Query("055"),
    doc_date: str = Query(""),
    filename: str = Query("ISO_Software_Quality_Assurance_Plan.docx")
):
    buffer = generate_sqap_bytes(
        project_title=project_title,
        customer_name=customer_name,
        sanction_letter_no=sanction_letter_no,
        project_no=project_no,
        software_version=software_version,
        released_by_org=released_by_org,
        user_agency_org=user_agency_org,
        prepared_by_name=prepared_by_name,
        prepared_by_sig=prepared_by_sig,
        checked_by_name=checked_by_name,
        checked_by_sig=checked_by_sig,
        approved_by_name=approved_by_name,
        approved_by_sig=approved_by_sig,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/sqap/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO SQAP Document (.docx) via POST"
)
async def generate_sqap_doc_post(
    payload: SQAPRequest
):
    filename = payload.filename or "ISO_Software_Quality_Assurance_Plan.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_sqap_bytes(
        project_title=payload.project_title,
        customer_name=payload.customer_name,
        sanction_letter_no=payload.sanction_letter_no,
        project_no=payload.project_no,
        software_version=payload.software_version or "v1.0",
        released_by_org=payload.released_by_org or "CMTI",
        user_agency_org=payload.user_agency_org or "",
        prepared_by_name=payload.prepared_by_name or payload.prepared_by or "",
        prepared_by_sig=payload.prepared_by_sig or "",
        checked_by_name=payload.checked_by_name or "",
        checked_by_sig=payload.checked_by_sig or "",
        approved_by_name=payload.approved_by_name or payload.approved_by or "",
        approved_by_sig=payload.approved_by_sig or "",
        user_agency_rows=payload.user_agency_rows,
        prepared_by=payload.prepared_by,
        approved_by=payload.approved_by,
        group_name=payload.group_name,
        centre_dept=payload.centre_dept,
        doc_no=payload.doc_no or "055",
        doc_date=payload.doc_date
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )
