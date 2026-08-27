"""
Universal / Generic ISO Document Generator.
Generates clean, standardized Word (.docx) documents for ANY custom or user-defined ISO document.
Supports custom tables (dynamic headers and rows), sections, key-value checklists, conclusions, headers, and footers.
"""

from typing import List, Dict, Any, Optional
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from iso.header import add_header_table
from iso.finalfooter import add_footer_table
from iso.sqap import set_cell_shading, set_cell_border, set_cell_margins, add_text


def create_generic_iso_document(
    doc_title: str = "ISO DOCUMENT",
    doc_code: str = "",
    doc_no: str = "",
    doc_date: str = "",
    centre_dept: str = "SMPM",
    group_name: str = "SMPM",
    project_title: str = "",
    project_no: str = "",
    customer_name: str = "",
    description: str = "",
    custom_headers: Optional[List[str]] = None,
    custom_rows: Optional[List[List[str]]] = None,
    sections: Optional[List[Dict[str, Any]]] = None,
    checklist_points: Optional[List[Dict[str, Any]]] = None,
    conclusion: str = "",
    prepared_by: str = "",
    approved_by: str = ""
) -> Document:
    doc = Document()

    # Determine page orientation: Landscape if more than 5 table columns
    num_cols = len(custom_headers) if custom_headers else 0
    is_wide = num_cols >= 5

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        if is_wide:
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
        else:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)

    clean_title = (doc_title or "ISO DOCUMENT").upper()
    header_group = (group_name or centre_dept or "SMPM").strip().upper()
    if header_group.startswith("G-") or header_group.startswith("C-"):
        header_group = header_group[2:]

    # Header Table
    add_header_table(
        doc.sections[0],
        title=f"{clean_title}-{header_group}" if header_group else clean_title,
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=doc_no or doc_code or "ISO",
        date_str=doc_date
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1. Project Information Box (if available)
    if project_title or project_no or customer_name or description:
        proj_table = doc.add_table(rows=2, cols=2)
        proj_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        proj_table.autofit = False

        c_w = Inches(5.2) if is_wide else Inches(3.6)
        for row in proj_table.rows:
            for cell in row.cells:
                cell.width = c_w

        set_cell_margins(proj_table.cell(0, 0), top=60, bottom=60, left=100, right=100)
        set_cell_margins(proj_table.cell(0, 1), top=60, bottom=60, left=100, right=100)
        set_cell_margins(proj_table.cell(1, 0), top=60, bottom=60, left=100, right=100)
        set_cell_margins(proj_table.cell(1, 1), top=60, bottom=60, left=100, right=100)

        set_cell_shading(proj_table.cell(0, 0), "F0F4F8")
        set_cell_shading(proj_table.cell(0, 1), "F0F4F8")
        set_cell_shading(proj_table.cell(1, 0), "F0F4F8")
        set_cell_shading(proj_table.cell(1, 1), "F0F4F8")

        for r_idx in range(2):
            for c_idx in range(2):
                set_cell_border(proj_table.cell(r_idx, c_idx), top="CCCCCC", bottom="CCCCCC", left="CCCCCC", right="CCCCCC")

        add_text(proj_table.cell(0, 0), "Project Title: ", bold=True, size_pt=9.5)
        add_text(proj_table.cell(0, 0), project_title or "-", bold=False, size_pt=9.5)

        add_text(proj_table.cell(0, 1), "Project No: ", bold=True, size_pt=9.5)
        add_text(proj_table.cell(0, 1), project_no or "-", bold=False, size_pt=9.5)

        add_text(proj_table.cell(1, 0), "Customer: ", bold=True, size_pt=9.5)
        add_text(proj_table.cell(1, 0), customer_name or "-", bold=False, size_pt=9.5)

        add_text(proj_table.cell(1, 1), "Doc Code / No: ", bold=True, size_pt=9.5)
        add_text(proj_table.cell(1, 1), f"{doc_code or doc_no or '-'}", bold=False, size_pt=9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 2. Main Data Table (if headers and rows provided)
    if custom_headers and len(custom_headers) > 0:
        total_cols = len(custom_headers)
        table_rows = (custom_rows if custom_rows else [])
        t_rows_cnt = max(len(table_rows), 1)

        table = doc.add_table(rows=1 + t_rows_cnt, cols=total_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Width calculation
        avail_width = 10.5 if is_wide else 7.2
        col_w = Inches(avail_width / total_cols)

        # Header row
        hdr_row = table.rows[0]
        for col_idx, col_name in enumerate(custom_headers):
            cell = hdr_row.cells[col_idx]
            cell.width = col_w
            set_cell_shading(cell, "003366")
            set_cell_margins(cell, top=70, bottom=70, left=80, right=80)
            set_cell_border(cell, top="003366", bottom="003366", left="003366", right="003366")
            add_text(cell, str(col_name).strip(), bold=True, color_rgb=(255, 255, 255), size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Data rows
        for r_idx in range(t_rows_cnt):
            row_cells = table.rows[1 + r_idx].cells
            row_data = table_rows[r_idx] if r_idx < len(table_rows) else []
            bg_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"

            for c_idx in range(total_cols):
                cell = row_cells[c_idx]
                cell.width = col_w
                set_cell_shading(cell, bg_color)
                set_cell_margins(cell, top=50, bottom=50, left=70, right=70)
                set_cell_border(cell, top="CCCCCC", bottom="CCCCCC", left="CCCCCC", right="CCCCCC")
                val = str(row_data[c_idx]) if (isinstance(row_data, list) and c_idx < len(row_data)) else ""
                add_text(cell, val or "-", bold=False, size_pt=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 3. Dynamic Checklist / Review Points (if any)
    if checklist_points and len(checklist_points) > 0:
        cl_p = doc.add_paragraph()
        cl_p.paragraph_format.space_before = Pt(6)
        cl_p.paragraph_format.space_after = Pt(2)
        r = cl_p.add_run("Checklist / Review Points:")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0, 51, 102)

        cl_table = doc.add_table(rows=1 + len(checklist_points), cols=4)
        cl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cl_table.autofit = False

        cl_w = [Inches(0.6), Inches(5.0) if is_wide else Inches(3.4), Inches(1.2), Inches(3.7) if is_wide else Inches(2.0)]
        cl_headers = ["Sl No", "Review / Checklist Point", "Response (Yes/No/NA)", "Details / Remarks"]

        for c_idx, h_text in enumerate(cl_headers):
            cell = cl_table.rows[0].cells[c_idx]
            cell.width = cl_w[c_idx]
            set_cell_shading(cell, "003366")
            set_cell_margins(cell, top=60, bottom=60, left=70, right=70)
            set_cell_border(cell, top="003366", bottom="003366", left="003366", right="003366")
            add_text(cell, h_text, bold=True, color_rgb=(255, 255, 255), size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        for pt_idx, pt in enumerate(checklist_points):
            row_cells = cl_table.rows[1 + pt_idx].cells
            bg = "F9FAFB" if pt_idx % 2 == 1 else "FFFFFF"

            for c_idx in range(4):
                cell = row_cells[c_idx]
                cell.width = cl_w[c_idx]
                set_cell_shading(cell, bg)
                set_cell_margins(cell, top=45, bottom=45, left=60, right=60)
                set_cell_border(cell, top="CCCCCC", bottom="CCCCCC", left="CCCCCC", right="CCCCCC")

            add_text(row_cells[0], str(pt.get("sl_no") or (pt_idx + 1)), bold=False, size_pt=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(row_cells[1], pt.get("point") or pt.get("review_point") or pt.get("description") or "-", bold=False, size_pt=8.5)
            add_text(row_cells[2], pt.get("response") or pt.get("yes_no_na") or "-", bold=True, size_pt=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(row_cells[3], pt.get("details") or pt.get("remarks") or "-", bold=False, size_pt=8.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 4. Custom Sections (if any)
    if sections and isinstance(sections, list):
        for sec in sections:
            sec_title = sec.get("title", "").strip()
            sec_content = sec.get("content", "").strip()
            sec_headers = sec.get("headers")
            sec_rows = sec.get("rows")

            if sec_title:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(sec_title)
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0, 51, 102)

            if sec_content:
                p_c = doc.add_paragraph()
                p_c.paragraph_format.space_after = Pt(4)
                r_c = p_c.add_run(sec_content)
                r_c.font.size = Pt(9)

            if sec_headers and len(sec_headers) > 0:
                s_cols = len(sec_headers)
                s_rows_data = sec_rows if sec_rows else []
                s_rows_cnt = max(len(s_rows_data), 1)

                s_table = doc.add_table(rows=1 + s_rows_cnt, cols=s_cols)
                s_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                s_table.autofit = False

                col_w_s = Inches((10.5 if is_wide else 7.2) / s_cols)

                for c_i, h_name in enumerate(sec_headers):
                    cell = s_table.rows[0].cells[c_i]
                    cell.width = col_w_s
                    set_cell_shading(cell, "003366")
                    set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
                    set_cell_border(cell, top="003366", bottom="003366", left="003366", right="003366")
                    add_text(cell, str(h_name).strip(), bold=True, color_rgb=(255, 255, 255), size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)

                for r_i in range(s_rows_cnt):
                    cells = s_table.rows[1 + r_i].cells
                    r_d = s_rows_data[r_i] if r_i < len(s_rows_data) else []
                    bg_s = "F9FAFB" if r_i % 2 == 1 else "FFFFFF"

                    for c_i in range(s_cols):
                        cell = cells[c_i]
                        cell.width = col_w_s
                        set_cell_shading(cell, bg_s)
                        set_cell_margins(cell, top=45, bottom=45, left=60, right=60)
                        set_cell_border(cell, top="CCCCCC", bottom="CCCCCC", left="CCCCCC", right="CCCCCC")
                        val = str(r_d[c_i]) if (isinstance(r_d, list) and c_i < len(r_d)) else ""
                        add_text(cell, val or "-", bold=False, size_pt=8.5)

                doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 5. Conclusion / Remarks Box (if any)
    if conclusion:
        conc_table = doc.add_table(rows=1, cols=1)
        conc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        conc_table.autofit = False
        c_cell = conc_table.cell(0, 0)
        c_cell.width = Inches(10.5) if is_wide else Inches(7.2)
        set_cell_shading(c_cell, "F0F7FF")
        set_cell_margins(c_cell, top=60, bottom=60, left=100, right=100)
        set_cell_border(c_cell, top="003366", bottom="003366", left="003366", right="003366")
        add_text(c_cell, "Conclusion / Remarks: ", bold=True, color_rgb=(0, 51, 102), size_pt=9.5)
        add_text(c_cell, conclusion, bold=False, size_pt=9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 6. Standard Final Footer Table
    add_footer_table(
        doc,
        prepared_name=prepared_by or "-",
        approved_name=approved_by or "-",
        group_name=group_name or centre_dept or "SMPM",
        doc_code=doc_code or doc_no or "ISO",
        in_body=True
    )

    return doc
