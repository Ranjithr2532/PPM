import os
import io
from typing import Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel
from fastapi import APIRouter, status, Query, HTTPException, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from db import get_db
from models.model import Proposal
from iso.header import add_header_table
from iso.finalfooter import add_footer_table

# ============================================================
# FASTAPI ROUTER
# ============================================================

router = APIRouter(prefix="/iso", tags=["ISO Minutes of Meeting Generator"])


class SummaryPointRequest(BaseModel):
    sl_no: int
    points_discussed: str = ""
    responsibility: str = ""


class MomRequest(BaseModel):
    project_id: Optional[int] = None
    meeting_date_time: str = ""
    meeting_location: str = ""
    prev_mom_no_date: str = "-"
    prev_action_points: str = "-"
    prev_status: str = "-"
    agenda: str = "Project kick off meeting"
    summary_points: Optional[List[SummaryPointRequest]] = None
    conclusion: str = ""

    centre_dept: str = ""
    group_name: str = ""
    doc_no: Optional[str] = ""
    doc_date: Optional[str] = ""
    prepared_by: Optional[str] = ""
    approved_by: Optional[str] = ""
    filename: Optional[str] = "CMTI_Minutes_of_Meeting.docx"


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_margins(cell, top=50, start=50, bottom=50, end=50):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")

    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")

    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_width(cell, width_inches):
    """Set fixed cell width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_row_height(row, height_inches):
    """Set fixed row height."""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_inches * 1440)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def add_text(
    cell,
    text,
    font_size=9,
    bold=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    font_name="Arial",
    color=None
):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if isinstance(color, str):
        run.font.color.rgb = RGBColor.from_string(color)
    elif isinstance(color, RGBColor):
        run.font.color.rgb = color

    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    return run


# ============================================================
# DOCUMENT GENERATION FUNCTION
# ============================================================

def create_mom_document(
    meeting_date_time: str = "",
    meeting_location: str = "",
    prev_mom_no_date: str = "-",
    prev_action_points: str = "-",
    prev_status: str = "-",
    agenda: str = "Project kick off meeting",
    summary_points: Optional[List[SummaryPointRequest]] = None,
    conclusion: str = "Clearance was given for design of fixtures and electrical design.",
    centre_dept: str = "SMPM",
    group_name: str = "SMPM",
    doc_no: str = "037/001",
    doc_date: str = "",
    prepared_by: str = "",
    approved_by: str = ""
) -> Document:
    doc = Document()

    # Page Setup (Margins: 1 inch top/bottom/left/right)
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.different_first_page_header_footer = False

    # Header & Footer Tables
    add_header_table(
        section,
        title="MINUTES OF MEETING",
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=doc_no,
        date_str=doc_date
    )
    add_footer_table(
        section,
        prepared_name=prepared_by,
        approved_name=approved_by,
        group_name=group_name,
        doc_code="037"
    )

    border_fmt = {"val": "single", "sz": 4, "color": "000000"}

    # 1. Meeting Date & Location Table (1 row x 4 cells: Meeting Date and time | val | Meeting Location | val)
    t1 = doc.add_table(rows=1, cols=4)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.autofit = False
    col_w1 = [1.8, 1.8, 1.3, 1.37]
    for cell_idx, width in enumerate(col_w1):
        set_cell_width(t1.rows[0].cells[cell_idx], width)
        set_cell_border(t1.rows[0].cells[cell_idx], top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
        set_cell_margins(t1.rows[0].cells[cell_idx], top=40, bottom=40, start=40, end=40)

    add_text(t1.cell(0, 0), "Meeting Date and time", font_size=9, bold=True)
    add_text(t1.cell(0, 1), meeting_date_time, font_size=9)
    add_text(t1.cell(0, 2), "Meeting Location", font_size=9, bold=True)
    add_text(t1.cell(0, 3), meeting_location, font_size=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 2. Previous Meeting MOM Table (2 rows x 3 cols)
    t2 = doc.add_table(rows=2, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.autofit = False
    col_w2 = [2.2, 2.3, 1.77]

    for row in t2.rows:
        for idx, width in enumerate(col_w2):
            set_cell_width(row.cells[idx], width)
            set_cell_border(row.cells[idx], top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            set_cell_margins(row.cells[idx], top=40, bottom=40, start=40, end=40)

    add_text(t2.cell(0, 0), "Previous Meeting MOM\nNumber with Date", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(t2.cell(0, 1), "Action Points", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(t2.cell(0, 2), "Status\n(Closed/Open with Justification)", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_text(t2.cell(1, 0), prev_mom_no_date or "-", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(t2.cell(1, 1), prev_action_points or "-", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(t2.cell(1, 2), prev_status or "-", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 3. Agenda Table (2 rows x 1 col)
    t3 = doc.add_table(rows=2, cols=1)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3.autofit = False
    for row in t3.rows:
        set_cell_width(row.cells[0], 6.27)
        set_cell_border(row.cells[0], top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
        set_cell_margins(row.cells[0], top=40, bottom=40, start=40, end=40)

    add_text(t3.cell(0, 0), "Agenda", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(t3.cell(1, 0), agenda or "Project kick off meeting", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 4. Summary of Meeting Table
    if not summary_points:
        summary_points = [
            SummaryPointRequest(sl_no=1, points_discussed="", responsibility="")
        ]


    t4 = doc.add_table(rows=2 + len(summary_points), cols=3)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    t4.autofit = False

    # Merge header row 0 (Summary Of the Meeting)
    hdr_cell = t4.cell(0, 0).merge(t4.cell(0, 2))
    set_cell_border(hdr_cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
    set_cell_margins(hdr_cell, top=40, bottom=40, start=40, end=40)
    add_text(hdr_cell, "Summary Of the Meeting", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 1 Column Headers
    col_w4 = [0.8, 3.87, 1.6]
    for idx, width in enumerate(col_w4):
        cell = t4.cell(1, idx)
        set_cell_width(cell, width)
        set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
        set_cell_margins(cell, top=40, bottom=40, start=40, end=40)

    add_text(t4.cell(1, 0), "Sl No", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(t4.cell(1, 1), "Points Discussed", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(t4.cell(1, 2), "Responsibility", font_size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data Rows
    for i, pt in enumerate(summary_points):
        row_idx = 2 + i
        for c_idx, width in enumerate(col_w4):
            cell = t4.cell(row_idx, c_idx)
            set_cell_width(cell, width)
            set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            set_cell_margins(cell, top=40, bottom=40, start=40, end=40)

        add_text(t4.cell(row_idx, 0), str(pt.sl_no), font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        pt_text = f"• {pt.points_discussed}" if pt.points_discussed else "• "
        add_text(t4.cell(row_idx, 1), pt_text, font_size=9)
        add_text(t4.cell(row_idx, 2), pt.responsibility, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 5. Conclusion Table (2 rows x 1 col)
    t5 = doc.add_table(rows=2, cols=1)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    t5.autofit = False
    for row in t5.rows:
        set_cell_width(row.cells[0], 6.27)
        set_cell_border(row.cells[0], top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
        set_cell_margins(row.cells[0], top=40, bottom=40, start=40, end=40)

    add_text(t5.cell(0, 0), "CONCLUSION", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_text(t5.cell(1, 0), conclusion or "Clearance was given for design of fixtures and electrical design.", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


# ============================================================
# FASTAPI ENDPOINT
# ============================================================

@router.post("/mom/generate-word")
def generate_mom_word(req: MomRequest, db: Session = Depends(get_db)):
    doc = create_mom_document(
        meeting_date_time=req.meeting_date_time,
        meeting_location=req.meeting_location,
        prev_mom_no_date=req.prev_mom_no_date,
        prev_action_points=req.prev_action_points,
        prev_status=req.prev_status,
        agenda=req.agenda,
        summary_points=req.summary_points,
        conclusion=req.conclusion,
        centre_dept=req.centre_dept,
        group_name=req.group_name,
        doc_no=req.doc_no or "037/001",
        doc_date=req.doc_date or "",
        prepared_by=req.prepared_by or "",
        approved_by=req.approved_by or ""
    )

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = req.filename if req.filename.endswith(".docx") else f"{req.filename}.docx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
