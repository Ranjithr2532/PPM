"""
FastAPI Router & Document Generator for ISO Document 085: Inspection Report.
Generates Word (.docx) document matching CMTI-QMS-085/Rev00 specification.

Header and footer are intentionally NOT built here — they are imported from
the shared iso.header / iso.finalfooter modules, same pattern as Doc 053.
"""

import io
from typing import List, Optional

from fastapi import APIRouter, status, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from iso.header import add_header_table
from iso.finalfooter import add_footer_table

router = APIRouter(prefix="/iso", tags=["ISO Inspection Report (Doc 085)"])


# ============================================================
# LOW-LEVEL DOCX HELPERS (same conventions as Doc 053 generator)
# ============================================================

def set_cell_shading(cell, color_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ("val", "color", "sz", "space"):
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_margins(cell, top=40, start=60, bottom=40, end=60):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m_name, m_val in [("top", top), ("left", start), ("bottom", bottom), ("right", end)]:
        node = OxmlElement(f"w:{m_name}")
        node.set(qn("w:w"), str(m_val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_width(cell, width_inches: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_table_indent(table, twips: int):
    """Negative indent used to match the original template's table position."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(twips))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)


def add_text(cell_or_paragraph, text: str, font_name: str = "Arial", font_size: int = 10,
             bold: bool = False, italic: bool = False, color: RGBColor = RGBColor(0, 0, 0),
             alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, space_after: int = 0):
    if hasattr(cell_or_paragraph, "paragraphs"):
        p = cell_or_paragraph.paragraphs[0]
    else:
        p = cell_or_paragraph

    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def add_label_value(cell, label: str, value: str, font_size: int = 10):
    """Renders 'Label: value' as two runs — bold label, regular value — in one paragraph."""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    r_label = p.add_run(f"{label}: ")
    r_label.font.name = "Arial"
    r_label.font.size = Pt(font_size)
    r_label.font.bold = True

    if value:
        r_val = p.add_run(str(value))
        r_val.font.name = "Arial"
        r_val.font.size = Pt(font_size)
        r_val.font.bold = False


# ============================================================
# REQUEST MODELS
# ============================================================

class InspectionRowRequest(BaseModel):
    """One row of the measurement table. The frontend's 'Add Row' button
    should append one of these to `rows` for every row the user adds."""
    sl_no: str = ""
    specified_dimensions: str = ""
    drawing_zone: str = ""
    measured_values: str = ""
    instrument_used: str = ""
    remarks: str = ""


class InspectionReportRequest(BaseModel):
    # Info block
    report_no: str = ""
    date: str = ""
    project_no: str = ""
    type: str = ""
    drawing_no: str = ""
    drawing_name: str = ""
    quantity: str = ""

    # Measurement table — dynamic, one entry per row added on the frontend
    rows: Optional[List[InspectionRowRequest]] = None

    # Header / footer passthrough (forwarded to iso.header / iso.finalfooter)
    prepared_by: str = ""
    approved_by: str = ""
    group_name: str = ""
    centre_dept: str = ""
    doc_no: str = "085"
    doc_date: str = ""

    filename: str = "ISO_Inspection_Report.docx"


DEFAULT_ROW_COUNT = 12  # matches the blank template when no rows are supplied

# Twips extracted from the reference template — keeps both tables visually
# aligned exactly as in 085-Inspection_report.docx
INFO_TABLE_COL_TWIPS = [3681, 2557, 4394]
DATA_TABLE_COL_TWIPS = [1276, 2197, 1208, 2262, 1567, 2122]
TABLE_INDENT_TWIPS = -601


# ============================================================
# DOCUMENT BUILDER
# ============================================================

def create_inspection_report_document(
    report_no: str = "",
    date: str = "",
    project_no: str = "",
    type: str = "",
    drawing_no: str = "",
    drawing_name: str = "",
    quantity: str = "",
    rows: Optional[List[dict]] = None,
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "085",
    doc_date: str = "",
) -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.69)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    header_group = (group_name or centre_dept or "SMC").upper()
    if header_group.startswith("G-") or header_group.startswith("C-"):
        header_group = header_group[2:]

    # Standard ISO Header (built elsewhere)
    add_header_table(
        section,
        title=f"INSPECTION REPORT-{header_group}",
        page_str="1 of 1",
        centre_dept=centre_dept,
        doc_no=doc_no or "085",
        date_str=doc_date,
    )

    border_fmt = {"val": "single", "sz": "4", "color": "000000"}

    # ---------- Table 0: Report info block (3 rows x 3 cols) ----------
    info_table = doc.add_table(rows=3, cols=3)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    set_table_indent(info_table, TABLE_INDENT_TWIPS)

    info_col_widths = [w / 1440 for w in INFO_TABLE_COL_TWIPS]
    for row in info_table.rows:
        for c_idx, cell in enumerate(row.cells):
            set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_width(cell, info_col_widths[c_idx])

    # Row 0: Report No (merged col 0-1) | Date
    c_report = info_table.cell(0, 0).merge(info_table.cell(0, 1))
    add_label_value(c_report, "Report No", report_no)
    add_label_value(info_table.cell(0, 2), "Date", date)

    # Row 1: Project No | Type | Drawing No
    add_label_value(info_table.cell(1, 0), "Project No", project_no)
    add_label_value(info_table.cell(1, 1), "Type", type)
    add_label_value(info_table.cell(1, 2), "Drawing No", drawing_no)

    # Row 2: Drawing Name (merged col 0-1) | Quantity
    c_drawing = info_table.cell(2, 0).merge(info_table.cell(2, 1))
    add_label_value(c_drawing, "Drawing Name", drawing_name)
    add_label_value(info_table.cell(2, 2), "Quantity", quantity)

    # ---------- Note paragraph ----------
    note_p = doc.add_paragraph()
    add_text(note_p, "(All dimensions in mm unless stated specifically)",
             font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # ---------- Table 1: Measurement table (dynamic rows) ----------
    row_list = rows if rows else [{} for _ in range(DEFAULT_ROW_COUNT)]
    total_rows = 1 + len(row_list)  # +1 for the header row

    data_table = doc.add_table(rows=total_rows, cols=6)
    data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data_table.autofit = False
    set_table_indent(data_table, TABLE_INDENT_TWIPS)

    data_col_widths = [w / 1440 for w in DATA_TABLE_COL_TWIPS]
    for row in data_table.rows:
        for c_idx, cell in enumerate(row.cells):
            set_cell_border(cell, top=border_fmt, bottom=border_fmt, left=border_fmt, right=border_fmt)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_width(cell, data_col_widths[c_idx])

    headers = ["Sl No", "Specified Dimensions", "Drawing Zone",
               "Measured Values", "Instrument used (with Serial No)", "Remarks"]
    for c_idx, h in enumerate(headers):
        add_text(data_table.cell(0, c_idx), h, font_size=10, bold=True,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, item in enumerate(row_list):
        r_idx = 1 + idx
        if isinstance(item, dict):
            sl_no = item.get("sl_no") or str(idx + 1)
            specified_dimensions = item.get("specified_dimensions", "")
            drawing_zone = item.get("drawing_zone", "")
            measured_values = item.get("measured_values", "")
            instrument_used = item.get("instrument_used", "")
            remarks = item.get("remarks", "")
        else:
            sl_no = getattr(item, "sl_no", "") or str(idx + 1)
            specified_dimensions = getattr(item, "specified_dimensions", "")
            drawing_zone = getattr(item, "drawing_zone", "")
            measured_values = getattr(item, "measured_values", "")
            instrument_used = getattr(item, "instrument_used", "")
            remarks = getattr(item, "remarks", "")

        row_values = [sl_no, specified_dimensions, drawing_zone,
                      measured_values, instrument_used, remarks]
        for c_idx, val in enumerate(row_values):
            add_text(data_table.cell(r_idx, c_idx), val, font_size=10,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Standard ISO Footer / signature block (built elsewhere)
    doc_code_footer = f"CMTI-{header_group}-QMS-{doc_no or '085'}/Rev00"
    add_footer_table(
        doc,
        prepared_name=prepared_by,
        approved_name=approved_by,
        group_name=header_group,
        doc_code=doc_code_footer,
        in_body=True,
    )

    return doc


def generate_inspection_report_bytes(
    report_no: str = "",
    date: str = "",
    project_no: str = "",
    type: str = "",
    drawing_no: str = "",
    drawing_name: str = "",
    quantity: str = "",
    rows: Optional[List[InspectionRowRequest]] = None,
    prepared_by: str = "",
    approved_by: str = "",
    group_name: str = "",
    centre_dept: str = "",
    doc_no: str = "085",
    doc_date: str = "",
) -> io.BytesIO:
    row_dicts = [r.dict() if hasattr(r, "dict") else r for r in rows] if rows else None

    doc = create_inspection_report_document(
        report_no=report_no,
        date=date,
        project_no=project_no,
        type=type,
        drawing_no=drawing_no,
        drawing_name=drawing_name,
        quantity=quantity,
        rows=row_dicts,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date,
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# API ROUTE HANDLERS
# ============================================================

@router.get(
    "/inspection-report/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Inspection Report (.docx) via GET",
)
async def generate_inspection_report_doc_get(
    report_no: str = Query(""),
    date: str = Query(""),
    project_no: str = Query(""),
    type: str = Query(""),
    drawing_no: str = Query(""),
    drawing_name: str = Query(""),
    quantity: str = Query(""),
    prepared_by: str = Query(""),
    approved_by: str = Query(""),
    group_name: str = Query(""),
    centre_dept: str = Query(""),
    doc_no: str = Query("085"),
    doc_date: str = Query(""),
    filename: str = Query("ISO_Inspection_Report.docx"),
):
    # GET has no body, so rows aren't supported here — use POST for dynamic rows.
    buffer = generate_inspection_report_bytes(
        report_no=report_no,
        date=date,
        project_no=project_no,
        type=type,
        drawing_no=drawing_no,
        drawing_name=drawing_name,
        quantity=quantity,
        prepared_by=prepared_by,
        approved_by=approved_by,
        group_name=group_name,
        centre_dept=centre_dept,
        doc_no=doc_no,
        doc_date=doc_date,
    )

    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.post(
    "/inspection-report/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Inspection Report (.docx) via POST",
)
async def generate_inspection_report_doc_post(payload: InspectionReportRequest):
    filename = payload.filename or "ISO_Inspection_Report.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_inspection_report_bytes(
        report_no=payload.report_no,
        date=payload.date,
        project_no=payload.project_no,
        type=payload.type,
        drawing_no=payload.drawing_no,
        drawing_name=payload.drawing_name,
        quantity=payload.quantity,
        rows=payload.rows,
        prepared_by=payload.prepared_by,
        approved_by=payload.approved_by,
        group_name=payload.group_name,
        centre_dept=payload.centre_dept,
        doc_no=payload.doc_no,
        doc_date=payload.doc_date,
    )

    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )