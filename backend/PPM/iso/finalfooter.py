import os
import io
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel
from fastapi import APIRouter, status, Query
from fastapi.responses import StreamingResponse

# ============================================================
# FASTAPI ROUTER
# ============================================================

router = APIRouter(prefix="/iso", tags=["ISO Footer Generator"])


class FooterInfoRequest(BaseModel):
    prepared_name: str = ""
    approved_name: str = ""
    group_name: str = ""
    filename: Optional[str] = "CMTI_Project_Team_Footer.docx"


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_margins(cell, top=50, start=50, bottom=50, end=50):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")

    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")

    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_width(cell, width_inches):
    """Set fixed cell width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_row_height(row, height_inches):
    """Set fixed row height."""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_inches * 1440)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def add_text(
    cell,
    text,
    font_size=10,
    bold=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    font_name="Arial",
    color=None
):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if isinstance(color, str):
        run.font.color.rgb = RGBColor.from_string(color)
    elif isinstance(color, RGBColor):
        run.font.color.rgb = color

    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    return run


# ============================================================
# DOCUMENT GENERATION FUNCTIONS
# ============================================================

def build_signature_table_element(
    container,
    prepared_name: str = "",
    approved_name: str = "",
    group_name: str = "",
    doc_code: str = "072"
):
    # Main Footer Table (3 rows x 4 cols)
    if hasattr(container, "is_linked_to_previous") or container.__class__.__name__ in ("_Footer", "_Header"):
        table = container.add_table(3, 4, Inches(6.27))
    else:
        table = container.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Column widths based on layout analysis (sum = 6.27 in)
    col_widths = [1.10, 2.00, 1.10, 2.07]

    # Set row heights
    row_heights = [0.25, 0.45, 0.45]
    for r_idx in range(3):
        row = table.rows[r_idx]
        set_row_height(row, row_heights[r_idx])
        for c_idx in range(4):
            set_cell_width(row.cells[c_idx], col_widths[c_idx])

    # Apply borders & vertical alignment & padding
    border_format = {"val": "single", "sz": 4, "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=border_format,
                bottom=border_format,
                left=border_format,
                right=border_format
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=25, start=45, bottom=25, end=45)

    # Merge cells
    # Prepared by: Column 0, Row 0 to 1
    prepared_cell = table.cell(0, 0).merge(table.cell(1, 0))
    # Approved by: Column 2, Row 0 to 1
    approved_cell = table.cell(0, 2).merge(table.cell(1, 2))
    # Institute Name/Address: Row 2, Column 0 to 3
    inst_cell = table.cell(2, 0).merge(table.cell(2, 3))

    # Add Prepared by text
    add_text(prepared_cell, "Prepared by", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Add Approved By text
    add_text(approved_cell, "Approved By", font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Prepared Name Cell (Row 0, Col 1)
    p_prep_name = table.cell(0, 1).paragraphs[0]
    p_prep_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_prep_name.paragraph_format.space_before = Pt(0)
    p_prep_name.paragraph_format.space_after = Pt(0)
    r_prep_lbl = p_prep_name.add_run("Name: ")
    r_prep_lbl.font.name = "Arial"
    r_prep_lbl.font.size = Pt(10)
    r_prep_lbl._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r_prep_lbl._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if prepared_name:
        r_prep_val = p_prep_name.add_run(prepared_name)
        r_prep_val.bold = True
        r_prep_val.font.name = "Arial"
        r_prep_val.font.size = Pt(10)
        r_prep_val._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        r_prep_val._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

    # Prepared Signature Cell (Row 1, Col 1)
    add_text(table.cell(1, 1), "Signature: ", font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Approved Name Cell (Row 0, Col 3)
    p_app_name = table.cell(0, 3).paragraphs[0]
    p_app_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_app_name.paragraph_format.space_before = Pt(0)
    p_app_name.paragraph_format.space_after = Pt(0)
    r_app_lbl = p_app_name.add_run("Name: ")
    r_app_lbl.font.name = "Arial"
    r_app_lbl.font.size = Pt(10)
    r_app_lbl._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r_app_lbl._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if approved_name:
        r_app_val = p_app_name.add_run(approved_name)
        r_app_val.bold = True
        r_app_val.font.name = "Arial"
        r_app_val.font.size = Pt(10)
        r_app_val._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        r_app_val._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

    # Approved Signature Cell (Row 1, Col 3)
    add_text(table.cell(1, 3), "Signature: ", font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Bottom Row: Address & Institute
    p_inst = inst_cell.paragraphs[0]
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_before = Pt(0)
    p_inst.paragraph_format.space_after = Pt(0)
    r_inst1 = p_inst.add_run("CENTRAL MANUFACTURING TECHNOLOGY INSTITUTE")
    r_inst1.bold = True
    r_inst1.font.name = "Arial"
    r_inst1.font.size = Pt(10)
    r_inst1._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r_inst1._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

    p_addr = inst_cell.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_addr.paragraph_format.space_before = Pt(0)
    p_addr.paragraph_format.space_after = Pt(0)
    r_inst2 = p_addr.add_run("TUMKUR ROAD, BANGALORE 560 022")
    r_inst2.bold = True
    r_inst2.font.name = "Arial"
    r_inst2.font.size = Pt(10)
    r_inst2._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r_inst2._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

    return table


def add_footer_table(
    target,
    prepared_name: str = "",
    approved_name: str = "",
    group_name: str = "",
    doc_code: str = "072",
    in_body: bool = False
):
    """
    Adds ISO footer / signature block.
    - If in_body=True or target is a Document object, the Prepared by / Approved by table is added 
      directly into the Document Body (so it appears ONLY on the LAST PAGE).
    - In addition, section.footer gets a clean ISO revision line on every page.
    """
    section = None
    container = None

    if hasattr(target, "sections"):
        # target is a Document object
        section = target.sections[0]
        container = target
        in_body = True
    elif hasattr(target, "footer"):
        # target is a Section object
        section = target
        container = target.footer if not in_body else target._document if hasattr(target, "_document") else target.footer
    else:
        container = target

    # 1. Resolve exact revision text (always full CMTI code)
    code_candidate = (doc_code or "").strip()
    group_candidate = (group_name or "").strip()

    if code_candidate.upper().startswith("CMTI"):
        rev_text = code_candidate
    elif group_candidate.upper().startswith("CMTI"):
        rev_text = group_candidate
    elif code_candidate and "/" in code_candidate:
        rev_text = code_candidate
    else:
        group_str = group_candidate.upper() if group_candidate else ""
        if group_str.startswith("C-") or group_str.startswith("G-"):
            group_str = group_str[2:]
        
        group_part = f"-{group_str}" if group_str else ""
        code_str = f"-{code_candidate}" if code_candidate else ""
        rev_text = f"CMTI-QMS{group_part}{code_str}/Rev00"

    if not in_body and section:
        footer = section.footer

        # Step 1: Add the signature table FIRST into the footer
        build_signature_table_element(
            container=footer,
            prepared_name=prepared_name,
            approved_name=approved_name,
            group_name=group_name,
            doc_code=doc_code
        )

        # Step 2: Add the ISO revision text paragraph AFTER the table
        rev_para = footer.add_paragraph()
        rev_para.paragraph_format.space_before = Pt(4)
        rev_para.paragraph_format.space_after = Pt(0)
        rev_para.paragraph_format.line_spacing = 1.0
        rev_para.text = f"ISO 9001-2015                                                                                     {rev_text}"
        rev_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if rev_para.runs:
            run_rev = rev_para.runs[0]
            run_rev.font.name = "Arial"
            run_rev.font.size = Pt(8.5)

        # Step 3: Remove the initial default empty paragraph that was before the table
        if footer.paragraphs and len(footer.paragraphs) > 1:
            first_p = footer.paragraphs[0]
            if not first_p.text.strip():
                try:
                    first_p._element.getparent().remove(first_p._element)
                except Exception:
                    pass

    elif container:
        # If in body (e.g. multi-page proposals/plans)
        build_signature_table_element(
            container=container,
            prepared_name=prepared_name,
            approved_name=approved_name,
            group_name=group_name,
            doc_code=doc_code
        )

        if hasattr(container, "add_paragraph"):
            rev_para = container.add_paragraph()
            rev_para.paragraph_format.space_before = Pt(4)
            rev_para.paragraph_format.space_after = Pt(0)
            rev_para.paragraph_format.line_spacing = 1.0
            rev_para.text = f"ISO 9001-2015                                                                                     {rev_text}"
            rev_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if rev_para.runs:
                run_rev = rev_para.runs[0]
                run_rev.font.name = "Arial"
                run_rev.font.size = Pt(8.5)


def create_iso_footer_document(
    prepared_name: str = "",
    approved_name: str = "",
    group_name: str = "",
    doc_code: str = ""
) -> Document:
    doc = Document()
    section = doc.sections[0]

    # A4 page dimensions
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # Margins
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    add_footer_table(
        section,
        prepared_name=prepared_name,
        approved_name=approved_name,
        group_name=group_name,
        doc_code=doc_code
    )

    return doc


def generate_iso_footer_bytes(
    prepared_name: str = "",
    approved_name: str = "",
    group_name: str = "",
    doc_code: str = "072"
) -> io.BytesIO:
    doc = create_iso_footer_document(
        prepared_name=prepared_name,
        approved_name=approved_name,
        group_name=group_name,
        doc_code=doc_code
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# API ENDPOINTS
# ============================================================

@router.get(
    "/footer/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Footer Word Document (.docx)"
)
async def generate_footer_doc_get(
    prepared_name: str = Query("", description="Prepared by name"),
    approved_name: str = Query("", description="Approved by name"),
    group_name: str = Query("", description="Group ISO code name"),
    doc_code: str = Query("072", description="Document reference code"),
    filename: str = Query("CMTI_Project_Team_Footer.docx", description="Output filename")
):
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_iso_footer_bytes(
        prepared_name=prepared_name,
        approved_name=approved_name,
        group_name=group_name,
        doc_code=doc_code
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.post(
    "/footer/generate",
    status_code=status.HTTP_200_OK,
    summary="Generate ISO Footer Word Document (.docx) via POST"
)
async def generate_footer_doc_post(payload: FooterInfoRequest):
    filename = payload.filename or "CMTI_Project_Team_Footer.docx"
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    buffer = generate_iso_footer_bytes(
        prepared_name=payload.prepared_name,
        approved_name=payload.approved_name,
        group_name=payload.group_name
    )

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


# ============================================================
# MAIN SCRIPT EXECUTION (Standalone Test)
# ============================================================

if __name__ == "__main__":
    doc = create_iso_footer_document(
        prepared_name="Ranjith Kumar",
        approved_name="SMC Director"
    )
    output_filename = "CMTI_Project_Team_Footer.docx"
    doc.save(output_filename)
    print(f"Document created successfully: {output_filename}")
